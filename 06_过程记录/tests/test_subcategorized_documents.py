"""Contract tests for the plain-language subcategory DOCX pipeline."""

from __future__ import annotations

import hashlib
import importlib.util
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[2]
TOOLS_DIR = PROJECT_ROOT / "06_过程记录" / "tools"
BUILDER_PATH = TOOLS_DIR / "build_subcategorized_documents.py"
VERIFIER_PATH = TOOLS_DIR / "verify_subcategorized_documents.py"
RENDERER_PATH = TOOLS_DIR / "render_subcategorized_documents.py"


def load_module(path: Path, name: str):
    """Load a task module only after asserting that the planned artifact exists."""
    if not path.exists():
        raise AssertionError(f"缺少计划脚本: {path.name}")
    spec = importlib.util.spec_from_file_location(name, path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def fixture_record(skill_id: str, *, subcategory_code: str = "05-05", suffix: str = "") -> dict:
    number = int(skill_id.rsplit("-", 1)[1])
    long_note = (
        "先在一份可公开的样例任务上记录基线耗时，再逐项调整处理器数量、显存占用和批次大小；"
        "每次只改变一个条件，并把结果写进同一份比较记录。"
    )
    return {
        "id": skill_id,
        "name": f"parallel-tool-{number}{suffix}",
        "cn": f"并行计算工具 {number}",
        "cat": "05",
        "repo": f"example/repository-{number}",
        "path": f"skills/parallel-tool-{number}/SKILL.md",
        "ecosystem": "Agent Skills（兼容 Codex）",
        "form": "社区 skill、开源仓库",
        "tags": "并行计算、性能分析",
        "summary": "使用多个处理器缩短计算时间。",
        "detail": "按任务特点分配计算资源，并保留可比较的运行记录。",
        "roles": f"适合科研人员、数据分析人员 {number} 使用。",
        "scenario": "批量模拟、模型训练、数据转换",
        "compat": "A",
        "adapt": "先用小样例核对结果，再逐步增加工作量。",
        "deps": "可公开的样例数据、可用的处理器或显卡、基线运行记录",
        "risk": "并行任务可能增加内存占用；结果仍需与单任务版本比较。",
        "verify": "二级包内容验证",
        "priority": "高" if number % 2 else "中",
        "related": "",
        "repo_url": f"https://github.com/example/repository-{number}?tab=readme-ov-file&source=very-long-address-{number}",
        "skill_url": f"https://github.com/example/repository-{number}/blob/main/skills/parallel-tool-{number}/SKILL.md?plain=1&source=very-long-address-{number}",
        "stars": number * 10,
        "repo_pushed": "2026-08-03",
        "license": "MIT",
        "subcategory_code": subcategory_code,
        "subcategory_name": {
            "05-01": "计算设备与资源规划",
            "05-05": "并行计算与性能优化",
        }.get(subcategory_code, "并行计算与性能优化"),
        "plain_purpose": f"把可拆分的计算任务分给多个处理器，缩短等待时间 {number}。",
        "plain_outputs": f"可得到性能比较报告、资源使用记录和调整建议 {number}。",
        "plain_audience": f"适合科研人员、数据分析人员 {number} 使用。",
        "plain_when_to_use": f"当单次计算等待时间较长、任务能够拆分时值得使用 {number}。",
        "plain_prerequisites": f"需要准备：可公开的样例数据、基线耗时和可用计算资源 {number}。",
        "plain_limitations": f"需要注意：{long_note}{number}。",
        "plain_integration": f"需要中等调整：先接入一条任务，再比较结果一致性和资源占用 {number}。",
        "plain_verification": "包内容已核验：已核对用途说明和包内文件；本次未安装、未运行，不能据此判断实际效果。",
    }


OVERVIEW = {
    "code": "05",
    "name": "编程、数学、数据分析和可视化",
    "subcategories": [
        {
            "code": "05-01",
            "name": "计算设备与资源规划",
            "inclusion_focus": "检查处理器、显卡、内存和磁盘，并据此安排计算任务",
        },
        {
            "code": "05-05",
            "name": "并行计算与性能优化",
            "inclusion_focus": "使用多个处理器或显卡加快计算，并查找性能瓶颈",
        },
    ],
}

SUBCATEGORY = {
    "code": "05-05",
    "name": "并行计算与性能优化",
    "inclusion_focus": "使用多个处理器或显卡加快计算，并查找性能瓶颈",
    "big_category_name": "编程、数学、数据分析和可视化",
}


def document_xml(document: Document) -> str:
    with tempfile.TemporaryDirectory() as temporary_directory:
        path = Path(temporary_directory) / "artifact.docx"
        document.save(path)
        with zipfile.ZipFile(path) as package:
            return package.read("word/document.xml").decode("utf-8")


def all_document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        parts.extend(cell.text for row in table.rows for cell in row.cells)
    return "\n".join(parts)


def fixture_knowledge_base_markdown(category: dict, rows: list[dict]) -> str:
    lines = [
        f"# {category['code']} {category['name']}",
        "",
        "## 这类工具是做什么的",
        "",
        category["inclusion_focus"],
        "",
        "## 收录数量",
        "",
        f"共 {len(rows)} 项 Skill。",
        "",
        "## 收录条目",
        "",
        "| 内部编号 | 中文名称 | 主要用途 | 推荐程度 | 条目链接 |",
        "|---|---|---|---|---|",
    ]
    for row in rows:
        lines.append(
            f"| {row['id']} | {row['cn']} | {row['plain_purpose']} | {row['priority']} | "
            f"[查看](<../../skills/{row['id']}_{row['name']}.md>) |"
        )
    return "\n".join(lines) + "\n"


def write_fixture_knowledge_base(root: Path, category: dict, rows: list[dict]) -> Path:
    path = (
        root
        / "02_知识库/functional_domains/05_编程数学数据分析和可视化/subcategories"
        / f"{category['code']}_{category['name']}"
        / "INDEX.md"
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(fixture_knowledge_base_markdown(category, rows), encoding="utf-8")
    skills = path.parent.parent.parent / "skills"
    skills.mkdir(parents=True, exist_ok=True)
    for row in rows:
        (skills / f"{row['id']}_{row['name']}.md").write_text("# fixture\n", encoding="utf-8")
    return path


class PlannedScriptsTests(unittest.TestCase):
    def test_all_three_planned_scripts_exist(self):
        """Removing any pipeline stage must make the documented workflow unusable."""
        self.assertTrue(BUILDER_PATH.exists(), f"缺少 {BUILDER_PATH.name}")
        self.assertTrue(VERIFIER_PATH.exists(), f"缺少 {VERIFIER_PATH.name}")
        self.assertTrue(RENDERER_PATH.exists(), f"缺少 {RENDERER_PATH.name}")


class DocumentLayoutTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.builder = load_module(BUILDER_PATH, "build_subcategorized_documents_for_test")
        cls.records = [fixture_record("GH-05-0003"), fixture_record("GH-05-0024")]
        cls.document = cls.builder.build_subcategory_document(SUBCATEGORY, cls.records)

    def test_letter_geometry_and_compact_reference_styles_are_encoded(self):
        """A default Word template or style drift must fail the preset contract."""
        section = self.document.sections[0]
        self.assertEqual((section.page_width.twips, section.page_height.twips), (12240, 15840))
        self.assertEqual(
            tuple(value.twips for value in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin)),
            (1440, 1440, 1440, 1440),
        )
        normal = self.document.styles["Normal"]
        self.assertEqual(normal.font.name, "Calibri")
        self.assertEqual(normal._element.rPr.rFonts.get(qn("w:eastAsia")), "Microsoft YaHei")
        self.assertEqual(normal.font.size.pt, 11)
        self.assertEqual(normal.paragraph_format.space_after.pt, 6)
        self.assertEqual(normal.paragraph_format.line_spacing, 1.25)

        expected = {
            "Heading 1": (16, "2E74B5", 18, 10),
            "Heading 2": (13, "2E74B5", 14, 7),
            "Heading 3": (12, "1F4D78", 10, 5),
        }
        for style_name, (size, color, before, after) in expected.items():
            style = self.document.styles[style_name]
            with self.subTest(style=style_name):
                self.assertEqual(style.font.name, "Calibri")
                self.assertEqual(style._element.rPr.rFonts.get(qn("w:eastAsia")), "Microsoft YaHei")
                self.assertEqual(style.font.size.pt, size)
                self.assertEqual(str(style.font.color.rgb), color)
                self.assertEqual(style.paragraph_format.space_before.pt, before)
                self.assertEqual(style.paragraph_format.space_after.pt, after)

    def test_every_table_has_fixed_full_width_geometry_and_indent(self):
        """Autofit, percentage widths, or mismatched cells must fail XML inspection."""
        self.assertGreater(len(self.document.tables), 0)
        for table in self.document.tables:
            table_width = table._tbl.tblPr.find(qn("w:tblW"))
            table_indent = table._tbl.tblPr.find(qn("w:tblInd"))
            layout = table._tbl.tblPr.find(qn("w:tblLayout"))
            grid_widths = [int(node.get(qn("w:w"))) for node in table._tbl.tblGrid]
            with self.subTest(table=table):
                self.assertEqual((table_width.get(qn("w:type")), int(table_width.get(qn("w:w")))), ("dxa", 9360))
                self.assertEqual((table_indent.get(qn("w:type")), int(table_indent.get(qn("w:w")))), ("dxa", 120))
                self.assertEqual(layout.get(qn("w:type")), "fixed")
                self.assertEqual(sum(grid_widths), 9360)
                for row in table.rows:
                    self.assertIsNone(row._tr.get_or_add_trPr().find(qn("w:trHeight")))
                    for index, cell in enumerate(row.cells):
                        self.assertEqual(int(cell._tc.get_or_add_tcPr().tcW.get(qn("w:w"))), grid_widths[index])

    def test_editorial_cover_header_footer_and_page_number_are_present(self):
        """Losing the cover metadata, quiet header, or real PAGE field must fail."""
        cover_text = "\n".join(paragraph.text for paragraph in self.document.paragraphs[:12])
        self.assertIn("并行计算与性能优化", cover_text)
        self.assertIn("通俗版", cover_text)
        self.assertIn("2026-08-07", cover_text)
        self.assertIn("收录 2 项", cover_text)
        self.assertIn("未安装、未运行", cover_text)
        header_text = "".join(paragraph.text for paragraph in self.document.sections[0].header.paragraphs)
        self.assertEqual(header_text.strip(), "并行计算与性能优化")
        footer_xml = self.document.sections[0].footer._element.xml
        self.assertIn("PAGE", footer_xml)

    def test_all_user_facing_body_and_table_runs_are_exactly_11pt(self):
        """Body, summary tables, links, and technical trace may not be shrunk to fit pages."""
        body_started = False
        for paragraph in self.document.paragraphs:
            if paragraph.style.name == "Heading 1":
                body_started = True
            if not body_started or paragraph.style.name.startswith("Heading"):
                continue
            for run in paragraph._p.xpath(".//w:r[w:t]"):
                size = run.find(qn("w:rPr"))
                size = None if size is None else size.find(qn("w:sz"))
                if size is not None:
                    self.assertEqual(size.get(qn("w:val")), "22", paragraph.text)
        for table in self.document.tables:
            for run in table._tbl.xpath(".//w:r[w:t]"):
                size = run.find(qn("w:rPr"))
                size = None if size is None else size.find(qn("w:sz"))
                self.assertIsNotNone(size)
                self.assertEqual(size.get(qn("w:val")), "22")


class DocumentContentTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.builder = load_module(BUILDER_PATH, "build_subcategorized_documents_content_test")

    def test_overview_uses_plain_fields_in_the_required_section_order(self):
        """Reverting to technical source prose or rearranging reader questions must fail."""
        records = [
            fixture_record("GH-05-0001", subcategory_code="05-01"),
            fixture_record("GH-05-0003"),
            fixture_record("GH-05-0024"),
        ]
        document = self.builder.build_overview_document(OVERVIEW, records)
        headings = [paragraph.text for paragraph in document.paragraphs if paragraph.style.name == "Heading 1"]
        self.assertEqual(
            headings,
            [
                "这类工具解决什么问题",
                "适合哪些人",
                "小分类导航与数量",
                "怎样选择",
                "共同使用条件",
                "共同限制",
                "本次核验到哪一步",
                "来源说明",
            ],
        )
        text = all_document_text(document)
        for field in (
            "plain_purpose",
            "plain_audience",
            "plain_when_to_use",
            "plain_prerequisites",
            "plain_limitations",
            "plain_verification",
        ):
            self.assertIn(records[0][field], text, field)
        navigation = document.tables[0]
        self.assertEqual([cell.text for cell in navigation.rows[0].cells], ["代码与名称", "通俗定义", "数量"])
        self.assertIn("05-05 并行计算与性能优化", navigation.rows[2].cells[0].text)
        self.assertEqual(navigation.rows[2].cells[2].text, "2")
        source_link_paragraphs = [
            paragraph
            for paragraph in document.paragraphs
            if "w:hyperlink" in paragraph._p.xml
        ]
        self.assertEqual(
            source_link_paragraphs,
            [],
            "来源说明不得重复导航表已经逐行给出的仓库链接",
        )
        self.assertIn("仓库已随导航表逐行列出", text)
        navigation_link_counts = [
            sum(len(paragraph._p.findall(qn("w:hyperlink"))) for paragraph in row.cells[0].paragraphs)
            for row in navigation.rows[1:]
        ]
        self.assertEqual(navigation_link_counts, [2, 3])

    def test_overview_caps_limit_examples_to_keep_the_source_section_with_content(self):
        """A long overview must compile examples instead of creating a sparse source-only tail page."""
        records = [fixture_record(f"GH-05-{number:04d}") for number in range(1, 8)]
        document = self.builder.build_overview_document(OVERVIEW, records)
        self.assertEqual(all_document_text(document).count("注意事项"), 4)

    def test_subcategory_has_one_plain_reader_block_and_separate_trace_block_per_skill(self):
        """Omitting a user field, mixing trace facts into advice, or losing an ID must fail."""
        records = [fixture_record("GH-05-0003"), fixture_record("GH-05-0024")]
        document = self.builder.build_subcategory_document(SUBCATEGORY, records)
        text = all_document_text(document)
        skill_headings = [paragraph.text for paragraph in document.paragraphs if paragraph.style.name == "Heading 2"]
        self.assertEqual(skill_headings, ["并行计算工具 3", "并行计算工具 24"])
        for record in records:
            for field in (
                "plain_purpose", "plain_audience", "plain_when_to_use", "plain_prerequisites",
                "plain_limitations", "plain_integration", "plain_verification", "plain_outputs",
            ):
                self.assertIn(record[field], text, f"{record['id']}:{field}")
        self.assertEqual(
            [paragraph.text for paragraph in document.paragraphs if paragraph.style.name == "Heading 3"],
            ["技术追溯", "技术追溯"],
        )
        self.assertEqual(text.count("内部编号："), len(records))
        trace_paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.startswith("内部编号：")]
        self.assertEqual(len(trace_paragraphs), len(records))
        for paragraph in trace_paragraphs:
            self.assertEqual(paragraph.paragraph_format.space_after.pt, 0)
            paragraph_text = paragraph.text
            for label in ("功能标签：", "原生生态：", "来源形态：", "许可证：", "仓库最近更新："):
                self.assertIn(label, paragraph_text)
        for record in records:
            self.assertEqual(text.count(record["id"]), 1)

        hyperlink_paragraphs = [paragraph for paragraph in document.paragraphs if "w:hyperlink" in paragraph._p.xml]
        self.assertEqual(len(hyperlink_paragraphs), len(records))
        for paragraph in hyperlink_paragraphs:
            self.assertTrue(paragraph.text.startswith("内部编号："))
            self.assertIn("原始资料地址：", paragraph.text)
            self.assertEqual(len(paragraph._p.findall(qn("w:hyperlink"))), 2)

        self.assertEqual(len(document.tables), len(records))
        for table, record in zip(document.tables, records):
            labels = [row.cells[0].text for row in table.rows]
            values = [row.cells[1].text for row in table.rows]
            self.assertEqual(labels, ["英文名称", "推荐程度", "接入难度", "核验层级"])
            self.assertEqual(values[0], record["name"])
            self.assertEqual(values[1], record["priority"])
            self.assertEqual(values[3], record["verify"])
            for row in table.rows[:-1]:
                for cell in row.cells:
                    self.assertTrue(
                        cell.paragraphs[0].paragraph_format.keep_with_next,
                        "四行摘要表必须整块留在同一页，不能只把后两行挤到下一页",
                    )

    def test_short_value_table_does_not_add_an_empty_spacer_paragraph(self):
        """Repeated blank paragraphs waste enough height to create orphaned trace-only pages."""
        document = Document()
        self.builder._add_short_value_table(document, fixture_record("GH-05-0003"))
        self.assertEqual(document.paragraphs, [])

    def test_long_text_and_urls_remain_complete_without_full_url_display(self):
        """Truncation or dumping long addresses into tables must fail."""
        record = fixture_record("GH-05-0003")
        record["plain_limitations"] += "补充说明：" + "这是用于检查长文本分页和换行的完整句子。" * 30
        document = self.builder.build_subcategory_document(SUBCATEGORY, [record])
        text = all_document_text(document)
        self.assertIn(record["plain_limitations"], text)
        self.assertNotIn(record["repo_url"], text)
        self.assertNotIn(record["skill_url"], text)
        xml = document_xml(document)
        self.assertGreaterEqual(xml.count("w:hyperlink"), 4)
        self.assertIn("原始资料地址", text)
        for heading in [p for p in document.paragraphs if p.style.name in {"Heading 1", "Heading 2", "Heading 3"}]:
            self.assertTrue(heading.paragraph_format.keep_with_next or heading.style.paragraph_format.keep_with_next)

    def test_empty_subcategory_is_rejected_instead_of_emitting_a_hollow_report(self):
        """An empty taxonomy leaf must stop generation before a misleading DOCX exists."""
        with self.assertRaisesRegex(ValueError, "空小分类"):
            self.builder.build_subcategory_document(SUBCATEGORY, [])


class VerificationAndGenerationTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.builder = load_module(BUILDER_PATH, "build_subcategorized_documents_generation_test")
        cls.verifier = load_module(VERIFIER_PATH, "verify_subcategorized_documents_for_test")
        cls.renderer = load_module(RENDERER_PATH, "render_subcategorized_documents_for_test")

    def test_verifier_accepts_a_valid_document_and_rejects_structural_drift(self):
        """Bad margins, stale IDs, and placeholders must be reported independently."""
        records = [fixture_record("GH-05-0003"), fixture_record("GH-05-0024")]
        with tempfile.TemporaryDirectory() as temporary_directory:
            good_path = Path(temporary_directory) / "good.docx"
            self.builder.build_subcategory_document(SUBCATEGORY, records).save(good_path)
            self.assertEqual(
                self.verifier.verify_document(good_path, scope="subcategory", expected_records=records),
                [],
            )

            bad_document = Document(good_path)
            bad_document.sections[0].left_margin = 720
            bad_document.add_paragraph("TODO 已经运行成功 GH-05-9999")
            bad_path = Path(temporary_directory) / "bad.docx"
            bad_document.save(bad_path)
            issues = self.verifier.verify_document(bad_path, scope="subcategory", expected_records=records)
            self.assertTrue(any("左边距" in issue for issue in issues), issues)
            self.assertTrue(any("占位" in issue for issue in issues), issues)
            self.assertTrue(any("夸大" in issue for issue in issues), issues)
            self.assertTrue(any("Skill ID" in issue for issue in issues), issues)

    def test_verifier_rejects_direct_body_and_table_font_shrinking(self):
        """Direct 8.5/9.5 pt formatting must fail even when Normal remains 11 pt."""
        records = [fixture_record("GH-05-0003")]
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            document = self.builder.build_subcategory_document(SUBCATEGORY, records)
            trace = next(p for p in document.paragraphs if p.text.startswith("内部编号："))
            trace.runs[0].font.size = Pt(8.5)
            trace_path = root / "bad-trace.docx"
            document.save(trace_path)
            trace_issues = self.verifier.verify_document(
                trace_path, scope="subcategory", expected_records=records
            )
            self.assertTrue(any("正文直接字号" in issue for issue in trace_issues), trace_issues)

            document = self.builder.build_subcategory_document(SUBCATEGORY, records)
            table_run = next(
                run for run in document.tables[0].cell(0, 0).paragraphs[0].runs if run.text
            )
            table_run.font.size = Pt(9.5)
            table_path = root / "bad-table.docx"
            document.save(table_path)
            table_issues = self.verifier.verify_document(
                table_path, scope="subcategory", expected_records=records
            )
            self.assertTrue(any("表格直接字号" in issue for issue in table_issues), table_issues)

    def test_manifest_selection_is_sorted_safe_and_scope_aware(self):
        """Unsafe paths, unstable ordering, or ambiguous selectors must fail before writes."""
        manifest = [
            {"path": "05_交付物/通俗细分版_2026-08-07/05_编程数学数据分析和可视化/05-05_并行计算与性能优化/05-05_并行计算与性能优化_GitHub技能调研.docx", "format": "docx", "scope": "subcategory", "big_category_code": "05", "subcategory_code": "05-05", "subcategory_name": "并行计算与性能优化"},
            {"path": "05_交付物/通俗细分版_2026-08-07/05_编程数学数据分析和可视化/00_大分类总览.docx", "format": "docx", "scope": "overview", "big_category_code": "05"},
            {"path": "05_交付物/通俗细分版_2026-08-07/05_编程数学数据分析和可视化/00_大分类总览.xlsx", "format": "xlsx", "scope": "overview", "big_category_code": "05"},
        ]
        selected = self.builder.select_manifest_items(manifest, ["05-overview", "05-05"])
        self.assertEqual([(item["scope"], item.get("subcategory_code")) for item in selected], [("overview", None), ("subcategory", "05-05")])
        with self.assertRaisesRegex(ValueError, "未知 --only"):
            self.builder.select_manifest_items(manifest, ["99-99"])
        unsafe = [{**manifest[0], "path": "../escape.docx"}]
        with self.assertRaisesRegex(ValueError, "不安全"):
            self.builder.select_manifest_items(unsafe, None)

    def test_manifest_contract_rejects_cross_category_names_paths_and_unpaired_outputs(self):
        """Manifest metadata and paired Office paths must be an exact projection of taxonomy."""
        taxonomy = [OVERVIEW["subcategories"][1]]
        base = "05_交付物/通俗细分版_2026-08-07/05_编程数学数据分析和可视化/05-05_并行计算与性能优化/05-05_并行计算与性能优化_GitHub技能调研"
        manifest = [
            {"path": f"{base}.docx", "format": "docx", "scope": "subcategory", "big_category_code": "05", "subcategory_code": "05-05", "subcategory_name": "并行计算与性能优化"},
            {"path": f"{base}.xlsx", "format": "xlsx", "scope": "subcategory", "big_category_code": "05", "subcategory_code": "05-05", "subcategory_name": "并行计算与性能优化"},
        ]
        self.builder.validate_manifest_contract(manifest, taxonomy, require_complete=False)
        mutations = {
            "跨大类代码": [{**manifest[0], "big_category_code": "03"}, manifest[1]],
            "错名称": [{**manifest[0], "subcategory_name": "错误名称"}, manifest[1]],
            "错路径": [{**manifest[0], "path": manifest[0]["path"].replace("05_编程数学数据分析和可视化", "03_文献检索与学术研究")}, manifest[1]],
            "缺配对": [manifest[0]],
            "scope不一致": [manifest[0], {**manifest[1], "scope": "overview"}],
        }
        for label, mutated in mutations.items():
            with self.subTest(label=label), self.assertRaises(ValueError):
                self.builder.validate_manifest_contract(mutated, taxonomy, require_complete=False)

        overview_pair = [
            {"path": "05_交付物/通俗细分版_2026-08-07/05_编程数学数据分析和可视化/00_大分类总览.docx", "format": "docx", "scope": "overview", "big_category_code": "05"},
            {"path": "05_交付物/通俗细分版_2026-08-07/05_编程数学数据分析和可视化/00_大分类总览.xlsx", "format": "xlsx", "scope": "overview", "big_category_code": "05"},
        ]
        with self.assertRaisesRegex(ValueError, "未完整覆盖 taxonomy"):
            self.builder.validate_manifest_contract(
                [*overview_pair, *manifest], OVERVIEW["subcategories"]
            )

    def test_source_contract_rejects_assignment_name_and_knowledge_base_drift(self):
        """Catalog rows, assignments, taxonomy, repositories, and actual KB indexes stay aligned."""
        record = fixture_record("GH-05-0003")
        taxonomy = [OVERVIEW["subcategories"][1]]
        assignments = {record["id"]: "05-05"}
        repositories = {record["repo"]: {"license": "MIT"}}
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            kb = write_fixture_knowledge_base(root, taxonomy[0], [record])
            self.builder.validate_source_contract(
                [record], taxonomy, assignments, repositories, root
            )
            with self.assertRaisesRegex(ValueError, "名称"):
                self.builder.validate_source_contract(
                    [{**record, "subcategory_name": "错误名称"}], taxonomy, assignments, repositories, root
                )
            with self.assertRaisesRegex(ValueError, "归属"):
                self.builder.validate_source_contract(
                    [record], taxonomy, {record["id"]: "03-02"}, repositories, root
                )
            kb.write_text(kb.read_text(encoding="utf-8").replace("共 1 项", "共 2 项"), encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "知识库"):
                self.builder.validate_source_contract(
                    [record], taxonomy, assignments, repositories, root
                )

    def test_knowledge_base_table_rejects_extra_duplicate_and_every_field_or_link_drift(self):
        """Leaf indexes are exact structured ledgers, not bags of reassuring substrings."""
        record = fixture_record("GH-05-0003")
        category = OVERVIEW["subcategories"][1]
        assignments = {record["id"]: category["code"]}
        repositories = {record["repo"]: {"license": "MIT"}}
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            kb = write_fixture_knowledge_base(root, category, [record])
            baseline = kb.read_text(encoding="utf-8")
            row = next(line for line in baseline.splitlines() if line.startswith(f"| {record['id']} "))
            extra = row.replace(record["id"], "GH-05-9999").replace(record["cn"], "额外条目")
            mutations = {
                "额外行": baseline.replace(row, f"{row}\n{extra}"),
                "重复行": baseline.replace(row, f"{row}\n{row}"),
                "中文名": baseline.replace(record["cn"], "错误中文名"),
                "用途": baseline.replace(record["plain_purpose"], "错误用途"),
                "推荐程度": baseline.replace(f"| {record['priority']} | [查看]", "| 低 | [查看]"),
                "条目链接": baseline.replace(
                    f"../../skills/{record['id']}_{record['name']}.md",
                    f"../../skills/{record['id']}_wrong.md",
                ),
            }
            for label, content in mutations.items():
                with self.subTest(label=label):
                    kb.write_text(content, encoding="utf-8")
                    with self.assertRaisesRegex(ValueError, "知识库"):
                        self.builder.validate_source_contract(
                            [record], [category], assignments, repositories, root
                        )

    def test_assignment_contract_rejects_missing_and_extra_members(self):
        """Assignments must be a per-category and global bijection with the plain catalog."""
        record = fixture_record("GH-05-0003")
        category = OVERVIEW["subcategories"][1]
        repositories = {record["repo"]: {"license": "MIT"}}
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            write_fixture_knowledge_base(root, category, [record])
            for label, assignments in {
                "缺失成员": {},
                "额外成员": {record["id"]: category["code"], "GH-05-9999": category["code"]},
            }.items():
                with self.subTest(label=label), self.assertRaisesRegex(ValueError, "归属台账"):
                    self.builder.validate_source_contract(
                        [record], [category], assignments, repositories, root
                    )

    def test_json_loaders_reject_duplicate_keys_in_raw_assignment_text(self):
        """Duplicate assignment keys must be rejected before JSON decoding can overwrite them."""
        raw_json = (
            '{"taxonomy":[],"assignments":{'
            '"GH-05-0001":"05-01","GH-05-0001":"05-02"}}'
        )
        with tempfile.TemporaryDirectory() as temporary_directory:
            path = Path(temporary_directory) / "subcategory_assignments.json"
            path.write_text(raw_json, encoding="utf-8")
            for label, loader in {
                "生成器": self.builder._load_json,
                "验证器": self.verifier._load_json,
            }.items():
                with self.subTest(label=label), self.assertRaisesRegex(ValueError, "重复 JSON 键.*GH-05-0001"):
                    loader(path)

    def test_overview_verifier_rejects_missing_wrong_duplicate_extra_rows_and_links(self):
        """Every overview navigation row and link must exactly match taxonomy and members."""
        records = [
            fixture_record("GH-05-0001", subcategory_code="05-01"),
            fixture_record("GH-05-0003"),
            fixture_record("GH-05-0024"),
        ]
        records[0]["subcategory_name"] = "计算设备与资源规划"
        expected_item = {"scope": "overview", "big_category_code": "05"}
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)

            def issues_for(document):
                path = root / "overview.docx"
                document.save(path)
                return self.verifier.verify_document(
                    path,
                    scope="overview",
                    expected_records=records,
                    expected_taxonomy=OVERVIEW["subcategories"],
                    expected_item=expected_item,
                )

            self.assertEqual(issues_for(self.builder.build_overview_document(OVERVIEW, records)), [])

            missing = self.builder.build_overview_document(OVERVIEW, records)
            missing.tables[0]._tbl.remove(missing.tables[0].rows[-1]._tr)
            self.assertTrue(any("概览行" in issue for issue in issues_for(missing)))

            wrong_count = self.builder.build_overview_document(OVERVIEW, records)
            wrong_count.tables[0].cell(2, 2).text = "99"
            self.assertTrue(any("概览行" in issue for issue in issues_for(wrong_count)))

            duplicate = self.builder.build_overview_document(OVERVIEW, records)
            duplicate_row = duplicate.tables[0].add_row()
            for index, value in enumerate(("05-05 并行计算与性能优化", OVERVIEW["subcategories"][1]["inclusion_focus"], "2")):
                duplicate_row.cells[index].text = value
            self.assertTrue(any("概览行" in issue for issue in issues_for(duplicate)))

            wrong_link = self.builder.build_overview_document(OVERVIEW, records)
            relationship = next(
                rel for rel in wrong_link.part.rels.values()
                if rel.reltype.endswith("/hyperlink") and rel.target_ref == records[0]["repo_url"]
            )
            relationship._target = "https://example.invalid/wrong"
            self.assertTrue(any("概览链接" in issue for issue in issues_for(wrong_link)))

    def test_overview_verifier_rejects_duplicate_and_row_swapped_knowledge_links(self):
        """Per-row relationship order must catch duplicate or swapped targets with unchanged sets."""
        records = [
            fixture_record("GH-05-0001", subcategory_code="05-01"),
            fixture_record("GH-05-0003"),
        ]
        expected_item = {"scope": "overview", "big_category_code": "05"}
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)

            def issues_for(document):
                path = root / "overview-row-links.docx"
                document.save(path)
                return self.verifier.verify_document(
                    path,
                    scope="overview",
                    expected_records=records,
                    expected_taxonomy=OVERVIEW["subcategories"],
                    expected_item=expected_item,
                )

            duplicated = self.builder.build_overview_document(OVERVIEW, records)
            first_cell_paragraph = duplicated.tables[0].cell(1, 0).paragraphs[0]
            first_target = self.builder._knowledge_base_target(OVERVIEW["subcategories"][0])
            self.builder._add_hyperlink(
                first_cell_paragraph, "重复知识库入口", first_target, allow_relative=True
            )
            duplicate_issues = issues_for(duplicated)
            self.assertTrue(any("概览逐行链接" in issue for issue in duplicate_issues), duplicate_issues)

            swapped = self.builder.build_overview_document(OVERVIEW, records)
            relationships = []
            for row in swapped.tables[0].rows[1:3]:
                hyperlink = row.cells[0]._tc.xpath(".//w:hyperlink")[0]
                relationships.append(swapped.part.rels[hyperlink.get(qn("r:id"))])
            relationships[0]._target, relationships[1]._target = (
                relationships[1]._target,
                relationships[0]._target,
            )
            swapped_issues = issues_for(swapped)
            self.assertTrue(any("概览逐行链接" in issue for issue in swapped_issues), swapped_issues)

    def test_overview_verifier_rejects_existing_link_duplicated_outside_navigation(self):
        """A duplicate body link must fail even when the unique target set stays unchanged."""
        records = [
            fixture_record("GH-05-0001", subcategory_code="05-01"),
            fixture_record("GH-05-0003"),
        ]
        expected_item = {"scope": "overview", "big_category_code": "05"}
        targets = {
            "知识库": self.builder._knowledge_base_target(OVERVIEW["subcategories"][0]),
            "仓库": records[0]["repo_url"],
        }
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            for label, target in targets.items():
                with self.subTest(label=label):
                    document = self.builder.build_overview_document(OVERVIEW, records)
                    paragraph = document.add_paragraph("表外重复链接：")
                    self.builder._add_hyperlink(
                        paragraph,
                        label,
                        target,
                        allow_relative=label == "知识库",
                    )
                    path = root / f"overview-extra-{label}.docx"
                    document.save(path)
                    issues = self.verifier.verify_document(
                        path,
                        scope="overview",
                        expected_records=records,
                        expected_taxonomy=OVERVIEW["subcategories"],
                        expected_item=expected_item,
                    )
                    self.assertTrue(any("概览链接" in issue for issue in issues), issues)

    def test_manifest_driven_generation_is_idempotent_for_overview_and_subcategory(self):
        """Repeated generation with reversed input must preserve paths, order, and bytes."""
        records = [fixture_record("GH-05-0001", subcategory_code="05-01"), fixture_record("GH-05-0003")]
        taxonomy = OVERVIEW["subcategories"]
        manifest = [
            {"path": "05_交付物/通俗细分版_2026-08-07/05_编程数学数据分析和可视化/00_大分类总览.docx", "format": "docx", "scope": "overview", "big_category_code": "05"},
            {"path": "05_交付物/通俗细分版_2026-08-07/05_编程数学数据分析和可视化/00_大分类总览.xlsx", "format": "xlsx", "scope": "overview", "big_category_code": "05"},
            {"path": "05_交付物/通俗细分版_2026-08-07/05_编程数学数据分析和可视化/05-01_计算设备与资源规划/05-01_计算设备与资源规划_GitHub技能调研.docx", "format": "docx", "scope": "subcategory", "big_category_code": "05", "subcategory_code": "05-01", "subcategory_name": "计算设备与资源规划"},
            {"path": "05_交付物/通俗细分版_2026-08-07/05_编程数学数据分析和可视化/05-01_计算设备与资源规划/05-01_计算设备与资源规划_GitHub技能调研.xlsx", "format": "xlsx", "scope": "subcategory", "big_category_code": "05", "subcategory_code": "05-01", "subcategory_name": "计算设备与资源规划"},
            {"path": "05_交付物/通俗细分版_2026-08-07/05_编程数学数据分析和可视化/05-05_并行计算与性能优化/05-05_并行计算与性能优化_GitHub技能调研.docx", "format": "docx", "scope": "subcategory", "big_category_code": "05", "subcategory_code": "05-05", "subcategory_name": "并行计算与性能优化"},
            {"path": "05_交付物/通俗细分版_2026-08-07/05_编程数学数据分析和可视化/05-05_并行计算与性能优化/05-05_并行计算与性能优化_GitHub技能调研.xlsx", "format": "xlsx", "scope": "subcategory", "big_category_code": "05", "subcategory_code": "05-05", "subcategory_name": "并行计算与性能优化"},
        ]
        with tempfile.TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            first = self.builder.generate_documents(records, taxonomy, manifest, root, only=["05-overview", "05-05"])
            first_hashes = {path.relative_to(root).as_posix(): hashlib.sha256(path.read_bytes()).hexdigest() for path in first}
            second = self.builder.generate_documents(list(reversed(records)), list(reversed(taxonomy)), list(reversed(manifest)), root, only=["05-05", "05-overview"])
            second_hashes = {path.relative_to(root).as_posix(): hashlib.sha256(path.read_bytes()).hexdigest() for path in second}
            self.assertEqual(first_hashes, second_hashes)

    def test_renderer_plans_one_isolated_output_directory_per_document(self):
        """Two DOCX files must never overwrite each other's page-N PNGs."""
        items = [
            {"path": "05_交付物/通俗细分版_2026-08-07/05_编程数学数据分析和可视化/00_大分类总览.docx", "format": "docx", "scope": "overview", "big_category_code": "05"},
            {"path": "05_交付物/通俗细分版_2026-08-07/05_编程数学数据分析和可视化/05-05_并行计算与性能优化/05-05_并行计算与性能优化_GitHub技能调研.docx", "format": "docx", "scope": "subcategory", "big_category_code": "05", "subcategory_code": "05-05"},
        ]
        plan = self.renderer.build_render_plan(items, PROJECT_ROOT, PROJECT_ROOT / "06_过程记录" / "renders" / "subcategorized_docx")
        self.assertEqual([item["key"] for item in plan], ["05-overview", "05-05"])
        self.assertEqual(len({item["output_dir"] for item in plan}), 2)
        self.assertTrue(all(Path(item["docx_path"]).is_absolute() for item in plan))

    def test_renderer_rejects_a_missing_packaged_render_script(self):
        """A missing documents-skill renderer must fail before any false QA claim."""
        with self.assertRaisesRegex(FileNotFoundError, "documents 技能渲染器"):
            self.renderer.render_plan([], render_script=PROJECT_ROOT / "missing-render-docx.py")

    def test_renderer_path_supports_explicit_injection_and_clear_missing_error(self):
        """Canonical renderer discovery must be portable across users and plugin versions."""
        with tempfile.TemporaryDirectory() as temporary_directory:
            render_script = Path(temporary_directory) / "render_docx.py"
            render_script.write_text("# canonical test renderer\n", encoding="utf-8")
            self.assertEqual(
                self.renderer.resolve_render_script(render_script, environ={}),
                render_script.resolve(),
            )
            with self.assertRaisesRegex(FileNotFoundError, "--render-script|DOCUMENTS_RENDER_DOCX"):
                self.renderer.resolve_render_script(
                    Path(temporary_directory) / "missing.py", environ={}
                )


if __name__ == "__main__":
    unittest.main()
