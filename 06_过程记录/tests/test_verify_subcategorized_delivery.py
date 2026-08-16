"""Task 9 contracts for the single project-level delivery verifier."""

from __future__ import annotations

import copy
import hashlib
import importlib.util
import json
import shutil
import subprocess
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from PIL import Image, ImageDraw


PROJECT_ROOT = Path(__file__).resolve().parents[2]
VERIFIER_PATH = PROJECT_ROOT / "06_过程记录" / "tools" / "verify_subcategorized_delivery.py"
ORIGINAL_BASENAMES = {
    "01": "01_学术写作、引用与出版_GitHub技能调研",
    "02": "02_文档、表格、演示文稿与办公自动化_GitHub技能调研",
    "03": "03_文献检索与学术研究_GitHub技能调研",
    "04": "04_图书馆与信息素养_GitHub技能调研",
    "05": "05_编程、数学、数据分析和可视化_GitHub技能调研",
}


def load_verifier():
    if not VERIFIER_PATH.is_file():
        raise AssertionError(f"缺少计划脚本: {VERIFIER_PATH.name}")
    spec = importlib.util.spec_from_file_location("verify_subcategorized_delivery", VERIFIER_PATH)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def copy_contract_data(root: Path) -> None:
    source = PROJECT_ROOT / "03_候选池"
    target = root / "03_候选池"
    (target / "deduplicated").mkdir(parents=True)
    (target / "derived").mkdir(parents=True)
    for code in range(1, 6):
        shutil.copy2(
            source / "deduplicated" / f"category_{code:02d}.json",
            target / "deduplicated" / f"category_{code:02d}.json",
        )
    for name in (
        "subcategory_assignments.json",
        "plain_output_contract.json",
        "plain_language_catalog.json",
        "subcategory_manifest.json",
    ):
        shutil.copy2(source / "derived" / name, target / "derived" / name)


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, payload: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def make_delivery_tree(root: Path) -> list[dict]:
    manifest = load_json(PROJECT_ROOT / "03_候选池/derived/subcategory_manifest.json")
    write_json(root / "03_候选池/derived/subcategory_manifest.json", manifest)
    for item in manifest:
        path = root / Path(*item["path"].split("/"))
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(f"fixture:{item['path']}".encode("utf-8"))
    return manifest


def make_originals_and_archive(root: Path) -> None:
    delivery = root / "05_交付物"
    archive = delivery / "原始版_2026-08-06"
    archive.mkdir(parents=True, exist_ok=True)
    for big, basename in ORIGINAL_BASENAMES.items():
        for suffix in ("docx", "xlsx"):
            path = delivery / f"{basename}.{suffix}"
            path.write_bytes(f"original-{big}-{suffix}".encode("utf-8"))
            shutil.copy2(path, archive / path.name)
    for suffix in ("docx", "xlsx"):
        (delivery / f"0809_计算机类_跨平台技能调研.{suffix}").write_bytes(
            f"0809-{suffix}".encode("utf-8")
        )


def write_png(path: Path, size: tuple[int, int] = (900, 500)) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((80, 80, size[0] - 80, min(size[1] - 40, 180)), fill="#315A78")
    image.save(path)


def write_review_log(path: Path, records: list[dict]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        "".join(json.dumps(row, ensure_ascii=False, sort_keys=True) + "\n" for row in records),
        encoding="utf-8",
    )


def make_visual_fixture(root: Path, verifier) -> tuple[list[dict], list[dict]]:
    manifest = [
        {
            "path": "05_交付物/通俗细分版_2026-08-07/01_示例/00_大分类总览.docx",
            "format": "docx",
            "scope": "overview",
            "big_category_code": "01",
        },
        {
            "path": "05_交付物/通俗细分版_2026-08-07/01_示例/00_大分类总览.xlsx",
            "format": "xlsx",
            "scope": "overview",
            "big_category_code": "01",
        },
    ]
    for item in manifest:
        path = root / Path(*item["path"].split("/"))
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(b"fixture")
    catalog = [{"cat": "01"} for _ in range(11)]
    docx_dir = root / "06_过程记录/renders/subcategorized_docx/01-overview"
    write_png(docx_dir / "page-1.png", (800, 1100))
    write_json(docx_dir / "rendered-pages.json", {"pages": ["page-1.png"]})
    xlsx_dir = root / "06_过程记录/renders/subcategorized_xlsx/01-overview"
    for name in ("1_使用说明.png", "2_AI技能清单.png", "3_分类统计.png", "4_来源清单.png"):
        write_png(xlsx_dir / name)
    for label in ("title-header", "longest-text", "longest-url", "last-row"):
        write_png(xlsx_dir / f"2_AI技能清单_segment_{label}_A1-V5.png", (2600, 240))

    visual = verifier.load_visual_module(root)
    docx = visual.collect_docx_render_inventory(
        root / "06_过程记录/renders/subcategorized_docx", manifest
    )
    xlsx = visual.collect_xlsx_render_inventory(
        root / "06_过程记录/renders/subcategorized_xlsx",
        manifest,
        required_segment_keys={"01-overview"},
    )
    inventory = visual.build_inventory_document(docx + xlsx)
    inventory_path = root / "06_过程记录/visual_review/task-7-inventory.json"
    write_json(inventory_path, inventory)
    images = inventory["images"]
    records = [
        {
            "record_type": "session",
            "schema_version": 1,
            "session_id": "fixture-session",
            "inventory_digest": inventory["inventory_digest"],
            "time_basis": "ordered fixture markers",
        },
        {
            "record_type": "batch",
            "sequence": 1,
            "batch_id": "B1",
            "reviewer_id": "fixture-reviewer",
            "session_id": "fixture-session",
            "started_at": "2026-08-09T00:01:00+08:00",
            "ended_at": "2026-08-09T00:02:00+08:00",
            "inspection_criteria": ["截断", "重叠", "空白"],
        },
        {
            "record_type": "batch",
            "sequence": 2,
            "batch_id": "B2",
            "reviewer_id": "fixture-reviewer",
            "session_id": "fixture-session",
            "started_at": "2026-08-09T00:03:00+08:00",
            "ended_at": "2026-08-09T00:04:00+08:00",
            "inspection_criteria": ["截断", "重叠", "空白"],
        },
    ]
    for index, row in enumerate(images):
        records.append(
            {
                "record_type": "image",
                "relative_path": row["relative_path"],
                "image_sha256": row["image_sha256"],
                "batch_id": "B1" if index < len(images) // 2 else "B2",
                "status": "pass",
                "issues": [],
            }
        )
    review_path = root / "06_过程记录/visual_review/task-7-review-log.jsonl"
    write_review_log(review_path, records)
    review_hash = hashlib.sha256(review_path.read_bytes()).hexdigest()
    finalized = visual.finalize_review_inventory(
        root, inventory, records, review_log_sha256=review_hash
    )
    write_json(root / "06_过程记录/visual_review/task-7-finalized.json", finalized)
    return manifest, catalog


class DataContractTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.verifier = load_verifier()

    def test_real_contract_is_157_to_61_with_frozen_assignment_and_plain_facts(self):
        summary = self.verifier.verify_data_contract(PROJECT_ROOT)
        self.assertEqual(summary["source_records"], 157)
        self.assertEqual(summary["subcategories"], 61)
        self.assertEqual(summary["big_category_records"], {"01": 20, "02": 22, "03": 31, "04": 29, "05": 55})
        self.assertEqual(summary["fact_drift"], 0)
        self.assertEqual(summary["readability_issues"], 0)

    def test_source_missing_and_duplicate_are_rejected(self):
        for mutation in ("missing", "duplicate"):
            with self.subTest(mutation=mutation), tempfile.TemporaryDirectory() as temporary:
                root = Path(temporary)
                copy_contract_data(root)
                if mutation == "missing":
                    (root / "03_候选池/deduplicated/category_05.json").unlink()
                else:
                    path = root / "03_候选池/deduplicated/category_01.json"
                    payload = load_json(path)
                    payload["records"].append(copy.deepcopy(payload["records"][0]))
                    write_json(path, payload)
                with self.assertRaisesRegex(ValueError, "源数据|缺失|重复"):
                    self.verifier.verify_data_contract(root)

    def test_taxonomy_count_and_assignment_duplicate_or_wrong_owner_are_rejected(self):
        for mutation in ("taxonomy", "duplicate-key", "wrong-owner", "reordered"):
            with self.subTest(mutation=mutation), tempfile.TemporaryDirectory() as temporary:
                root = Path(temporary)
                copy_contract_data(root)
                path = root / "03_候选池/derived/subcategory_assignments.json"
                if mutation == "duplicate-key":
                    text = path.read_text(encoding="utf-8")
                    text = text.replace(
                        '"GH-01-0001": "01-02"',
                        '"GH-01-0001": "01-02",\n    "GH-01-0001": "01-01"',
                        1,
                    )
                    path.write_text(text, encoding="utf-8")
                else:
                    payload = load_json(path)
                    if mutation == "taxonomy":
                        payload["taxonomy"].pop()
                    elif mutation == "reordered":
                        first_key = next(iter(payload["assignments"]))
                        first_value = payload["assignments"].pop(first_key)
                        payload["assignments"][first_key] = first_value
                    else:
                        payload["assignments"]["GH-01-0001"] = "02-01"
                    write_json(path, payload)
                with self.assertRaisesRegex(ValueError, "taxonomy|小分类|重复 JSON|归属|冻结|顺序"):
                    self.verifier.verify_data_contract(root)


class DeliveryArchiveAndLinkTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.verifier = load_verifier()
        cls.taxonomy = load_json(PROJECT_ROOT / "03_候选池/derived/subcategory_assignments.json")["taxonomy"]

    def test_manifest_tree_rejects_missing_extra_noise_and_empty(self):
        for mutation in ("missing", "extra", "noise", "empty", "empty-directory"):
            with self.subTest(mutation=mutation), tempfile.TemporaryDirectory() as temporary:
                root = Path(temporary)
                manifest = make_delivery_tree(root)
                if mutation == "missing":
                    manifest.pop()
                    write_json(root / "03_候选池/derived/subcategory_manifest.json", manifest)
                elif mutation == "extra":
                    manifest.append({**manifest[-1], "path": manifest[-1]["path"].replace(".xlsx", "_extra.xlsx")})
                    write_json(root / "03_候选池/derived/subcategory_manifest.json", manifest)
                elif mutation == "noise":
                    (root / "05_交付物/通俗细分版_2026-08-07/noise.tmp").write_bytes(b"noise")
                elif mutation == "empty":
                    (root / Path(*manifest[0]["path"].split("/"))).write_bytes(b"")
                else:
                    (root / "05_交付物/通俗细分版_2026-08-07/empty-noise").mkdir()
                with self.assertRaisesRegex(ValueError, "manifest|交付|额外|空文件|132"):
                    self.verifier.verify_manifest_and_delivery(root, self.taxonomy)

    def test_archive_hash_mismatch_and_0809_intrusion_are_rejected(self):
        for mutation in ("hash", "0809-in-archive", "renamed-pair"):
            with self.subTest(mutation=mutation), tempfile.TemporaryDirectory() as temporary:
                root = Path(temporary)
                make_originals_and_archive(root)
                archive = root / "05_交付物/原始版_2026-08-06"
                if mutation == "hash":
                    next(archive.glob("01_*.docx")).write_bytes(b"changed")
                elif mutation == "0809-in-archive":
                    (archive / "0809_计算机类_跨平台技能调研.docx").write_bytes(b"intrusion")
                else:
                    original = next((root / "05_交付物").glob("01_*.docx"))
                    archived = archive / original.name
                    original.rename(original.with_name("01_同步改名.docx"))
                    archived.rename(archive / "01_同步改名.docx")
                with self.assertRaisesRegex(ValueError, "归档|SHA|0809|文件集合|批准文件名"):
                    self.verifier.verify_originals_and_archive(root)

    def test_commonmark_link_audit_rejects_a_broken_or_bare_space_target(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            page = root / "INDEX.md"
            page.write_text("[坏链接](missing file.md)\n", encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "CommonMark|坏链|空格"):
                self.verifier.verify_markdown_links(root, [page])

    def test_link_audit_ignores_fenced_code_and_task_list_syntax_but_rejects_real_broken_links(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            page = root / "INDEX.md"
            page.write_text(
                "- [ ] Run the task\n"
                "[ ]: missing-task.md\n\n"
                "```javascript\n"
                "const formula = values[0][0];\n"
                "[code-item]: missing-code.md\n"
                "[code-item]\n"
                "```\n",
                encoding="utf-8",
            )
            self.assertEqual(
                self.verifier.verify_markdown_links(root, [page]),
                {"markdown_files": 1, "links": 0, "local_links": 0},
            )
            page.write_text(
                page.read_text(encoding="utf-8") + "\n[real broken](missing-real.md)\n",
                encoding="utf-8",
            )
            with self.assertRaisesRegex(ValueError, "坏链"):
                self.verifier.verify_markdown_links(root, [page])

    def test_reference_style_bad_link_and_untracked_markdown_are_not_skipped(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            subprocess.run(["git", "init", "-q"], cwd=root, check=True)
            page = root / "UNTRACKED.md"
            page.write_text("[坏链接][missing]\n\n[missing]: <missing file.md>\n", encoding="utf-8")
            self.assertIn(page, self.verifier._tracked_markdown(root))
            internal = root / ".worktrees/other/BROKEN.md"
            internal.parent.mkdir(parents=True)
            internal.write_text("[坏](missing.md)\n", encoding="utf-8")
            self.assertNotIn(internal.resolve(), self.verifier._tracked_markdown(root))
            with self.assertRaisesRegex(ValueError, "引用|坏链|CommonMark"):
                self.verifier.verify_markdown_links(root, [page])
            for name, text in (
                ("IMAGE.md", "![图][asset]\n\n[asset]: missing.png\n"),
                ("SHORTCUT.md", "[资料]\n\n[资料]: missing.md\n"),
            ):
                with self.subTest(name=name):
                    candidate = root / name
                    candidate.write_text(text, encoding="utf-8")
                    with self.assertRaisesRegex(ValueError, "引用|坏链|CommonMark"):
                        self.verifier.verify_markdown_links(root, [candidate])

        # 发现范围的预期在测试中独立写死：正式项目页必须纳入，内部工作目录必须排除。
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            subprocess.run(["git", "init", "-q"], cwd=root, check=True)
            tracked = root / "TRACKED.md"
            untracked = root / "UNTRACKED.md"
            tracked.write_text("正式用户页面\n", encoding="utf-8")
            untracked.write_text("未跟踪的正式用户页面\n", encoding="utf-8")

            internal_paths = (
                root / ".superpowers/sdd/tracked.md",
                root / ".superpowers/sdd/untracked.md",
                root / ".worktrees/other/tracked.md",
                root / ".worktrees/other/untracked.md",
            )
            for path in internal_paths:
                path.parent.mkdir(parents=True, exist_ok=True)
                path.write_text(f"内部文件：{path.name}\n", encoding="utf-8")

            nested_repo_page = root / "vendor/nested-repository/README.md"
            nested_repo_page.parent.mkdir(parents=True, exist_ok=True)
            (nested_repo_page.parent / ".git").mkdir()
            nested_repo_page.write_text("[外部仓库链接](missing.md)\n", encoding="utf-8")
            subprocess.run(
                [
                    "git", "add", "--", str(tracked),
                    str(internal_paths[0]), str(internal_paths[2]),
                ],
                cwd=root,
                check=True,
            )

            expected_paths = [tracked.resolve(), untracked.resolve()]

            def inventory(paths: list[Path]) -> tuple[int, str]:
                bindings = [
                    (
                        path.relative_to(root).as_posix(),
                        hashlib.sha256(path.read_bytes()).hexdigest(),
                    )
                    for path in paths
                ]
                digest = hashlib.sha256(
                    json.dumps(bindings, ensure_ascii=False, separators=(",", ":")).encode("utf-8")
                ).hexdigest()
                return len(paths), digest

            discovered = self.verifier._tracked_markdown(root)
            expected_inventory = inventory(expected_paths)
            self.assertEqual(discovered, expected_paths)
            self.assertEqual(inventory(discovered), expected_inventory)
            for path in internal_paths:
                path.write_text(f"已修改的内部文件：{path.name}\n", encoding="utf-8")
            discovered_after = self.verifier._tracked_markdown(root)
            self.assertEqual(discovered_after, expected_paths)
            self.assertEqual(inventory(discovered_after), expected_inventory)
            self.assertNotIn(nested_repo_page.resolve(), discovered_after)

    def test_real_navigation_is_an_exact_projection_of_catalog_and_manifest(self):
        assignment_data = load_json(
            PROJECT_ROOT / "03_候选池/derived/subcategory_assignments.json"
        )
        catalog = load_json(PROJECT_ROOT / "03_候选池/derived/plain_language_catalog.json")
        manifest = load_json(PROJECT_ROOT / "03_候选池/derived/subcategory_manifest.json")
        summary = self.verifier.verify_navigation(
            PROJECT_ROOT,
            assignment_data["taxonomy"],
            assignment_data["assignments"],
            catalog=catalog,
            manifest=manifest,
        )
        self.assertEqual(summary["leaf_pages"], 61)
        self.assertEqual(summary["domain_indexes"], 5)
        self.assertEqual(summary["total_indexes"], 1)


class VisualAndTextTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.verifier = load_verifier()

    def expectations(self):
        return self.verifier.VisualExpectations(
            docx_pages=1,
            worksheets=4,
            segments=4,
            images=9,
            batches=2,
        )

    def test_visual_gate_rejects_stale_docx_marker_and_missing_xlsx_sheet_or_segment(self):
        for mutation in ("stale", "worksheet", "segment"):
            with self.subTest(mutation=mutation), tempfile.TemporaryDirectory() as temporary:
                root = Path(temporary)
                manifest, catalog = make_visual_fixture(root, self.verifier)
                if mutation == "stale":
                    write_png(root / "06_过程记录/renders/subcategorized_docx/01-overview/page-2.png")
                elif mutation == "worksheet":
                    (root / "06_过程记录/renders/subcategorized_xlsx/01-overview/4_来源清单.png").unlink()
                else:
                    next((root / "06_过程记录/renders/subcategorized_xlsx/01-overview").glob("*segment_last-row*.png")).unlink()
                with self.assertRaisesRegex(ValueError, "残留|预期集合|渲染图|last-row|缺少"):
                    self.verifier.verify_visual_evidence(root, manifest, catalog, self.expectations())

    def test_visual_gate_rejects_inventory_hash_digest_review_log_and_finalized_drift(self):
        for mutation in ("inventory-digest", "image-hash", "review-log", "finalized"):
            with self.subTest(mutation=mutation), tempfile.TemporaryDirectory() as temporary:
                root = Path(temporary)
                manifest, catalog = make_visual_fixture(root, self.verifier)
                inventory_path = root / "06_过程记录/visual_review/task-7-inventory.json"
                review_path = root / "06_过程记录/visual_review/task-7-review-log.jsonl"
                finalized_path = root / "06_过程记录/visual_review/task-7-finalized.json"
                if mutation == "inventory-digest":
                    payload = load_json(inventory_path)
                    payload["inventory_digest"] = "0" * 64
                    write_json(inventory_path, payload)
                elif mutation == "image-hash":
                    payload = load_json(inventory_path)
                    payload["images"][0]["image_sha256"] = "1" * 64
                    write_json(inventory_path, payload)
                elif mutation == "review-log":
                    lines = review_path.read_text(encoding="utf-8").splitlines()
                    review_path.write_text("\n".join(lines[:-1]) + "\n", encoding="utf-8")
                else:
                    payload = load_json(finalized_path)
                    payload["review_complete"] = False
                    write_json(finalized_path, payload)
                with self.assertRaisesRegex(ValueError, "digest|hash|review|复核|库存|finalized|完成"):
                    self.verifier.verify_visual_evidence(root, manifest, catalog, self.expectations())

    def test_placeholder_overclaim_and_discipline_misstatement_are_rejected(self):
        samples = {
            "placeholder": "这里是 TODO 模板标记。",
            "template": "这里仍有模板占位。",
            "internal": "请运行 06_过程记录/tools/internal.py 后交付。",
            "overclaim": "这个候选已经安装并已运行成功。",
            "overclaim-available": "这个候选已经验证可用。",
            "overclaim-verified": "经验证可以正常使用。",
            "overclaim-test": "测试通过可正常运行。",
            "discipline": "本轮按照学科分类建立小分类。",
            "fixme": "FIXME: 稍后补写。",
            "template-en": "TEMPLATE",
            "temporary-path": r"请读取 D:\临时\tmp\result.txt。",
            "mixed-overclaim": "A 未安装；B 已经安装并已运行成功。",
            "mixed-discipline": "本轮不涉及学科分类；附录按照学科分类建立。",
            "comma-overclaim": "A 未安装，但 B 已经安装并已运行成功。",
            "comma-discipline": "本轮不涉及学科分类，但附录按照学科分类建立。",
        }
        for label, text in samples.items():
            with self.subTest(label=label):
                issues = self.verifier.text_quality_issues(text, source=f"fixture:{label}")
                self.assertTrue(issues, label)
        self.assertEqual(
            self.verifier.text_quality_issues("本轮不涉及专业或学科分类；候选未安装、未运行。", source="safe"),
            [],
        )

    def test_office_visible_text_rejoins_rich_text_runs_and_cells(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            docx_path = root / "rich.docx"
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("候选已经")
            paragraph.add_run("验证可用。")
            document.save(docx_path)
            self.assertIn("候选已经验证可用", self.verifier._docx_text(docx_path))

            xlsx_path = root / "rich.xlsx"
            with zipfile.ZipFile(xlsx_path, "w") as archive:
                archive.writestr(
                    "xl/sharedStrings.xml",
                    '<?xml version="1.0" encoding="UTF-8"?>'
                    '<sst xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
                    '<si><r><t>候选已经</t></r><r><t>验证可用。</t></r></si></sst>',
                )
                archive.writestr(
                    "xl/worksheets/sheet1.xml",
                    '<?xml version="1.0" encoding="UTF-8"?>'
                    '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
                    '<sheetData><row><c t="s"><v>0</v></c></row></sheetData></worksheet>',
                )
            self.assertIn("候选已经验证可用", self.verifier._xlsx_text(xlsx_path))
        self.assertEqual(
            self.verifier.text_quality_issues(
                "说明已核验不代表依赖已安装或任务已成功运行；两者都不等于安装或运行成功。",
                source="safe-boundary",
            ),
            [],
        )
        self.assertEqual(
            self.verifier.text_quality_issues(
                "本表不把候选 Skill 写成已部署工具。", source="safe-workbook"
            ),
            [],
        )
        self.assertEqual(
            self.verifier.text_quality_issues(
                "候选也没有被写成已经证明有效。", source="safe-proof-boundary"
            ),
            [],
        )

    def test_semantic_result_digest_excludes_only_check_time(self):
        payload = {
            "version": 1,
            "checked_at": "first",
            "complete": True,
            "counts": {"source_records": 157},
            "gates": {"data": {"status": "pass"}},
        }
        first = self.verifier.semantic_result_digest(payload)
        payload["checked_at"] = "second"
        self.assertEqual(self.verifier.semantic_result_digest(payload), first)
        payload["counts"]["source_records"] = 156
        self.assertNotEqual(self.verifier.semantic_result_digest(payload), first)


class ProjectResultAndSafetyTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.verifier = load_verifier()

    def test_output_is_restricted_to_the_single_verification_json(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            expected = root / "06_过程记录/verification/subcategorized_delivery_verification.json"
            self.assertEqual(self.verifier.validated_output_path(root, None), expected.resolve())
            self.assertEqual(self.verifier.validated_output_path(root, expected), expected.resolve())
            for protected in (
                root / "03_候选池/source.json",
                root / "05_交付物/report.docx",
                root / "06_过程记录/visual_review/audit.json",
                root / "06_过程记录/verification/not-the-result.json",
            ):
                with self.subTest(path=protected), self.assertRaisesRegex(ValueError, "输出|验收结果|保护"):
                    self.verifier.validated_output_path(root, protected)

    def test_end_to_end_failure_writes_false_result_without_touching_protected_file(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            copy_contract_data(root)
            (root / "03_候选池/deduplicated/category_05.json").unlink()
            sentinel = root / "05_交付物/通俗细分版_2026-08-07/sentinel.docx"
            sentinel.parent.mkdir(parents=True)
            sentinel.write_bytes(b"protected")
            before = hashlib.sha256(sentinel.read_bytes()).hexdigest()
            exit_code = self.verifier.main(["--project-root", str(root)])
            self.assertNotEqual(exit_code, 0)
            result = load_json(
                root / "06_过程记录/verification/subcategorized_delivery_verification.json"
            )
            self.assertIs(result["complete"], False)
            self.assertEqual(set(result["gates"]), {"data_contract"})
            self.assertEqual(result["gates"]["data_contract"]["status"], "fail")
            self.assertEqual(
                result["known_legacy_test_exceptions"],
                list(self.verifier.KNOWN_LEGACY_TEST_EXCEPTIONS),
            )
            self.assertEqual(hashlib.sha256(sentinel.read_bytes()).hexdigest(), before)
            self.assertEqual(list((root / "06_过程记录/verification").glob("*.tmp")), [])

    def test_formal_result_schema_counts_digest_and_manual_boundaries(self):
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            copy_contract_data(root)
            delivery = {
                "manifest": [], "delivery_files": 132,
                "document_files": 66, "spreadsheet_files": 66,
            }
            navigation = {
                "leaf_pages": 61, "domain_indexes": 5, "total_indexes": 1,
                "markdown_files": 272, "links": 1261, "local_links": 670,
            }
            visual = {
                "docx_pages": 259, "worksheets": 264, "segments": 20,
                "review_images": 543, "review_batches": 19,
                "inventory_digest": "80888dde132a1b3e0ef6069458efd3f6e1f5cde813b04f97efa895e91ca6d0f2",
            }
            with (
                patch.object(self.verifier, "verify_manifest_and_delivery", return_value=delivery),
                patch.object(self.verifier, "verify_originals_and_archive", return_value={"archived": 10}),
                patch.object(self.verifier, "verify_navigation", return_value=navigation),
                patch.object(self.verifier, "verify_office_structure", return_value={"docx_verified": 66, "xlsx_verified": 66, "worksheets_verified": 264}),
                patch.object(self.verifier, "verify_visual_evidence", return_value=visual),
                patch.object(self.verifier, "verify_visible_text", return_value={"text_sources": 518, "text_issues": 0}),
            ):
                result = self.verifier.verify_project(root)
        self.assertIs(result["complete"], True)
        self.assertEqual(
            list(result["counts"].values()),
            [157, 61, 132, 66, 66, 10, 259, 264, 20, 543, 19],
        )
        self.assertEqual(set(result["gates"]), {
            "data_contract", "manifest_delivery", "original_archive", "navigation_links",
            "office_structure", "visual_evidence", "visible_text",
        })
        self.assertEqual(
            result["gates"]["visual_evidence"]["summary"]["inventory_digest"],
            "80888dde132a1b3e0ef6069458efd3f6e1f5cde813b04f97efa895e91ca6d0f2",
        )
        self.assertEqual(
            result["gates"]["navigation_links"]["summary"],
            {
                "leaf_pages": 61,
                "domain_indexes": 5,
                "total_indexes": 1,
                "markdown_files": 272,
                "links": 1261,
                "local_links": 670,
            },
        )
        self.assertEqual(
            result["gates"]["visible_text"]["summary"],
            {"text_sources": 518, "text_issues": 0},
        )
        self.assertEqual(set(result["manual_review"]["category_spot_check_ids"]), set("0102030405"[i:i+2] for i in range(0, 10, 2)))
        self.assertEqual(len(result["manual_review"]["technical_term_checks"]), 10)


if __name__ == "__main__":
    unittest.main()
