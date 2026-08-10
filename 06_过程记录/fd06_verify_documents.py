#!/usr/bin/env python3
"""Reopen and structurally verify all formal FD06 Word reports."""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
CATALOG = json.loads((ROOT / "06_过程记录" / "fd06_catalog.json").read_text(encoding="utf-8"))
DELIVERY_ROOT = ROOT / "05_交付物" / "06_课程设计、教学材料与教学评估_全网公开技能调研"
MANIFEST = json.loads((DELIVERY_ROOT / "DOCUMENT_MANIFEST.json").read_text(encoding="utf-8"))
PLACEHOLDERS = ("TODO", "TBD", "待补充", "Lorem ipsum", "占位符")


def external_hyperlinks(document: Document) -> int:
    return sum(1 for relationship in document.part.rels.values() if relationship.is_external)


def resolve_manifest_path(stored_path: str) -> Path:
    relative_path = Path(stored_path)
    assert not relative_path.is_absolute(), f"Word 清单必须使用相对路径：{stored_path}"
    resolved = (DELIVERY_ROOT / relative_path).resolve()
    assert resolved.is_relative_to(DELIVERY_ROOT.resolve()), f"Word 清单路径超出交付目录：{stored_path}"
    return resolved


def verify() -> dict:
    files = sorted(DELIVERY_ROOT.rglob("*.docx"))
    assert len(MANIFEST) == len(files) == 13, f"Word 文件数应为 13，实际 {len(files)}"
    expected_ids = {item["skill_id"] for item in CATALOG}
    seen_ids: set[str] = set()
    total_pages_unknown = True

    for item in MANIFEST:
        path = resolve_manifest_path(item["path"])
        assert path.exists(), f"缺少 Word 文件：{path}"
        document = Document(path)
        section = document.sections[0]
        assert abs(section.page_width.inches - 8.5) < 0.01
        assert abs(section.page_height.inches - 11) < 0.01
        assert abs(section.top_margin.inches - 1.0) < 0.01
        assert abs(section.bottom_margin.inches - 1.0) < 0.01
        assert abs(section.left_margin.inches - 1.0) < 0.01
        assert abs(section.right_margin.inches - 1.0) < 0.01
        assert abs(section.header_distance.inches - 0.492) < 0.01
        assert abs(section.footer_distance.inches - 0.492) < 0.01

        assert document.styles["Normal"].font.name == "Calibri"
        assert round(document.styles["Normal"].font.size.pt, 1) == 11.0
        assert round(document.styles["Heading 1"].font.size.pt, 1) == 16.0
        assert round(document.styles["Heading 2"].font.size.pt, 1) == 13.0
        assert round(document.styles["Heading 3"].font.size.pt, 1) == 12.0

        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        assert not any(marker in text for marker in PLACEHOLDERS), f"{item['key']}: 发现占位内容"
        assert "未安装" in text and "未运行" in text, f"{item['key']}: 缺少验证边界"
        assert "全网公开" in text, f"{item['key']}: 缺少来源范围"

        headings = [paragraph.text for paragraph in document.paragraphs if paragraph.style.name == "Heading 2"]
        if item["key"] == "00":
            assert len(headings) == 0, "总览不应重复 298 项详细说明"
        else:
            assert len(headings) == item["count"], f"{item['key']}: 技能标题数量不一致"
            ids = {match.group(0) for heading in headings if (match := re.match(r"FD-06-\d{4}", heading))}
            assert len(ids) == item["count"], f"{item['key']}: Skill ID 标题不完整"
            seen_ids.update(ids)
            assert external_hyperlinks(document) >= item["count"], f"{item['key']}: 固定版本链接不足"

        for table in document.tables:
            properties = table._tbl.tblPr
            width = properties.find(qn("w:tblW"))
            indent = properties.find(qn("w:tblInd"))
            layout = properties.find(qn("w:tblLayout"))
            assert width is not None and width.get(qn("w:w")) == "9360", f"{item['key']}: 表格宽度不正确"
            assert indent is not None and indent.get(qn("w:w")) == "120", f"{item['key']}: 表格缩进不正确"
            assert layout is not None and layout.get(qn("w:type")) == "fixed", f"{item['key']}: 表格未固定列宽"

        print(f"{item['key']}: reopen OK; tables={len(document.tables)}; skill_headings={len(headings)}; links={external_hyperlinks(document)}")

    assert seen_ids == expected_ids, "十二个详细报告的 Skill ID 与正式目录不一致"
    return {"docx": len(files), "skill_headings": len(seen_ids), "page_count_known": not total_pages_unknown}


if __name__ == "__main__":
    print(json.dumps(verify(), ensure_ascii=False, sort_keys=True))
