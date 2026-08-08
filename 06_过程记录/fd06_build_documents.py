#!/usr/bin/env python3
"""Build the 13 formal FD06 Word reports from the frozen catalog."""

from __future__ import annotations

import json
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
CATALOG_PATH = ROOT / "06_过程记录" / "fd06_catalog.json"
DELIVERY_ROOT = ROOT / "05_交付物" / "06_课程设计、教学材料与教学评估_全网公开技能调研"
DATE = "2026-08-09"

NAVY = "17324D"
TEAL = "187B80"
PALE_TEAL = "DDEEEF"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF3D6"
PALE_RED = "FBE5E1"
INK = "22313F"
MUTED = "5B6872"
GRID = "C9D2DA"
WHITE = "FFFFFF"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"

SUBCATEGORIES = [
    ("06-01", "课程体系、目标与能力设计"),
    ("06-02", "教学大纲、教案与课时规划"),
    ("06-03", "讲义、课件与阅读材料"),
    ("06-04", "案例、实验、讨论与课堂活动"),
    ("06-05", "作业、测验与考试命题"),
    ("06-06", "作业批改与形成性反馈"),
    ("06-07", "评分量规与评价方案"),
    ("06-08", "考试评卷、成绩分析与学情诊断"),
    ("06-09", "个性化、无障碍与多语言教学适配"),
    ("06-10", "课程质量、教学反思与持续改进"),
    ("06-11", "课程论文与毕业论文评阅"),
    ("06-12", "期刊与会议论文同行评审"),
]
SUBCATEGORY_NAMES = dict(SUBCATEGORIES)


def build_delivery_specs(catalog: list[dict]) -> list[dict]:
    specs = [
        {
            "key": "00",
            "title": "课程设计、教学材料与教学评估",
            "records": catalog,
            "output_path": DELIVERY_ROOT / "00_大分类总览.docx",
            "kind": "overview",
        }
    ]
    for code, name in SUBCATEGORIES:
        specs.append(
            {
                "key": code,
                "title": f"{code} {name}",
                "records": [item for item in catalog if item["primary_subcategory"] == code],
                "output_path": DELIVERY_ROOT / f"{code}_{name}" / f"{code}_详细说明.docx",
                "kind": "subcategory",
            }
        )
    return specs


def skill_text_blocks(skill: dict) -> dict[str, str]:
    adaptation = skill["adaptation_requirements"] or ["无额外强制改造项；采用前仍应按本校制度复核。"]
    return {
        "overview": f"{skill['plain_function']} {skill['detailed_function']}",
        "preparation": f"使用前需要准备：{skill['inputs']}",
        "result": f"通常会得到：{skill['outputs']}",
        "boundary": skill["limitations"],
        "safety": f"安全等级 {skill['security_grade']}，采用建议为“{skill['adoption_level']}”。{skill['security_plain']}",
        "adaptation": "；".join(adaptation),
        "verification": f"{skill['verification_depth']} 本报告中的结论来自固定版本静态检查，未安装、未运行。",
    }


def set_run_font(run, *, size: float = 11, bold: bool = False, italic: bool = False, color: str = INK) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell, text, *, size: float = 9.2, bold: bool = False, color: str = INK, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    set_run_font(paragraph.add_run(str(text)), size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    assert sum(widths_dxa) == 9360
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), "9360")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(indent_dxa))
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), GRID)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        row_properties = row._tr.get_or_add_trPr()
        if row_properties.find(qn("w:cantSplit")) is None:
            row_properties.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(widths_dxa[index]))
            margins = cell_properties.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell_properties.append(margins)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                margin = margins.find(qn(f"w:{side}"))
                if margin is None:
                    margin = OxmlElement(f"w:{side}")
                    margins.append(margin)
                margin.set(qn("w:w"), str(value))
                margin.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def add_hyperlink(paragraph, text: str, url: str, *, size: float = 9.2) -> None:
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(round(size * 2)))
    properties.extend([fonts, color, underline, font_size])
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(paragraph.add_run("第 "), size=9, color=MUTED)
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)
    instruction_run = paragraph.add_run()
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    instruction_run._r.append(instruction)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)
    set_run_font(paragraph.add_run(" 页"), size=9, color=MUTED)


def configure_document(document: Document, running_title: str) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run(f"高校 AI 技能库调研｜{running_title}"), size=9, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.paragraph_format.space_before = Pt(0)
    add_page_number(footer)
    document.core_properties.title = running_title
    document.core_properties.subject = "FD06 全网公开技能调研"
    document.core_properties.author = "高校 AI 技能库调研项目"
    document.core_properties.keywords = "AI skills, 教学, 安全审查, 高校"


def add_cover(document: Document, title: str, subtitle: str, count: int) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(82)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    set_run_font(kicker.add_run("UNIVERSITY AI SKILLS LIBRARY · FD06"), size=10.5, bold=True, color=TEAL)
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(12)
    title_size = 22 if len(title) >= 22 else 25 if len(title) >= 18 else 27
    set_run_font(heading.add_run(title), size=title_size, bold=True, color=NAVY)
    subheading = document.add_paragraph()
    subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheading.paragraph_format.space_after = Pt(30)
    set_run_font(subheading.add_run(subtitle), size=15, color=TEAL)
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(8)
    set_run_font(meta.add_run(f"正式收录 {count} 项｜全网公开来源｜截至 {DATE}"), size=11, bold=True, color=NAVY)
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(note.add_run("固定版本静态审查；未安装、未运行；输出必须由有权限人员复核"), size=9.5, italic=True, color=MUTED)
    document.add_page_break()


def add_label_paragraph(document: Document, label: str, text: str, *, color: str = INK, size: float = 11):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run_font(paragraph.add_run(f"{label}："), size=size, bold=True, color=NAVY)
    set_run_font(paragraph.add_run(str(text)), size=size, color=color)
    return paragraph


def add_callout(document: Document, text: str, *, fill: str = PALE_BLUE, color: str = NAVY) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.keep_together = True
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)
    set_run_font(paragraph.add_run(text), size=10.5, bold=True, color=color)


def add_summary_metrics(document: Document, records: list[dict]) -> None:
    grades = Counter(item["security_grade"] for item in records)
    platforms = Counter(item["source_label"] for item in records)
    table = document.add_table(rows=2, cols=4)
    values = [
        ("正式数量", len(records)),
        ("安全等级", f"SA {grades['SA']}｜SB {grades['SB']}｜SB-A {grades['SB-A']}"),
        ("公开来源", f"{len(platforms)} 类平台"),
        ("运行状态", "未安装、未运行"),
    ]
    for index, (label, value) in enumerate(values):
        set_cell_text(table.cell(0, index), label, size=9, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.cell(0, index), TEAL)
        set_cell_text(table.cell(1, index), value, size=10, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(table, [2340, 2340, 2340, 2340])


def add_skill_section(document: Document, skill: dict) -> None:
    blocks = skill_text_blocks(skill)
    heading = document.add_heading(f"{skill['skill_id']}｜{skill['name']}", level=2)
    heading.paragraph_format.page_break_before = False
    add_callout(document, skill["plain_function"])
    add_label_paragraph(document, "适合谁使用", "、".join(skill["audience"]))
    add_label_paragraph(document, "什么时候使用", skill["when_to_use"])
    add_label_paragraph(document, "使用前准备", skill["inputs"])
    add_label_paragraph(document, "通常会得到", skill["outputs"])
    add_label_paragraph(document, "详细功能", skill["detailed_function"])
    add_label_paragraph(document, "使用边界", blocks["boundary"])
    safety_color = PALE_RED if skill["security_grade"] == "SB-A" else PALE_GOLD if skill["security_grade"] == "SB" else PALE_TEAL
    add_callout(document, blocks["safety"], fill=safety_color, color=NAVY)
    add_label_paragraph(document, "改造要求", blocks["adaptation"])
    add_label_paragraph(document, "人工复核", skill["human_review"])
    add_label_paragraph(document, "公平、无障碍与学术诚信", f"{skill['fairness_accessibility']} {skill['academic_integrity']}")
    add_label_paragraph(document, "学生材料、论文和个人信息", f"{skill['untrusted_input']} {skill['sensitive_data']}")
    add_label_paragraph(document, "联网、账号与文件处理", f"联网：{skill['network_behavior']} 账号或密钥：{skill['credential_behavior']} 文件：{skill['file_behavior']}")
    add_label_paragraph(document, "核验说明", blocks["verification"])

    source = document.add_paragraph()
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(10)
    set_run_font(source.add_run("公开来源："), size=9.2, bold=True, color=NAVY)
    add_hyperlink(source, "打开固定版本", skill["canonical_url"], size=9.2)
    set_run_font(
        source.add_run(
            f"｜{skill['source_label']}｜维护者 {skill['maintainer']}｜版本 {skill['fixed_version']}｜许可证 {skill['license']}"
        ),
        size=9.2,
        color=MUTED,
    )


def build_overview_document(records: list[dict]) -> Document:
    document = Document()
    configure_document(document, "FD06 大分类总览")
    add_cover(document, "课程设计、教学材料与教学评估", "全网公开 AI 技能调研总览", len(records))

    document.add_heading("调研结论", level=1)
    add_label_paragraph(document, "结果", "本轮共形成 298 项正式技能，覆盖课程设计、材料制作、作业与考试、教学反馈、论文评阅和同行评审等十二类任务。")
    add_label_paragraph(document, "来源", "范围为全网可公开访问来源，包括但不限于 GitHub，并保留每项技能的固定版本地址和许可证信息。")
    add_label_paragraph(document, "准入", "正式报告只包含 SA、SB 和 SB-A。落选、重复、SC 和 SX 均留在内部过程记录，不进入正式 Excel、Word 或正式技能页。")
    add_label_paragraph(document, "验证边界", "本轮只读取说明、许可证和包内文件并开展静态安全审查，未安装、未运行，也没有向外部系统写入数据。")
    add_summary_metrics(document, records)

    document.add_heading("怎样理解安全等级", level=1)
    grade_table = document.add_table(rows=4, cols=4)
    grade_rows = [
        ("等级", "采用建议", "通俗含义", "必须注意"),
        ("SA", "可直接使用", "未发现阻断性问题。", "仍要按本校制度复核实际输出。"),
        ("SB", "需要少量调整", "需要替换少量工具、路径、模板或规则。", "完成调整并进行人工检查后再使用。"),
        ("SB-A", "需要重新改造", "原项目不能直接接入学校环境。", "先删除或改写高风险步骤，再测试和审批。"),
    ]
    for row_index, values in enumerate(grade_rows):
        for column, value in enumerate(values):
            set_cell_text(grade_table.cell(row_index, column), value, size=9.2, bold=row_index == 0, color=WHITE if row_index == 0 else INK)
            if row_index == 0:
                shade_cell(grade_table.cell(row_index, column), TEAL)
            elif column == 0:
                shade_cell(grade_table.cell(row_index, column), PALE_BLUE)
    set_table_geometry(grade_table, [800, 1800, 2800, 3960])
    repeat_table_header(grade_table.rows[0])
    add_label_paragraph(document, "共同边界", "涉及成绩、录取、学生权益、未公开论文、个人信息、版权材料或最终评审决定时，AI 只能准备草稿和核查线索，不能替代有权限人员作出决定。")

    document.add_heading("十二个小分类", level=1)
    category_table = document.add_table(rows=1, cols=5)
    for column, label in enumerate(("编号", "小分类", "正式数量", "安全等级", "采用建议")):
        set_cell_text(category_table.cell(0, column), label, size=9, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(category_table.cell(0, column), TEAL)
    repeat_table_header(category_table.rows[0])
    for code, name in SUBCATEGORIES:
        rows = [item for item in records if item["primary_subcategory"] == code]
        grades = Counter(item["security_grade"] for item in rows)
        adoption = Counter(item["adoption_level"] for item in rows)
        cells = category_table.add_row().cells
        values = (
            code,
            name,
            len(rows),
            f"SA {grades['SA']}｜SB {grades['SB']}｜SB-A {grades['SB-A']}",
            f"直接 {adoption['可直接使用']}｜少调 {adoption['需要少量调整']}｜重改 {adoption['需要重新改造']}",
        )
        for column, value in enumerate(values):
            set_cell_text(cells[column], value, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER if column in {0, 2} else WD_ALIGN_PARAGRAPH.LEFT)
        shade_cell(cells[0], PALE_BLUE)
    set_table_geometry(category_table, [900, 3300, 900, 2060, 2200])

    platforms = Counter(item["source_label"] for item in records)
    licenses = Counter(item["license"] for item in records)
    document.add_heading("来源与许可证概况", level=1)
    add_label_paragraph(document, "来源平台", "；".join(f"{name} {count} 项" for name, count in sorted(platforms.items())))
    add_label_paragraph(document, "许可证", "；".join(f"{name} {count} 项" for name, count in sorted(licenses.items(), key=lambda item: (-item[1], item[0]))))
    add_label_paragraph(document, "解释", "许可证不明确或带非商业限制的技能仍可用于调研参考，但在学校部署、对外服务、收费项目或二次发布前，必须重新确认许可条件。")

    document.add_heading("正式技能索引", level=1)
    intro = document.add_paragraph("以下索引用于快速定位；每项的详细功能、输入输出、安全边界和固定版本来源，请查看对应小分类的独立说明文档。")
    intro.paragraph_format.keep_with_next = True
    for code, name in SUBCATEGORIES:
        category_records = [item for item in records if item["primary_subcategory"] == code]
        document.add_heading(f"{code} {name}（{len(category_records)} 项）", level=3)
        index_table = document.add_table(rows=1, cols=4)
        for column, label in enumerate(("Skill ID", "技能名称", "采用建议", "安全等级")):
            set_cell_text(index_table.cell(0, column), label, size=8.6, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
            shade_cell(index_table.cell(0, column), TEAL)
        repeat_table_header(index_table.rows[0])
        for item in category_records:
            cells = index_table.add_row().cells
            values = (item["skill_id"], item["name"], item["adoption_level"], item["security_grade"])
            for column, value in enumerate(values):
                set_cell_text(cells[column], value, size=8.1, align=WD_ALIGN_PARAGRAPH.CENTER if column in {0, 2, 3} else WD_ALIGN_PARAGRAPH.LEFT)
            shade_cell(cells[0], PALE_BLUE)
        set_table_geometry(index_table, [1500, 4560, 1800, 1500])

    document.add_heading("使用建议", level=1)
    add_label_paragraph(document, "先筛任务", "从十二个小分类中选择与当前工作最接近的一类，再在独立清单中比较采用建议、输入输出、许可证和来源。")
    add_label_paragraph(document, "再看风险", "优先选择 SA；SB 先完成少量调整；SB-A 只能进入改造评估，不能直接接入真实学生、课程或论文数据。")
    add_label_paragraph(document, "最后复核", "所有生成内容都应由教师、导师、教学管理人员、审稿人或编辑中的有权限人员确认，保留判断依据和必要的申诉、纠错与审计记录。")
    return document


def build_subcategory_document(code: str, title: str, records: list[dict]) -> Document:
    document = Document()
    configure_document(document, title)
    add_cover(document, title, "全网公开 AI 技能详细说明", len(records))

    document.add_heading("本小分类概览", level=1)
    add_label_paragraph(document, "收录范围", f"本报告聚焦“{SUBCATEGORY_NAMES[code]}”，共收录 {len(records)} 项通过静态安全准入的正式技能。")
    add_label_paragraph(document, "如何使用", "先按自己的任务阅读“什么时候使用、使用前准备、通常会得到”，再核对采用建议、安全等级、改造要求和许可证。")
    add_label_paragraph(document, "共同边界", "正式技能只能辅助准备草稿、材料、反馈和核查线索。涉及成绩、学生权益、未公开论文、个人信息或最终评价决定时，必须由有权限人员复核。")
    add_summary_metrics(document, records)

    grades = Counter(item["security_grade"] for item in records)
    platforms = Counter(item["source_label"] for item in records)
    licenses = Counter(item["license"] for item in records)
    document.add_heading("结构与来源", level=1)
    add_label_paragraph(document, "安全等级", f"SA {grades['SA']} 项；SB {grades['SB']} 项；SB-A {grades['SB-A']} 项。")
    add_label_paragraph(document, "来源平台", "；".join(f"{name} {count} 项" for name, count in sorted(platforms.items())))
    add_label_paragraph(document, "许可证", "；".join(f"{name} {count} 项" for name, count in sorted(licenses.items(), key=lambda item: (-item[1], item[0]))))

    document.add_heading("正式技能详细说明", level=1)
    introduction = document.add_paragraph("以下条目按固定 Skill ID 排序。每项都给出通俗功能、适用对象、输入输出、采用建议、具体安全边界、改造要求和固定版本来源。")
    introduction.paragraph_format.keep_with_next = True
    for skill in records:
        add_skill_section(document, skill)

    document.add_heading("本小分类使用建议", level=1)
    add_label_paragraph(document, "选型顺序", "优先比较任务匹配程度和安全等级，再看许可证、来源维护情况以及是否需要联网、账号、外部写入或文件处理。")
    add_label_paragraph(document, "运行验证", "本报告不包含运行效果。若后续需要验证某一技能，应由用户明确指定，再设计可回滚、最少数据、无外部写入或经授权写入的测试方案。")
    return document


def build_documents() -> list[dict]:
    catalog = json.loads(CATALOG_PATH.read_text(encoding="utf-8"))
    manifest = []
    for spec in build_delivery_specs(catalog):
        if spec["kind"] == "overview":
            document = build_overview_document(spec["records"])
        else:
            document = build_subcategory_document(spec["key"], spec["title"], spec["records"])
        output_path = spec["output_path"]
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        manifest.append({"key": spec["key"], "title": spec["title"], "count": len(spec["records"]), "path": str(output_path)})
        print(f"{spec['key']}: {len(spec['records'])} skills -> {output_path}")
    (DELIVERY_ROOT / "DOCUMENT_MANIFEST.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return manifest


if __name__ == "__main__":
    build_documents()
