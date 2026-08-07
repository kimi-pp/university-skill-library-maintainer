"""生成 0809 计算机类跨平台开源 Skill 详细调研报告。"""

from __future__ import annotations

import json
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
DATA_PATH = PROJECT_ROOT / "03_候选池" / "deduplicated" / "0809_computer_science.json"
OUTPUT_PATH = PROJECT_ROOT / "05_交付物" / "0809_计算机类_跨平台技能调研.docx"

NAVY = "18324A"
TEAL = "167D89"
BLUE = "2E74B5"
LIGHT_TEAL = "E8F3F4"
PALE_BLUE = "EAF2F8"
PALE = "F6F8FA"
INK = "18212B"
MUTED = "536273"
GRID = "C9D3DC"
WHITE = "FFFFFF"
GREEN = "DDF3E4"
GREEN_INK = "176B3A"
AMBER = "FFF0C7"
AMBER_INK = "8A5A00"
ORANGE = "FBE1D1"
ORANGE_INK = "9A3412"

GROUP_ORDER = ["A 计算基础", "B 软件工程", "C 网络安全", "D 物联网", "E AI与数据", "F 空间信息", "G 数字媒体", "H 保密治理"]
MAJORS = [
    ("080901K", "计算机科学与技术"), ("080902", "软件工程"), ("080903", "网络工程"),
    ("080904K", "信息安全"), ("080905T", "物联网工程"), ("080906T", "数字媒体技术"),
    ("080907T", "智能科学与技术"), ("080908T", "空间信息与数字技术"),
    ("080909T", "电子与计算机工程"), ("080910T", "数据科学与大数据技术"),
    ("080911T", "网络空间安全"), ("080912T", "新媒体技术"), ("080913T", "电影制作"),
    ("080914T", "保密管理"),
]


def set_run_font(run, size=11, bold=False, italic=False, color=INK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill, left=120, right=120, top=70, bottom=70):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left_border = OxmlElement("w:left")
    left_border.set(qn("w:val"), "single")
    left_border.set(qn("w:sz"), "18")
    left_border.set(qn("w:space"), "6")
    left_border.set(qn("w:color"), TEAL)
    borders.append(left_border)
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), str(top))
    spacing.set(qn("w:after"), str(bottom))
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:right"), str(right))


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, size=9.2, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1.5)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), GRID)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                margin = margins.find(qn(f"w:{side}"))
                if margin is None:
                    margin = OxmlElement(f"w:{side}")
                    margins.append(margin)
                margin.set(qn("w:w"), str(value))
                margin.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(repeat)


def add_hyperlink(paragraph, text, url, size=9.2):
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(size * 2)))
    r_pr.extend([fonts, color, underline, size_node])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lead = paragraph.add_run("0809 计算机类跨平台技能调研  ·  ")
    set_run_font(lead, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    field_run.extend([begin, instr, separate, value, end])
    paragraph._p.append(field_run)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("高校 AI 技能库调研  |  0809 计算机类  |  跨平台开源"), size=8.5, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_cover(document, count):
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(78)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(15)
    set_run_font(kicker.add_run("UNIVERSITY AI SKILLS LIBRARY · PILOT REPORT"), size=10, bold=True, color=TEAL)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("0809 计算机类"), size=30, bold=True, color=NAVY)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    set_run_font(subtitle.add_run("跨平台开源 Skill 初步调研报告"), size=16, color=BLUE)
    accent = document.add_paragraph()
    accent.alignment = WD_ALIGN_PARAGRAPH.CENTER
    accent.paragraph_format.space_after = Pt(26)
    set_run_font(accent.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━"), size=10, color=TEAL)
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(8)
    set_run_font(meta.add_run(f"正式候选 {count} 项  |  14 个专业  |  8 个能力群"), size=11, bold=True, color=NAVY)
    meta2 = document.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta2.paragraph_format.space_after = Pt(22)
    set_run_font(meta2.add_run("审查日期：2026-08-07  |  方法：说明读取、拆包与静态审查"), size=9.5, color=MUTED)
    warning = document.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning.paragraph_format.space_before = Pt(10)
    warning.paragraph_format.space_after = Pt(0)
    shade_paragraph(warning, AMBER)
    set_run_font(warning.add_run("未安装、未运行候选包；SB-A 原包不可直接接入"), size=10, bold=True, color=AMBER_INK)
    document.add_page_break()


def add_static_contents(document):
    document.add_heading("目录与阅读路径", level=1)
    intro = document.add_paragraph("本页提供静态导航；正式条目按能力群组织，条目标题含稳定 ID，便于与 Excel 和 Markdown 索引互相定位。")
    intro.paragraph_format.space_after = Pt(10)
    sections = [
        "1  执行摘要", "2  调研范围与安全准入", "3  专业与能力群覆盖",
        "4  正式候选详细说明", "5  来源仓库与版本固定", "6  结论与打样建议", "附录  方法与结论边界",
    ]
    for item in sections:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.15)
        paragraph.paragraph_format.space_after = Pt(4)
        set_run_font(paragraph.add_run(item), size=11, bold=True, color=NAVY)
    sub = document.add_paragraph()
    sub.paragraph_format.left_indent = Inches(0.35)
    sub.paragraph_format.space_before = Pt(4)
    sub.paragraph_format.space_after = Pt(0)
    set_run_font(sub.add_run("能力群：A 计算基础 · B 软件工程 · C 网络安全 · D 物联网 · E AI与数据 · F 空间信息 · G 数字媒体 · H 保密治理"), size=9.5, color=MUTED)
    document.add_page_break()


def add_label_paragraph(document, label, text, color=INK, keep=False):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.22
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = keep
    set_run_font(paragraph.add_run(f"{label}："), size=10.2, bold=True, color=NAVY)
    set_run_font(paragraph.add_run(str(text)), size=10.2, color=color)
    return paragraph


def add_callout(document, text, grade=None):
    fill, ink = PALE_BLUE, NAVY
    if grade == "SA": fill, ink = GREEN, GREEN_INK
    elif grade == "SB": fill, ink = AMBER, AMBER_INK
    elif grade == "SB-A": fill, ink = ORANGE, ORANGE_INK
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = True
    shade_paragraph(paragraph, fill)
    set_run_font(paragraph.add_run(text), size=10.2, bold=True, color=ink)
    return paragraph


def add_summary_table(document, records):
    security = Counter(row["security_grade"] for row in records)
    priority = Counter(row["priority"] for row in records)
    platform = Counter(row["platform"] for row in records)
    values = [
        ("正式候选", len(records)), ("安全 SA / SB / SB-A", f"{security['SA']} / {security['SB']} / {security['SB-A']}"),
        ("高优先级", priority["高"]), ("来源平台标注", len(platform)),
    ]
    table = document.add_table(rows=2, cols=4)
    for column, (label, value) in enumerate(values):
        set_cell_text(table.cell(0, column), label, size=8.8, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.cell(0, column), TEAL)
        set_cell_text(table.cell(1, column), value, size=11, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(table, [2340, 2340, 2340, 2340])


def add_security_table(document):
    rows = [
        ("等级", "准入形式", "原包边界", "最低控制"),
        ("SA", "低风险知识/流程包", "可作为候选，但不是绝对安全证明", "按校内制度与版本复核"),
        ("SB", "限制使用", "只在明确限制下采用", "隔离目录、固定依赖、最小权限、人工确认"),
        ("SB-A", "仅适配后使用", "原包不可安装或整体复制", "剥离自动联网、凭据、上传、发布、部署和高权限路径"),
    ]
    table = document.add_table(rows=len(rows), cols=4)
    for r_idx, values in enumerate(rows):
        for c_idx, value in enumerate(values):
            color = WHITE if r_idx == 0 else INK
            set_cell_text(table.cell(r_idx, c_idx), value, size=8.8, bold=r_idx == 0 or c_idx == 0, color=color)
            if r_idx == 0:
                shade_cell(table.cell(r_idx, c_idx), TEAL)
            elif c_idx == 0:
                shade_cell(table.cell(r_idx, c_idx), PALE_BLUE)
    set_table_geometry(table, [900, 1900, 2500, 4060])
    repeat_table_header(table.rows[0])


def add_major_coverage_table(document, records):
    rows = [("专业代码", "专业名称", "候选数", "高优先级", "SA / SB / SB-A")]
    for code, major in MAJORS:
        selected = [row for row in records if major in row["majors"]]
        grades = Counter(row["security_grade"] for row in selected)
        rows.append((code, major, len(selected), sum(row["priority"] == "高" for row in selected), f"{grades['SA']} / {grades['SB']} / {grades['SB-A']}"))
    table = document.add_table(rows=len(rows), cols=5)
    for r_idx, values in enumerate(rows):
        for c_idx, value in enumerate(values):
            set_cell_text(table.cell(r_idx, c_idx), value, size=8.4, bold=r_idx == 0, color=WHITE if r_idx == 0 else INK, align=WD_ALIGN_PARAGRAPH.CENTER if c_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT)
            if r_idx == 0:
                shade_cell(table.cell(r_idx, c_idx), TEAL)
            elif r_idx % 2 == 0:
                shade_cell(table.cell(r_idx, c_idx), PALE)
    set_table_geometry(table, [1050, 2870, 1100, 1250, 3090])
    repeat_table_header(table.rows[0])


def add_group_summary_table(document, records):
    counts = Counter(row["primary_group"] for row in records)
    table = document.add_table(rows=1 + len(GROUP_ORDER), cols=3)
    for column, label in enumerate(("能力群", "候选数", "主要支撑方向")):
        set_cell_text(table.cell(0, column), label, size=9, bold=True, color=WHITE)
        shade_cell(table.cell(0, column), TEAL)
    descriptions = {
        "A 计算基础": "资源、图网络、仿真、符号计算与形式化基础",
        "B 软件工程": "需求、设计、实现、测试、评审、调试与持续交付",
        "C 网络安全": "威胁建模、攻防验证、漏洞与基础设施安全",
        "D 物联网": "嵌入式、Arduino、机器人、设备与系统集成",
        "E AI与数据": "机器学习、数据处理、训练、评估与可视化",
        "F 空间信息": "地理空间计算、遥感与空间数据工作流",
        "G 数字媒体": "前端、交互、图像、视频、3D 与内容制作",
        "H 保密治理": "合规、隐私、秘密与凭据治理",
    }
    for index, group in enumerate(GROUP_ORDER, start=1):
        values = (group, counts[group], descriptions[group])
        for column, value in enumerate(values):
            set_cell_text(table.cell(index, column), value, size=8.8, bold=column == 0)
            if index % 2 == 0:
                shade_cell(table.cell(index, column), PALE)
    set_table_geometry(table, [2100, 1100, 6160])
    repeat_table_header(table.rows[0])


def add_skill_section(document, row):
    heading = document.add_heading(f"{row['id']}｜{row['cn_name']}（{row['name']}）", level=3)
    heading.paragraph_format.page_break_before = False
    add_callout(document, row["summary"], row["security_grade"])
    source = document.add_paragraph()
    source.paragraph_format.space_after = Pt(5)
    source.paragraph_format.keep_together = True
    source.paragraph_format.keep_with_next = True
    set_run_font(source.add_run("来源："), size=9.5, bold=True, color=NAVY)
    set_run_font(source.add_run(f"{row['platform']} · {row['repo']} · {row['ecosystem']}  |  "), size=9.5, color=INK)
    add_hyperlink(source, "打开固定版本 Skill 源文件", row["skill_url"], size=9.5)
    set_run_font(source.add_run(f"  |  提交 {row['review_commit'][:12]}"), size=9.2, color=MUTED)

    function_text = f"{row['summary']} 适合用于{row['scenarios'].replace('0809 计算机类', '本学科')}。定位为“{row['coverage_type']}”，可服务{row['roles']}。"
    add_label_paragraph(document, "功能与使用", function_text, keep=True)
    add_label_paragraph(document, "专业覆盖", "、".join(row["majors"]), keep=True)
    evidence_text = f"审查到包文件 {row['package_files']} 个、脚本文件 {row['script_files']} 个、二进制文件 {row['binary_files']} 个；{row['executable_behavior']}；{row['network_data_behavior']}。"
    add_label_paragraph(document, "静态证据", evidence_text, keep=True)
    safety_text = f"{row['security_grade']}｜{row['admission_form']}。{row['security_restrictions']}"
    add_label_paragraph(document, "安全结论", safety_text, color=ORANGE_INK if row["security_grade"] == "SB-A" else INK, keep=True)
    adapt_text = f"保留：{row['adapt_keep']}；剥离/禁用：{row['adapt_strip']}。"
    add_label_paragraph(document, "适配边界", adapt_text, keep=True)
    add_label_paragraph(document, "依赖与验证", f"{row['dependencies']} {row['verification']}")


def add_sources_table(document, records):
    repo_records = defaultdict(list)
    for row in records:
        repo_records[row["repo"]].append(row)
    table = document.add_table(rows=1, cols=6)
    for column, label in enumerate(("平台", "仓库", "入选数", "许可证", "固定提交", "提交日期")):
        set_cell_text(table.cell(0, column), label, size=8.4, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.cell(0, column), TEAL)
    repeat_table_header(table.rows[0])
    for index, repo in enumerate(sorted(repo_records), start=1):
        sample = repo_records[repo][0]
        cells = table.add_row().cells
        values = (sample["platform"], repo, len(repo_records[repo]), sample["license"], sample["review_commit"], sample["review_commit_date"])
        for column, value in enumerate(values):
            set_cell_text(cells[column], value, size=7.9, align=WD_ALIGN_PARAGRAPH.CENTER if column in (0, 2, 5) else WD_ALIGN_PARAGRAPH.LEFT)
            if index % 2 == 0:
                shade_cell(cells[column], PALE)
    set_table_geometry(table, [1200, 2380, 850, 1350, 2580, 1000])


def build():
    payload = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    records = payload["records"]
    document = Document()
    configure_document(document)
    add_cover(document, len(records))
    add_static_contents(document)

    document.add_heading("1  执行摘要", level=1)
    add_label_paragraph(document, "结论", "本轮在公开开源平台中形成 88 项正式候选，覆盖 0809 计算机类全部 14 个专业和 8 个能力群。候选并非全部可直接接入：27 项为 SA，35 项为 SB，26 项为 SB-A。")
    add_label_paragraph(document, "跨平台范围", "正式来源包含 GitHub、GitLab、Gitee，以及由 Hugging Face 官方维护但上游托管在 GitHub 的技能包。Codeberg、SourceHut 等平台也进入检索矩阵，但没有为平台多样性而降低准入标准。")
    add_label_paragraph(document, "安全结论", "静态审查只说明当前固定提交中已发现的行为与准入边界。它不能证明运行时绝对安全；部署前仍需重新核验依赖、权限、网络、凭据、数据流和外部状态改变。")
    add_summary_table(document, records)

    document.add_heading("2  调研范围与安全准入", level=1)
    add_label_paragraph(document, "对象", "0809 计算机类 14 个本科专业。本报告不外延到其他一级学科；下一轮范围继续由用户按任务指定。")
    add_label_paragraph(document, "正式收录", "只纳入功能相关、来源可定位、包结构可读取且静态风险可被明确约束的候选。落选项目不进入本报告和配套 Excel，只留存在内部过程记录。")
    add_label_paragraph(document, "验证方式", "逐项读取 SKILL.md 或等价入口，枚举包内文件，检查脚本、依赖、网络访问、凭据语境、文件写入/删除、子进程、高权限动作和外部状态改变。未安装、未运行任何候选脚本。")
    add_security_table(document)
    warning = document.add_paragraph()
    shade_paragraph(warning, ORANGE)
    set_run_font(warning.add_run("SB-A 的“安全”只属于适配版：必须保留明确的保留清单与剥离清单；原包不得直接安装、整体复制或按原说明执行。"), size=10, bold=True, color=ORANGE_INK)

    document.add_heading("3  专业与能力群覆盖", level=1)
    add_label_paragraph(document, "解释", "一项 Skill 可以服务多个专业，因此专业覆盖数量之和会高于正式候选总数。覆盖代表教学、课程设计、科研开发或高校技术岗位中的潜在适用性，不代表所有专业都应部署同一套默认触发器。")
    add_major_coverage_table(document, records)
    document.add_heading("3.1  能力群结构", level=2)
    add_group_summary_table(document, records)

    document.add_page_break()
    document.add_heading("4  正式候选详细说明", level=1)
    intro = document.add_paragraph("以下 88 项按能力群和稳定 ID 排序。每项均给出功能定位、来源、专业覆盖、静态证据、安全等级、适配边界与验证限制；地址固定到本轮审查提交。")
    intro.paragraph_format.keep_with_next = True
    by_group = defaultdict(list)
    for record in records:
        by_group[record["primary_group"]].append(record)
    for group in GROUP_ORDER:
        document.add_heading(f"4.{GROUP_ORDER.index(group) + 1}  {group}", level=2)
        group_records = by_group[group]
        add_label_paragraph(document, "本组概览", f"共 {len(group_records)} 项。该组条目按稳定 ID 排列；部署时应按具体课程或岗位任务选取最小组合，避免多个相邻 Skill 重复触发。")
        for record in group_records:
            add_skill_section(document, record)

    document.add_page_break()
    document.add_heading("5  来源仓库与版本固定", level=1)
    add_label_paragraph(document, "版本策略", "每项地址与仓库提交均固定到 2026-08-07 本轮审查所见版本。上游更新、许可证变化或依赖变化后，当前安全结论自动失效，需要重新审查。")
    add_sources_table(document, records)

    document.add_heading("6  结论与打样建议", level=1)
    high = [row for row in records if row["priority"] == "高"]
    add_label_paragraph(document, "打样价值", "0809 计算机类能够同时检验通用软件工程技能、专业工具链技能、硬件/设备工作流、数据与媒体能力，以及保密和合规边界，适合作为后续学科调研的结构样板。")
    add_label_paragraph(document, "优先顺序", f"先从 {len(high)} 项高优先级候选中选择 SA 与可清晰隔离的 SB 做制度和场景评审；SB-A 先产出静态适配稿，不进入运行试点。")
    add_label_paragraph(document, "部署前门槛", "确认许可证与上游维护状态；固定依赖和哈希；采用隔离目录、最小权限、网络白名单和人工确认；学生作业、员工数据、科研材料与管理数据按校内分级制度处理。")
    add_label_paragraph(document, "下一步", "本报告只完成打样学科的静态调研。若用户指定某项候选进入验证，再单独设计最小、可回滚、无敏感数据、无默认外部写入的验证任务。")

    document.add_heading("附录  方法与结论边界", level=1)
    add_label_paragraph(document, "审查证据", "入口说明、包结构、脚本和静态命中保存在项目知识库与验证记录中；配套 Excel 提供 88 项机器可筛选清单、来源地址、专业映射和公式统计。")
    add_label_paragraph(document, "未做事项", "未安装候选、未解析其依赖树到所有传递依赖、未调用外部 API、未访问账号、未写入第三方系统、未运行硬件或模型训练，也未声称真实任务效果。")
    add_label_paragraph(document, "安全保证边界", "静态审查能发现显式风险，但无法排除运行环境差异、依赖供应链、动态下载、上游账户被接管、外部服务行为或恶意输入导致的新风险。")
    add_label_paragraph(document, "复用建议", "后续一级学科继续复用“范围—来源—静态安全—专业映射—Excel 清单—Word 详解—Markdown 索引”的结构；候选数量、平台和风险阈值仍按每轮任务单独确定。")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(f"{len(records)} skills -> {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
