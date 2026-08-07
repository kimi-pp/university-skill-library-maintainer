"""对最终 DOCX 做内容完整性、版式 token 和 OOXML 结构审计。"""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


project_root = Path(__file__).resolve().parents[2]
data_dir = project_root / "03_候选池" / "deduplicated"
output_dir = project_root / "05_交付物"
manifest = json.loads((data_dir / "manifest.json").read_text(encoding="utf-8"))

requested_categories = tuple(sys.argv[1:])
known_categories = tuple(dict.fromkeys(item["category"] for item in manifest))
verify_categories = requested_categories or known_categories
invalid_categories = set(verify_categories) - set(known_categories)
if invalid_categories:
    raise SystemExit(f"未知分类：{', '.join(sorted(invalid_categories))}")

for category in verify_categories:
    payload = json.loads((data_dir / f"category_{category}.json").read_text(encoding="utf-8"))
    expected_ids = [row["id"] for row in payload["records"]]
    item = next(entry for entry in manifest if entry["category"] == category and entry["format"] == "docx")
    document = Document(output_dir / item["path"])

    all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    heading_ids = [
        paragraph.text.split("｜", 1)[0]
        for paragraph in document.paragraphs
        if paragraph.style.name == "Heading 2" and paragraph.text.startswith("GH-")
    ]
    assert heading_ids == expected_ids, f"{category}: Heading 2 Skill ID 不完整或顺序错误"
    assert all(token not in all_text for token in ("TODO", "TBD", "{{", "}}", "PLACEHOLDER")), f"{category}: 存在占位符"
    assert "未安装" in all_text and "未运行" in all_text, f"{category}: 缺少验证边界说明"

    section = document.sections[0]
    assert section.page_width.twips == 12240 and section.page_height.twips == 15840, f"{category}: 页面不是 Letter"
    assert all(
        value.twips == 1440
        for value in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin)
    ), f"{category}: 页边距不是 1 英寸"
    assert section.header_distance.twips in (708, 709), f"{category}: 页眉距离不符合 preset"
    assert section.footer_distance.twips in (708, 709), f"{category}: 页脚距离不符合 preset"

    normal = document.styles["Normal"]
    assert normal.font.name == "Calibri" and round(normal.font.size.pt, 2) == 11, f"{category}: Normal 字体错误"
    assert round(normal.paragraph_format.space_after.pt, 2) == 6, f"{category}: Normal 段后错误"
    assert float(normal.paragraph_format.line_spacing) == 1.25, f"{category}: Normal 行距错误"
    expected_heading_tokens = {
        "Heading 1": (16, "2E74B5", 18, 10),
        "Heading 2": (13, "2E74B5", 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
    }
    for style_name, (size, color, before, after) in expected_heading_tokens.items():
        style = document.styles[style_name]
        assert round(style.font.size.pt, 2) == size, f"{category}: {style_name} 字号错误"
        assert str(style.font.color.rgb) == color, f"{category}: {style_name} 颜色错误"
        assert round(style.paragraph_format.space_before.pt, 2) == before, f"{category}: {style_name} 段前错误"
        assert round(style.paragraph_format.space_after.pt, 2) == after, f"{category}: {style_name} 段后错误"

    assert len(document.tables) == len(expected_ids) + 3, f"{category}: 表格数量不符合报告结构"
    for table_index, table in enumerate(document.tables):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        assert tbl_w is not None and int(tbl_w.get(qn("w:w"))) == 9360, f"{category}: 表 {table_index} 宽度错误"
        assert tbl_ind is not None and int(tbl_ind.get(qn("w:w"))) == 120, f"{category}: 表 {table_index} 缩进错误"
        grid_widths = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
        assert sum(grid_widths) == 9360, f"{category}: 表 {table_index} 网格宽度总和错误"
        for row in table.rows:
            for column, cell in enumerate(row.cells):
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                assert tc_w is not None and int(tc_w.get(qn("w:w"))) == grid_widths[column], f"{category}: 表 {table_index} 单元格宽度漂移"

    hyperlinks = [relationship for relationship in document.part.rels.values() if relationship.reltype.endswith("/hyperlink")]
    unique_repositories = len({row["repo"] for row in payload["records"]})
    assert len(hyperlinks) >= len(expected_ids) + unique_repositories, f"{category}: GitHub 超链接数量不足"
    print(
        f"{category}: docx reopen OK; skill_headings={len(heading_ids)}; "
        f"tables={len(document.tables)}; hyperlinks={len(hyperlinks)}; preset audit OK"
    )
