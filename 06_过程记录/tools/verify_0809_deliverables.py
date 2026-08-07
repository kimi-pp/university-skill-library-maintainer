"""校验 0809 DOCX 的内容完整性、样式令牌和 OOXML 结构。"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
DATA = json.loads((ROOT / "03_候选池" / "deduplicated" / "0809_computer_science.json").read_text(encoding="utf-8"))
DOCX = ROOT / "05_交付物" / "0809_计算机类_跨平台技能调研.docx"

document = Document(DOCX)
records = DATA["records"]
expected_ids = [row["id"] for row in records]
all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
actual_ids = [
    paragraph.text.split("｜", 1)[0]
    for paragraph in document.paragraphs
    if paragraph.style.name == "Heading 3" and paragraph.text.startswith("DISC-0809-")
]

assert actual_ids == expected_ids, "Heading 3 的 88 个 Skill ID 不完整或顺序错误"
assert all(token not in all_text for token in ("TODO", "TBD", "{{", "}}", "PLACEHOLDER")), "发现占位符"
for required in ("未安装", "未运行", "SB-A", "仅适配后使用", "14 个专业", "8 个能力群"):
    assert required in all_text, f"缺少关键边界：{required}"
for row in records:
    assert row["cn_name"] in all_text, f"缺少中文名称：{row['id']}"
    assert row["admission_form"] in all_text, f"缺少准入形式：{row['id']}"

section = document.sections[0]
assert section.page_width.twips == 12240 and section.page_height.twips == 15840, "页面不是 Letter"
assert all(value.twips == 1440 for value in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin)), "页边距不是 1 英寸"
assert section.header_distance.twips in (708, 709), "页眉距离不符合 preset"
assert section.footer_distance.twips in (708, 709), "页脚距离不符合 preset"

normal = document.styles["Normal"]
assert normal.font.name == "Calibri" and round(normal.font.size.pt, 2) == 11, "Normal 字体错误"
assert round(normal.paragraph_format.space_after.pt, 2) == 6, "Normal 段后错误"
assert float(normal.paragraph_format.line_spacing) == 1.25, "Normal 行距错误"
expected_heading_tokens = {
    "Heading 1": (16, "2E74B5", 18, 10),
    "Heading 2": (13, "2E74B5", 14, 7),
    "Heading 3": (12, "1F4D78", 10, 5),
}
for style_name, (size, color, before, after) in expected_heading_tokens.items():
    style = document.styles[style_name]
    assert round(style.font.size.pt, 2) == size, f"{style_name} 字号错误"
    assert str(style.font.color.rgb) == color, f"{style_name} 颜色错误"
    assert round(style.paragraph_format.space_before.pt, 2) == before, f"{style_name} 段前错误"
    assert round(style.paragraph_format.space_after.pt, 2) == after, f"{style_name} 段后错误"

assert len(document.tables) == 5, f"报告应有 5 张紧凑汇总表，实际 {len(document.tables)}"
for table_index, table in enumerate(document.tables):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    assert tbl_w is not None and int(tbl_w.get(qn("w:w"))) == 9360, f"表 {table_index} 宽度错误"
    assert tbl_ind is not None and int(tbl_ind.get(qn("w:w"))) == 120, f"表 {table_index} 缩进错误"
    grid_widths = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
    assert sum(grid_widths) == 9360, f"表 {table_index} 网格宽度错误"
    for row in table.rows:
        for column, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            assert tc_w is not None and int(tc_w.get(qn("w:w"))) == grid_widths[column], f"表 {table_index} 单元格宽度漂移"

hyperlinks = [rel for rel in document.part.rels.values() if rel.reltype.endswith("/hyperlink")]
assert len(hyperlinks) >= len(records), "Skill 超链接不足"
assert all(row["skill_url"] in {rel.target_ref for rel in hyperlinks} for row in records), "固定版本 Skill 地址未全部写入"

print(
    f"DOCX reopen OK; skills={len(actual_ids)}; tables={len(document.tables)}; "
    f"hyperlinks={len(hyperlinks)}; preset/IDs/security boundaries OK"
)
