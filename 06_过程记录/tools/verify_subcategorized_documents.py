"""Structurally verify manifest-driven plain-language DOCX reports."""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


PROJECT_ROOT = Path(__file__).resolve().parents[2]
TOOLS_DIR = Path(__file__).resolve().parent
if str(TOOLS_DIR) not in sys.path:
    sys.path.insert(0, str(TOOLS_DIR))

from build_subcategorized_documents import (  # noqa: E402
    ASSIGNMENT_FILE,
    DOCX_TOKENS,
    MANIFEST_FILE,
    PLAIN_CATALOG_FILE,
    PLAIN_FIELDS,
    REPOSITORIES_FILE,
    _knowledge_base_target,
    _parse_only,
    select_manifest_items,
    validate_manifest_contract,
    validate_source_contract,
)


FORBIDDEN_PLACEHOLDERS = re.compile(r"\b(?:TODO|TBD|PLACEHOLDER)\b", re.IGNORECASE)
FORBIDDEN_INTERNAL_REFERENCES = (
    "render_docx.py",
    "build_subcategorized_documents.py",
    "verify_subcategorized_documents.py",
    "render_subcategorized_documents.py",
    ".codex",
    "C:\\Users\\",
)
SKILL_ID_PATTERN = re.compile(r"GH-\d{2}-\d{4}")


def _load_json(path: Path) -> object:
    return json.loads(path.read_text(encoding="utf-8"))


def _document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for section in document.sections:
        parts.extend(paragraph.text for paragraph in section.header.paragraphs)
        parts.extend(paragraph.text for paragraph in section.footer.paragraphs)
    for table in document.tables:
        parts.extend(cell.text for row in table.rows for cell in row.cells)
    return "\n".join(parts)


def _font_name(style, attribute: str) -> str | None:
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    return fonts.get(qn(f"w:{attribute}"))


def _audit_styles(document: Document) -> list[str]:
    issues: list[str] = []
    normal = document.styles["Normal"]
    if normal.font.name != DOCX_TOKENS["body_font_western"]:
        issues.append("Normal 西文字体不符合 Calibri")
    if _font_name(normal, "eastAsia") != DOCX_TOKENS["body_font_east_asia"]:
        issues.append("Normal 中文字体不符合微软雅黑")
    if normal.font.size is None or round(normal.font.size.pt, 3) != DOCX_TOKENS["body_size_pt"]:
        issues.append("Normal 字号不符合 11 pt")
    if normal.paragraph_format.space_after is None or round(normal.paragraph_format.space_after.pt, 3) != DOCX_TOKENS["body_after_pt"]:
        issues.append("Normal 段后间距不符合 6 pt")
    if normal.paragraph_format.line_spacing != DOCX_TOKENS["body_line_spacing"]:
        issues.append("Normal 行距不符合 1.25 倍")

    for style_name, token_name in (
        ("Heading 1", "heading_1"),
        ("Heading 2", "heading_2"),
        ("Heading 3", "heading_3"),
    ):
        expected_size, expected_color, expected_before, expected_after = DOCX_TOKENS[token_name]
        style = document.styles[style_name]
        actual_color = str(style.font.color.rgb) if style.font.color and style.font.color.rgb else None
        checks = (
            (style.font.name == DOCX_TOKENS["body_font_western"], f"{style_name} 西文字体错误"),
            (_font_name(style, "eastAsia") == DOCX_TOKENS["body_font_east_asia"], f"{style_name} 中文字体错误"),
            (style.font.size is not None and round(style.font.size.pt, 3) == expected_size, f"{style_name} 字号错误"),
            (actual_color == expected_color, f"{style_name} 颜色错误"),
            (style.paragraph_format.space_before is not None and round(style.paragraph_format.space_before.pt, 3) == expected_before, f"{style_name} 段前错误"),
            (style.paragraph_format.space_after is not None and round(style.paragraph_format.space_after.pt, 3) == expected_after, f"{style_name} 段后错误"),
            (bool(style.paragraph_format.keep_with_next), f"{style_name} 缺少与下段同页保护"),
        )
        issues.extend(message for passed, message in checks if not passed)
    return issues


def _audit_sections(document: Document) -> list[str]:
    issues: list[str] = []
    for index, section in enumerate(document.sections, start=1):
        if (section.page_width.twips, section.page_height.twips) != (
            DOCX_TOKENS["page_width_twips"],
            DOCX_TOKENS["page_height_twips"],
        ):
            issues.append(f"第 {index} 节不是 Letter 纵向页面")
        for label, value in (
            ("上边距", section.top_margin),
            ("右边距", section.right_margin),
            ("下边距", section.bottom_margin),
            ("左边距", section.left_margin),
        ):
            if value.twips != DOCX_TOKENS["margin_twips"]:
                issues.append(f"第 {index} 节{label}不是 1440 DXA")
        if abs(section.header_distance.twips - DOCX_TOKENS["header_footer_distance_twips"]) > 1:
            issues.append(f"第 {index} 节页眉距离错误")
        if abs(section.footer_distance.twips - DOCX_TOKENS["header_footer_distance_twips"]) > 1:
            issues.append(f"第 {index} 节页脚距离错误")
        header_text = "".join(paragraph.text for paragraph in section.header.paragraphs).strip()
        if not header_text or "|" in header_text or "GitHub" in header_text:
            issues.append(f"第 {index} 节页眉不是单一分类名称")
        if "PAGE" not in section.footer._element.xml:
            issues.append(f"第 {index} 节页脚缺少 PAGE 字段")
    return issues


def _audit_tables(document: Document) -> list[str]:
    issues: list[str] = []
    for table_index, table in enumerate(document.tables, start=1):
        properties = table._tbl.tblPr
        table_width = properties.find(qn("w:tblW"))
        table_indent = properties.find(qn("w:tblInd"))
        layout = properties.find(qn("w:tblLayout"))
        if table_width is None or table_width.get(qn("w:type")) != "dxa" or table_width.get(qn("w:w")) != str(DOCX_TOKENS["content_width_dxa"]):
            issues.append(f"表格 {table_index} 宽度不是固定 9360 DXA")
        if table_indent is None or table_indent.get(qn("w:type")) != "dxa" or table_indent.get(qn("w:w")) != str(DOCX_TOKENS["table_indent_dxa"]):
            issues.append(f"表格 {table_index} 左缩进不是 120 DXA")
        if layout is None or layout.get(qn("w:type")) != "fixed":
            issues.append(f"表格 {table_index} 未使用 fixed 布局")
        grid_widths = [int(column.get(qn("w:w"))) for column in table._tbl.tblGrid]
        if not grid_widths or sum(grid_widths) != DOCX_TOKENS["content_width_dxa"]:
            issues.append(f"表格 {table_index} tblGrid 总宽度错误")
            continue
        for row_index, row in enumerate(table.rows, start=1):
            if row._tr.get_or_add_trPr().find(qn("w:trHeight")) is not None:
                issues.append(f"表格 {table_index} 第 {row_index} 行使用固定行高")
            if len(row.cells) != len(grid_widths):
                issues.append(f"表格 {table_index} 第 {row_index} 行列数与 tblGrid 不一致")
                continue
            for cell_index, cell in enumerate(row.cells):
                properties = cell._tc.get_or_add_tcPr()
                width = properties.find(qn("w:tcW"))
                if width is None or width.get(qn("w:type")) != "dxa" or int(width.get(qn("w:w"))) != grid_widths[cell_index]:
                    issues.append(f"表格 {table_index} 第 {row_index} 行第 {cell_index + 1} 列宽度错误")
                margins = properties.find(qn("w:tcMar"))
                if margins is None:
                    issues.append(f"表格 {table_index} 第 {row_index} 行第 {cell_index + 1} 列缺少单元格边距")
                    continue
                for side, expected in DOCX_TOKENS["cell_margins_dxa"].items():
                    margin = margins.find(qn(f"w:{side}"))
                    if margin is None or margin.get(qn("w:w")) != str(expected):
                        issues.append(f"表格 {table_index} 第 {row_index} 行第 {cell_index + 1} 列 {side} 边距错误")
    return issues


def _direct_run_size_issues(document: Document) -> list[str]:
    """Audit direct OOXML sizes after the cover; inheritance remains governed by Normal."""
    issues: list[str] = []
    expected = str(int(DOCX_TOKENS["body_size_pt"] * 2))
    body_started = False
    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        if paragraph.style.name == "Heading 1":
            body_started = True
        if not body_started or paragraph.style.name.startswith("Heading"):
            continue
        for run in paragraph._p.xpath(".//w:r[w:t]"):
            properties = run.find(qn("w:rPr"))
            if properties is None:
                continue
            for tag in ("w:sz", "w:szCs"):
                size = properties.find(qn(tag))
                if size is not None and size.get(qn("w:val")) != expected:
                    issues.append(
                        f"正文直接字号不是 11 pt: paragraph={paragraph_index} value={size.get(qn('w:val'))}"
                    )
                    break
    for table_index, table in enumerate(document.tables, start=1):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph._p.xpath(".//w:r[w:t]"):
                        properties = run.find(qn("w:rPr"))
                        if properties is None:
                            continue
                        for tag in ("w:sz", "w:szCs"):
                            size = properties.find(qn(tag))
                            if size is not None and size.get(qn("w:val")) != expected:
                                issues.append(
                                    f"表格直接字号不是 11 pt: table={table_index} value={size.get(qn('w:val'))}"
                                )
                                break
    return issues


def _external_hyperlinks(document: Document) -> set[str]:
    return {
        relationship.target_ref
        for relationship in document.part.rels.values()
        if relationship.reltype.endswith("/hyperlink") and relationship.is_external
    }


def _audit_content(
    document: Document,
    scope: str,
    expected_records: list[dict],
    *,
    expected_taxonomy: list[dict] | None = None,
    expected_item: dict | None = None,
) -> list[str]:
    issues: list[str] = []
    text = _document_text(document)
    if FORBIDDEN_PLACEHOLDERS.search(text):
        issues.append("存在 TODO/TBD/PLACEHOLDER 占位文本")
    if any(reference in text for reference in FORBIDDEN_INTERNAL_REFERENCES):
        issues.append("存在内部工具或内部路径引用")
    if "已经运行成功" in text or re.search(r"(?<!未)运行成功", text):
        issues.append("存在对实际运行效果的夸大表述")
    if "未安装、未运行" not in text:
        issues.append("封面缺少未安装、未运行核验边界")

    expected_ids = [record["id"] for record in sorted(expected_records, key=lambda item: item["id"])]
    found_ids = SKILL_ID_PATTERN.findall(text)
    if scope == "subcategory":
        if expected_item:
            expected_code = expected_item.get("subcategory_code")
            expected_name = expected_item.get("subcategory_name")
            if expected_name:
                header_texts = {
                    "".join(paragraph.text for paragraph in section.header.paragraphs).strip()
                    for section in document.sections
                }
                if header_texts != {expected_name} or f"{expected_code} {expected_name}" not in text:
                    issues.append("小分类代码/名称与 manifest 路径归属不一致")
        if found_ids != expected_ids:
            issues.append(f"Skill ID 与派生数据不一致: expected={expected_ids} actual={found_ids}")
        for record in expected_records:
            for field in PLAIN_FIELDS:
                if record[field] not in text:
                    issues.append(f"{record['id']} 缺少通俗字段 {field}")
        links = _external_hyperlinks(document)
        for record in expected_records:
            for field in ("skill_url", "repo_url"):
                if record[field] not in links:
                    issues.append(f"{record['id']} 缺少真实超链接 {field}")
        expected_h2 = [record["cn"] for record in sorted(expected_records, key=lambda item: item["id"])]
        actual_h2 = [paragraph.text for paragraph in document.paragraphs if paragraph.style.name == "Heading 2"]
        if actual_h2 != expected_h2:
            issues.append("Skill Heading 2 顺序与派生数据不一致")
        trace_headings = [paragraph.text for paragraph in document.paragraphs if paragraph.style.name == "Heading 3"]
        if trace_headings != ["技术追溯"] * len(expected_records):
            issues.append("技术追溯分区数量错误")
    elif scope == "overview":
        required_headings = [
            "这类工具解决什么问题",
            "适合哪些人",
            "小分类导航与数量",
            "怎样选择",
            "共同使用条件",
            "共同限制",
            "本次核验到哪一步",
            "来源说明",
        ]
        actual = [paragraph.text for paragraph in document.paragraphs if paragraph.style.name == "Heading 1"]
        if actual != required_headings:
            issues.append("大分类总览章节顺序错误")
        for field in ("plain_purpose", "plain_audience", "plain_when_to_use", "plain_prerequisites", "plain_limitations", "plain_verification"):
            if expected_records and not any(record[field] in text for record in expected_records):
                issues.append(f"大分类总览未使用通俗字段 {field}")
        if not document.tables or [cell.text for cell in document.tables[0].rows[0].cells] != ["代码与名称", "通俗定义", "数量"]:
            issues.append("大分类总览缺少三列小分类导航")
        if expected_taxonomy is not None:
            categories = sorted(expected_taxonomy, key=lambda item: item["code"])
            counts = Counter(record["subcategory_code"] for record in expected_records)
            expected_rows = [
                (f"{category['code']} {category['name']}", category["inclusion_focus"], str(counts[category["code"]]))
                for category in categories
            ]
            actual_rows: list[tuple[str, str, str]] = []
            if document.tables:
                for row in document.tables[0].rows[1:]:
                    first = row.cells[0].text.splitlines()[0].strip()
                    actual_rows.append((first, row.cells[1].text.strip(), row.cells[2].text.strip()))
            if actual_rows != expected_rows:
                issues.append(f"概览行与 taxonomy/成员数量不一致: expected={expected_rows} actual={actual_rows}")
            links = _external_hyperlinks(document)
            expected_links = {
                *{record["repo_url"] for record in expected_records},
                *{_knowledge_base_target(category) for category in categories},
            }
            if links != expected_links:
                issues.append(f"概览链接与仓库/知识库不一致: expected={sorted(expected_links)} actual={sorted(links)}")
            text_count = len(expected_records)
            if f"收录 {text_count} 项 Skill" not in text or f"本总览汇总 {text_count} 项 Skill" not in text:
                issues.append(f"概览总成员数不准确: expected={text_count}")
            if expected_item and any(category["code"][:2] != expected_item.get("big_category_code") for category in categories):
                issues.append("概览 taxonomy 含有跨大分类行")
    else:
        issues.append(f"未知验证 scope: {scope}")
    return issues


def verify_document(
    path: Path,
    *,
    scope: str,
    expected_records: list[dict],
    expected_taxonomy: list[dict] | None = None,
    expected_item: dict | None = None,
) -> list[str]:
    """Return all structural/content issues for one DOCX; an empty list means pass."""
    if not path.exists() or path.stat().st_size == 0:
        return [f"DOCX 不存在或为空: {path}"]
    try:
        document = Document(path)
    except Exception as exc:  # pragma: no cover - defensive reopen boundary
        return [f"DOCX 无法重新打开: {exc}"]
    issues = []
    issues.extend(_audit_sections(document))
    issues.extend(_audit_styles(document))
    issues.extend(_audit_tables(document))
    issues.extend(_direct_run_size_issues(document))
    issues.extend(
        _audit_content(
            document,
            scope,
            expected_records,
            expected_taxonomy=expected_taxonomy,
            expected_item=expected_item,
        )
    )
    return issues


def verify_selected_documents(
    records: list[dict],
    manifest: list[dict],
    project_root: Path,
    *,
    only: list[str] | None = None,
    taxonomy: list[dict] | None = None,
    assignments: dict[str, str] | None = None,
    repositories: dict | None = None,
) -> dict[str, list[str]]:
    if taxonomy is not None:
        validate_manifest_contract(manifest, taxonomy)
    if taxonomy is not None and assignments is not None and repositories is not None:
        validate_source_contract(records, taxonomy, assignments, repositories, project_root)
    taxonomy_by_big: dict[str, list[dict]] = defaultdict(list)
    for category in taxonomy or []:
        taxonomy_by_big[category["code"][:2]].append(category)
    grouped: dict[str, list[dict]] = defaultdict(list)
    for record in sorted(records, key=lambda item: item["id"]):
        grouped[record["subcategory_code"]].append(record)
    results: dict[str, list[str]] = {}
    for item in select_manifest_items(manifest, only):
        if item["scope"] == "overview":
            expected = [record for record in records if record["subcategory_code"].startswith(f"{item['big_category_code']}-")]
        else:
            expected = grouped[item["subcategory_code"]]
        path = project_root / Path(*Path(item["path"]).parts)
        results[item["key"]] = verify_document(
            path,
            scope=item["scope"],
            expected_records=expected,
            expected_taxonomy=taxonomy_by_big[item["big_category_code"]] if item["scope"] == "overview" and taxonomy is not None else None,
            expected_item=item,
        )
    return results


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--only", nargs="*", help="仅验证键，如 05-overview 05-05")
    args = parser.parse_args(argv)
    records = _load_json(PLAIN_CATALOG_FILE)
    assignment = _load_json(ASSIGNMENT_FILE)
    manifest = _load_json(MANIFEST_FILE)
    repositories = _load_json(REPOSITORIES_FILE)
    if not isinstance(records, list) or not isinstance(assignment, dict) or not isinstance(manifest, list) or not isinstance(repositories, dict):
        raise ValueError("验证输入格式错误")
    taxonomy = assignment.get("taxonomy")
    assignments = assignment.get("assignments")
    if not isinstance(taxonomy, list) or not isinstance(assignments, dict):
        raise ValueError("验证归属输入格式错误")
    only = _parse_only(args.only)
    results = verify_selected_documents(
        records,
        manifest,
        PROJECT_ROOT,
        only=only,
        taxonomy=taxonomy,
        assignments=assignments,
        repositories=repositories,
    )
    failed = {key: issues for key, issues in results.items() if issues}
    if failed:
        for key, issues in failed.items():
            print(f"FAIL {key}")
            for issue in issues:
                print(f"  - {issue}")
        return 1
    counts = Counter(item["scope"] for item in select_manifest_items(manifest, only))
    print(f"verified={len(results)} overview={counts['overview']} subcategory={counts['subcategory']} preset=OK content=OK hyperlinks=OK")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
