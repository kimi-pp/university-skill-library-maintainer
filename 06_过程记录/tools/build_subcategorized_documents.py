"""Build manifest-driven plain-language overview and subcategory DOCX reports."""

from __future__ import annotations

import argparse
import json
import math
import re
import tempfile
import unicodedata
import zipfile
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path, PurePosixPath
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
PLAIN_CATALOG_FILE = PROJECT_ROOT / "03_候选池" / "derived" / "plain_language_catalog.json"
ASSIGNMENT_FILE = PROJECT_ROOT / "03_候选池" / "derived" / "subcategory_assignments.json"
MANIFEST_FILE = PROJECT_ROOT / "03_候选池" / "derived" / "subcategory_manifest.json"
REPOSITORIES_FILE = PROJECT_ROOT / "03_候选池" / "deduplicated" / "repositories.json"
DATA_DATE = "2026-08-07"

BIG_CATEGORY_NAMES = {
    "01": "学术写作、引用与出版",
    "02": "文档、表格、演示文稿与办公自动化",
    "03": "文献检索与学术研究",
    "04": "图书馆与信息素养",
    "05": "编程、数学、数据分析和可视化",
}
BIG_CATEGORY_DIRECTORIES = {
    "01": "01_学术写作引用与出版",
    "02": "02_文档表格演示文稿与办公自动化",
    "03": "03_文献检索与学术研究",
    "04": "04_图书馆与信息素养",
    "05": "05_编程数学数据分析和可视化",
}
DELIVERY_ROOT = PurePosixPath("05_交付物/通俗细分版_2026-08-07")

DOCX_TOKENS = {
    "preset_name": "compact_reference_guide",
    "header_template": "editorial_cover",
    "page_width_twips": 12240,
    "page_height_twips": 15840,
    "margin_twips": 1440,
    "header_footer_distance_twips": 708,
    "content_width_dxa": 9360,
    "table_indent_dxa": 120,
    "cell_margins_dxa": {"top": 80, "bottom": 80, "start": 120, "end": 120},
    "body_font_western": "Calibri",
    "body_font_east_asia": "Microsoft YaHei",
    "body_size_pt": 11,
    "body_after_pt": 6,
    "body_line_spacing": 1.25,
    "heading_1": (16, "2E74B5", 18, 10),
    "heading_2": (13, "2E74B5", 14, 7),
    "heading_3": (12, "1F4D78", 10, 5),
    "table_header_fill": "E8EEF5",
    "table_grid": "B8C4D1",
    "ink": "243447",
    "muted": "5B6775",
    "link": "1F5F8B",
    "cover_title": "203748",
    "cover_accent": "8A6512",
}

PLAIN_FIELDS = (
    "plain_purpose",
    "plain_outputs",
    "plain_audience",
    "plain_when_to_use",
    "plain_prerequisites",
    "plain_limitations",
    "plain_integration",
    "plain_verification",
)


def _reject_duplicate_json_keys(pairs: list[tuple[str, object]]) -> dict:
    result: dict = {}
    for key, value in pairs:
        if key in result:
            raise ValueError(f"重复 JSON 键: {key}")
        result[key] = value
    return result


def _load_json(path: Path) -> object:
    return json.loads(
        path.read_text(encoding="utf-8"),
        object_pairs_hook=_reject_duplicate_json_keys,
    )


def _set_run_font(
    run,
    *,
    size: float = 11,
    bold: bool = False,
    italic: bool = False,
    color: str = DOCX_TOKENS["ink"],
) -> None:
    western = DOCX_TOKENS["body_font_western"]
    east_asia = DOCX_TOKENS["body_font_east_asia"]
    run.font.name = western
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), western)
    fonts.set(qn("w:hAnsi"), western)
    fonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def _set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 6, line: float = 1.25) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def _cover_title_font_size(title: str) -> float:
    """Choose a cover-title size from estimated rendered width, not raw character count."""
    width_units = 0.0
    for character in title:
        if character.isspace():
            width_units += 0.35
        elif unicodedata.east_asian_width(character) in {"W", "F"}:
            width_units += 1.0
        else:
            width_units += 0.55
    if width_units <= 0:
        return 27
    fitted = min(27.0, 445.0 / width_units)
    return max(22.0, math.floor(fitted * 2) / 2)


def _shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_text(cell, text: object, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    _set_paragraph_spacing(paragraph, after=2, line=1.15)
    run = paragraph.add_run(str(text))
    _set_run_font(run, size=DOCX_TOKENS["body_size_pt"], bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_table_geometry(table, widths_dxa: list[int]) -> None:
    """Apply the compact_reference_guide fixed-DXA table contract."""
    if sum(widths_dxa) != DOCX_TOKENS["content_width_dxa"]:
        raise ValueError(f"表格列宽总和必须为 {DOCX_TOKENS['content_width_dxa']} DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    properties = table._tbl.tblPr

    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(DOCX_TOKENS["content_width_dxa"]))

    table_indent = properties.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        properties.append(table_indent)
    table_indent.set(qn("w:type"), "dxa")
    table_indent.set(qn("w:w"), str(DOCX_TOKENS["table_indent_dxa"]))

    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), DOCX_TOKENS["table_grid"])

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)

    for row in table.rows:
        row_properties = row._tr.get_or_add_trPr()
        if row_properties.find(qn("w:cantSplit")) is None:
            row_properties.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(widths_dxa[index]))
            margins = cell_properties.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell_properties.append(margins)
            for side, value in DOCX_TOKENS["cell_margins_dxa"].items():
                margin = margins.find(qn(f"w:{side}"))
                if margin is None:
                    margin = OxmlElement(f"w:{side}")
                    margins.append(margin)
                margin.set(qn("w:type"), "dxa")
                margin.set(qn("w:w"), str(value))


def _repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _add_hyperlink(
    paragraph,
    label: str,
    url: str,
    *,
    size: float = DOCX_TOKENS["body_size_pt"],
    allow_relative: bool = False,
) -> None:
    is_https = isinstance(url, str) and bool(re.match(r"^https://", url))
    is_relative = (
        allow_relative
        and isinstance(url, str)
        and url.startswith("../../../02_知识库/")
        and "\\" not in url
    )
    if not (is_https or is_relative):
        raise ValueError(f"超链接必须使用 HTTPS 或批准的知识库相对路径: {url!r}")
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), DOCX_TOKENS["body_font_western"])
    fonts.set(qn("w:hAnsi"), DOCX_TOKENS["body_font_western"])
    fonts.set(qn("w:eastAsia"), DOCX_TOKENS["body_font_east_asia"])
    color = OxmlElement("w:color")
    color.set(qn("w:val"), DOCX_TOKENS["link"])
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size_element = OxmlElement("w:sz")
    size_element.set(qn("w:val"), str(int(size * 2)))
    complex_size_element = OxmlElement("w:szCs")
    complex_size_element.set(qn("w:val"), str(int(size * 2)))
    properties.extend((fonts, color, underline, size_element, complex_size_element))
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_spacing(paragraph, after=0, line=1.0)
    lead = paragraph.add_run("第 ")
    _set_run_font(lead, size=9, color=DOCX_TOKENS["muted"])
    field_run = OxmlElement("w:r")
    field_properties = OxmlElement("w:rPr")
    field_fonts = OxmlElement("w:rFonts")
    field_fonts.set(qn("w:ascii"), DOCX_TOKENS["body_font_western"])
    field_fonts.set(qn("w:hAnsi"), DOCX_TOKENS["body_font_western"])
    field_fonts.set(qn("w:eastAsia"), DOCX_TOKENS["body_font_east_asia"])
    field_properties.append(field_fonts)
    field_run.append(field_properties)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run.extend((begin, instruction, separate, value, end))
    paragraph._p.append(field_run)
    tail = paragraph.add_run(" 页")
    _set_run_font(tail, size=9, color=DOCX_TOKENS["muted"])


def _configure_document(document: Document, running_label: str) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = DOCX_TOKENS["body_font_western"]
    normal_fonts = normal._element.get_or_add_rPr().get_or_add_rFonts()
    normal_fonts.set(qn("w:ascii"), DOCX_TOKENS["body_font_western"])
    normal_fonts.set(qn("w:hAnsi"), DOCX_TOKENS["body_font_western"])
    normal_fonts.set(qn("w:eastAsia"), DOCX_TOKENS["body_font_east_asia"])
    normal.font.size = Pt(DOCX_TOKENS["body_size_pt"])
    normal.font.color.rgb = RGBColor.from_string(DOCX_TOKENS["ink"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(DOCX_TOKENS["body_after_pt"])
    normal.paragraph_format.line_spacing = DOCX_TOKENS["body_line_spacing"]

    for style_name, token_name in (
        ("Heading 1", "heading_1"),
        ("Heading 2", "heading_2"),
        ("Heading 3", "heading_3"),
    ):
        size, color, before, after = DOCX_TOKENS[token_name]
        style = document.styles[style_name]
        style.font.name = DOCX_TOKENS["body_font_western"]
        fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        fonts.set(qn("w:ascii"), DOCX_TOKENS["body_font_western"])
        fonts.set(qn("w:hAnsi"), DOCX_TOKENS["body_font_western"])
        fonts.set(qn("w:eastAsia"), DOCX_TOKENS["body_font_east_asia"])
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(header, after=0, line=1.0)
    _set_run_font(header.add_run(running_label), size=9, bold=True, color=DOCX_TOKENS["muted"])
    footer = section.footer.paragraphs[0]
    footer.text = ""
    _add_page_number(footer)

    fixed_time = datetime(2026, 8, 7, 0, 0, 0, tzinfo=timezone.utc)
    document.core_properties.created = fixed_time
    document.core_properties.modified = fixed_time
    document.core_properties.author = "高校 AI 技能库调研"
    document.core_properties.last_modified_by = "高校 AI 技能库调研"


def _add_cover(document: Document, *, title: str, subtitle: str, count: int) -> None:
    spacer = document.add_paragraph()
    _set_paragraph_spacing(spacer, before=0, after=74, line=1.0)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(kicker, after=16, line=1.0)
    _set_run_font(kicker.add_run("高校 AI 技能库调研"), size=10.5, bold=True, color=DOCX_TOKENS["cover_accent"])

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(title_paragraph, after=10, line=1.05)
    title_paragraph.paragraph_format.keep_with_next = True
    _set_run_font(
        title_paragraph.add_run(title),
        size=_cover_title_font_size(title),
        bold=True,
        color=DOCX_TOKENS["cover_title"],
    )

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(subtitle_paragraph, after=32, line=1.15)
    subtitle_paragraph.paragraph_format.keep_with_next = True
    _set_run_font(subtitle_paragraph.add_run(subtitle), size=14, color="2B5163")

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(metadata, after=8, line=1.15)
    _set_run_font(metadata.add_run(f"数据日期：{DATA_DATE}  |  收录 {count} 项 Skill"), size=10.5, bold=True, color=DOCX_TOKENS["cover_title"])

    boundary = document.add_paragraph()
    boundary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(boundary, after=60, line=1.2)
    _set_run_font(
        boundary.add_run("核验边界：已核对公开说明与包内容；本次未安装、未运行，不能据此判断实际效果。"),
        size=9.5,
        italic=True,
        color=DOCX_TOKENS["muted"],
    )
    document.add_page_break()


def _add_labeled_paragraph(document: Document, label: str, value: str) -> None:
    if not isinstance(value, str) or not value.strip():
        raise ValueError(f"{label}不能为空")
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_together = False
    label_run = paragraph.add_run(f"{label}：")
    _set_run_font(label_run, bold=True, color=DOCX_TOKENS["cover_title"])
    _set_run_font(paragraph.add_run(value.strip()))


def _unique_plain_values(records: Iterable[dict], field: str, *, limit: int = 6) -> list[str]:
    values: list[str] = []
    for record in sorted(records, key=lambda item: item["id"]):
        value = record.get(field)
        if not isinstance(value, str) or not value.strip():
            raise ValueError(f"{record.get('id', '未知 Skill')} 缺少 {field}")
        if value not in values:
            values.append(value)
        if len(values) == limit:
            break
    return values


def _add_plain_examples(
    document: Document,
    records: list[dict],
    field: str,
    label: str,
    *,
    limit: int = 6,
) -> None:
    for index, value in enumerate(_unique_plain_values(records, field, limit=limit), start=1):
        _add_labeled_paragraph(document, f"{label}{index}", value)


def _validate_record(record: dict) -> None:
    required = {
        "id", "name", "cn", "cat", "repo", "ecosystem", "form", "tags", "compat",
        "verify", "priority", "repo_url", "skill_url", "license", "repo_pushed",
        "subcategory_code", "subcategory_name", *PLAIN_FIELDS,
    }
    missing = sorted(key for key in required if not isinstance(record.get(key), str) or not record[key].strip())
    if missing:
        raise ValueError(f"{record.get('id', '未知 Skill')} 缺少字段: {missing}")
    if not re.fullmatch(r"GH-\d{2}-\d{4}", record["id"]):
        raise ValueError(f"Skill ID 格式错误: {record['id']}")


def _sorted_records(records: list[dict]) -> list[dict]:
    copied = [dict(record) for record in records]
    for record in copied:
        _validate_record(record)
    ids = [record["id"] for record in copied]
    if len(ids) != len(set(ids)):
        raise ValueError("重复 Skill ID")
    return sorted(copied, key=lambda record: record["id"])


def _integration_difficulty(record: dict) -> str:
    compatibility = record.get("compat", "").upper()
    if compatibility == "A":
        return "低"
    if compatibility == "B":
        return "中"
    return "高"


def _add_short_value_table(document: Document, record: dict) -> None:
    table = document.add_table(rows=4, cols=2)
    values = (
        ("英文名称", record["name"]),
        ("推荐程度", record["priority"]),
        ("接入难度", _integration_difficulty(record)),
        ("核验层级", record["verify"]),
    )
    for row, (label, value) in zip(table.rows, values):
        _set_cell_text(row.cells[0], label, bold=True)
        _shade_cell(row.cells[0], DOCX_TOKENS["table_header_fill"])
        _set_cell_text(row.cells[1], value)
    for row in table.rows[:-1]:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.keep_with_next = True
    _set_table_geometry(table, [2700, 6660])


def _add_source_links(document: Document, record: dict) -> None:
    paragraph = document.add_paragraph()
    label = paragraph.add_run("原始资料地址：")
    _set_run_font(label, bold=True, color=DOCX_TOKENS["cover_title"])
    _add_hyperlink(paragraph, "Skill 说明页", record["skill_url"])
    separator = paragraph.add_run("  |  ")
    _set_run_font(separator, color=DOCX_TOKENS["muted"])
    _add_hyperlink(paragraph, "GitHub 仓库", record["repo_url"])


def _add_compact_trace(document: Document, record: dict) -> None:
    """Keep source facts distinct from user guidance without wasting a page."""
    if not document.paragraphs:
        raise ValueError("技术追溯前缺少读者正文")
    document.paragraphs[-1].paragraph_format.keep_with_next = True
    paragraph = document.add_paragraph(style="Heading 3")
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = False
    _set_paragraph_spacing(paragraph, before=6, after=0, line=1.0)
    _set_run_font(
        paragraph.add_run("技术追溯｜"),
        size=DOCX_TOKENS["body_size_pt"],
        bold=True,
        color=DOCX_TOKENS["heading_3"][1],
    )
    fields = (
        ("内部编号", record["id"]),
        ("功能标签", record["tags"]),
        ("原生生态", record["ecosystem"]),
        ("来源形态", record["form"]),
        ("许可证", record["license"]),
        ("仓库最近更新", record["repo_pushed"]),
    )
    for index, (label, value) in enumerate(fields):
        if index:
            _set_run_font(paragraph.add_run("； "), color=DOCX_TOKENS["muted"])
        _set_run_font(paragraph.add_run(f"{label}："), bold=True, color=DOCX_TOKENS["cover_title"])
        _set_run_font(paragraph.add_run(value))
    _set_run_font(paragraph.add_run("； 原始资料地址："), bold=True, color=DOCX_TOKENS["cover_title"])
    _add_hyperlink(paragraph, "Skill 说明页", record["skill_url"])
    _set_run_font(paragraph.add_run(" | "), color=DOCX_TOKENS["muted"])
    _add_hyperlink(paragraph, "GitHub 仓库", record["repo_url"])


def _knowledge_base_target(category: dict) -> str:
    code = category["code"]
    name = category["name"]
    return (
        f"../../../02_知识库/functional_domains/{BIG_CATEGORY_DIRECTORIES[code[:2]]}/"
        f"subcategories/{code}_{name}/INDEX.md"
    )


def _subcategory_repository_links(code: str, records: list[dict]) -> list[str]:
    repositories: dict[str, str] = {}
    for record in records:
        if record["subcategory_code"] == code:
            repositories.setdefault(record["repo"], record["repo_url"])
    return [repositories[name] for name in sorted(repositories)]


def build_overview_document(big_category: dict, records: list[dict]) -> Document:
    """Return one big-category overview built from approved plain-language fields."""
    code = big_category.get("code")
    name = big_category.get("name")
    categories = sorted(big_category.get("subcategories", []), key=lambda item: item["code"])
    if code not in BIG_CATEGORY_NAMES or not isinstance(name, str) or not name.strip():
        raise ValueError("大分类信息不完整")
    if not categories:
        raise ValueError(f"{code} 缺少小分类")
    sorted_records = _sorted_records(records)
    if not sorted_records:
        raise ValueError(f"空大分类: {code}")
    if any(record["subcategory_code"][:2] != code for record in sorted_records):
        raise ValueError(f"{code} 包含其他大分类的 Skill")

    document = Document()
    _configure_document(document, name)
    _add_cover(document, title=f"{code} {name}", subtitle="大分类总览 · 通俗版选择指南", count=len(sorted_records))

    document.add_heading("这类工具解决什么问题", level=1)
    _add_plain_examples(document, sorted_records, "plain_purpose", "代表用途")

    document.add_heading("适合哪些人", level=1)
    _add_plain_examples(document, sorted_records, "plain_audience", "适用对象")

    document.add_heading("小分类导航与数量", level=1)
    table = document.add_table(rows=1, cols=3)
    for cell, label in zip(table.rows[0].cells, ("代码与名称", "通俗定义", "数量")):
        _set_cell_text(cell, label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(cell, DOCX_TOKENS["table_header_fill"])
    counts = Counter(record["subcategory_code"] for record in sorted_records)
    for category in categories:
        row = table.add_row()
        _set_cell_text(row.cells[0], f"{category['code']} {category['name']}")
        first_paragraph = row.cells[0].paragraphs[0]
        _set_run_font(first_paragraph.add_run("\n"), color=DOCX_TOKENS["muted"])
        _add_hyperlink(
            first_paragraph,
            "知识库入口",
            _knowledge_base_target(category),
            allow_relative=True,
        )
        for index, repository_url in enumerate(
            _subcategory_repository_links(category["code"], sorted_records), start=1
        ):
            _set_run_font(first_paragraph.add_run(" | "), color=DOCX_TOKENS["muted"])
            _add_hyperlink(first_paragraph, f"仓库{index}", repository_url)
        _set_cell_text(row.cells[1], category["inclusion_focus"])
        _set_cell_text(row.cells[2], counts[category["code"]], align=WD_ALIGN_PARAGRAPH.CENTER)
    _repeat_table_header(table.rows[0])
    _set_table_geometry(table, [2850, 5610, 900])

    document.add_heading("怎样选择", level=1)
    _add_plain_examples(document, sorted_records, "plain_when_to_use", "选择线索")

    document.add_heading("共同使用条件", level=1)
    _add_plain_examples(document, sorted_records, "plain_prerequisites", "准备事项")

    document.add_heading("共同限制", level=1)
    _add_plain_examples(document, sorted_records, "plain_limitations", "注意事项", limit=4)

    document.add_heading("本次核验到哪一步", level=1)
    _add_plain_examples(document, sorted_records, "plain_verification", "核验说明")

    document.add_heading("来源说明", level=1)
    repositories: dict[str, dict] = {}
    for record in sorted_records:
        repositories.setdefault(record["repo"], record)
    _add_labeled_paragraph(
        document,
        "范围",
        f"本总览汇总 {len(sorted_records)} 项 Skill，来自 {len(repositories)} 个 GitHub 仓库；成员归属以已批准的小分类台账为准。",
    )
    _add_labeled_paragraph(
        document,
        "链接位置",
        "每个小分类的知识库入口和相关 GitHub 仓库已随导航表逐行列出。",
    )
    return document


def build_subcategory_document(subcategory: dict, records: list[dict]) -> Document:
    """Return one independent reader-first report for a non-empty subcategory."""
    code = subcategory.get("code")
    name = subcategory.get("name")
    if not isinstance(code, str) or not re.fullmatch(r"\d{2}-\d{2}", code):
        raise ValueError("小分类代码格式错误")
    if not isinstance(name, str) or not name.strip():
        raise ValueError("小分类名称不能为空")
    sorted_records = _sorted_records(records)
    if not sorted_records:
        raise ValueError(f"空小分类: {code}")
    if any(record["subcategory_code"] != code for record in sorted_records):
        raise ValueError(f"{code} 包含其他小分类的 Skill")

    document = Document()
    _configure_document(document, name)
    _add_cover(document, title=f"{code} {name}", subtitle="小分类独立报告 · 通俗版", count=len(sorted_records))

    document.add_heading("阅读说明", level=1)
    _add_labeled_paragraph(document, "这类工具是做什么的", subcategory.get("inclusion_focus", "").strip())
    _add_labeled_paragraph(document, "收录方式", "以下条目按内部编号排序；先看用途和使用条件，再看独立列出的技术追溯信息。")

    for record in sorted_records:
        document.add_heading(record["cn"], level=2)
        _add_short_value_table(document, record)
        _add_labeled_paragraph(document, "它能帮您做什么", record["plain_purpose"])
        _add_labeled_paragraph(document, "可得到什么", record["plain_outputs"])
        _add_labeled_paragraph(document, "适合谁使用", record["plain_audience"])
        _add_labeled_paragraph(document, "什么情况下值得使用", record["plain_when_to_use"])
        _add_labeled_paragraph(document, "使用前需要准备什么", record["plain_prerequisites"])
        _add_labeled_paragraph(document, "可能遇到什么限制", record["plain_limitations"])
        _add_labeled_paragraph(document, "接入需要多少调整", record["plain_integration"])
        _add_labeled_paragraph(document, "本次核验到了哪一步", record["plain_verification"])
        _add_compact_trace(document, record)
    return document


def _manifest_key(item: dict) -> str:
    if item.get("scope") == "overview":
        return f"{item.get('big_category_code')}-overview"
    if item.get("scope") == "subcategory":
        return str(item.get("subcategory_code"))
    raise ValueError(f"未知 manifest scope: {item.get('scope')!r}")


def _safe_relative_docx_path(value: object) -> Path:
    if not isinstance(value, str) or not value:
        raise ValueError("不安全的 manifest 路径: 空路径")
    pure = PurePosixPath(value)
    if pure.is_absolute() or ".." in pure.parts or pure.suffix.lower() != ".docx":
        raise ValueError(f"不安全的 manifest 路径: {value!r}")
    if not pure.parts or pure.parts[0] != "05_交付物":
        raise ValueError(f"不安全的 manifest 路径: {value!r}")
    if any(re.search(r'[<>:"|?*\\]', part) for part in pure.parts):
        raise ValueError(f"不安全的 manifest 路径: {value!r}")
    return Path(*pure.parts)


def _taxonomy_by_code(taxonomy: list[dict]) -> dict[str, dict]:
    result: dict[str, dict] = {}
    for original in taxonomy:
        item = dict(original)
        code = item.get("code")
        name = item.get("name")
        focus = item.get("inclusion_focus")
        if not isinstance(code, str) or not re.fullmatch(r"\d{2}-\d{2}", code):
            raise ValueError(f"小分类代码格式错误: {code!r}")
        if code[:2] not in BIG_CATEGORY_DIRECTORIES:
            raise ValueError(f"未知大分类: {code}")
        if not isinstance(name, str) or not name.strip():
            raise ValueError(f"小分类名称不能为空: {code}")
        if not isinstance(focus, str) or not focus.strip():
            raise ValueError(f"小分类收录重点不能为空: {code}")
        if code in result:
            raise ValueError(f"小分类代码重复: {code}")
        result[code] = item
    return result


def _expected_manifest_path(item: dict, taxonomy_by_code: dict[str, dict]) -> str:
    big_code = item["big_category_code"]
    directory = PurePosixPath(BIG_CATEGORY_DIRECTORIES[big_code])
    suffix = item["format"]
    if item["scope"] == "overview":
        return str(DELIVERY_ROOT / directory / f"00_大分类总览.{suffix}")
    code = item["subcategory_code"]
    name = taxonomy_by_code[code]["name"]
    stem = f"{code}_{name}_GitHub技能调研"
    return str(DELIVERY_ROOT / directory / f"{code}_{name}" / f"{stem}.{suffix}")


def validate_manifest_contract(
    manifest: list[dict], taxonomy: list[dict], *, require_complete: bool = True
) -> None:
    """Reject metadata/path drift and require one DOCX/XLSX pair per logical output."""
    taxonomy_map = _taxonomy_by_code(taxonomy)
    pairs: dict[tuple[str, str], list[dict]] = defaultdict(list)
    paths: set[str] = set()
    for original in manifest:
        item = dict(original)
        scope = item.get("scope")
        file_format = item.get("format")
        big_code = item.get("big_category_code")
        if scope not in {"overview", "subcategory"} or file_format not in {"docx", "xlsx"}:
            raise ValueError(f"manifest scope/format 错误: {scope!r}/{file_format!r}")
        if big_code not in BIG_CATEGORY_DIRECTORIES:
            raise ValueError(f"manifest 大分类错误: {big_code!r}")
        path = item.get("path")
        if not isinstance(path, str):
            raise ValueError("manifest 路径缺失")
        if path in paths:
            raise ValueError(f"重复 manifest 路径: {path}")
        paths.add(path)
        if scope == "subcategory":
            code = item.get("subcategory_code")
            if not isinstance(code, str) or code[:2] != big_code:
                raise ValueError(f"小分类与大分类代码归属不一致: {code!r}/{big_code!r}")
            category = taxonomy_map.get(code)
            if category is None:
                raise ValueError(f"manifest 使用未知小分类: {code}")
            if item.get("subcategory_name") != category["name"]:
                raise ValueError(f"manifest 小分类名称不一致: {code}")
            logical_key = (scope, code)
        else:
            logical_key = (scope, big_code)
        expected_path = _expected_manifest_path(item, taxonomy_map)
        if path != expected_path:
            raise ValueError(f"manifest 路径与元数据不一致: expected={expected_path} actual={path}")
        pairs[logical_key].append(item)
    for logical_key, items in pairs.items():
        formats = {item["format"] for item in items}
        if len(items) != 2 or formats != {"docx", "xlsx"}:
            raise ValueError(f"DOCX/XLSX 配对不完整: {logical_key}")
        metadata = [
            {key: value for key, value in item.items() if key not in {"path", "format"}}
            for item in items
        ]
        if metadata[0] != metadata[1]:
            raise ValueError(f"DOCX/XLSX scope 或元数据不一致: {logical_key}")
    if require_complete:
        expected_keys = {
            *{("overview", code[:2]) for code in taxonomy_map},
            *{("subcategory", code) for code in taxonomy_map},
        }
        if set(pairs) != expected_keys:
            raise ValueError(
                f"manifest 未完整覆盖 taxonomy: missing={sorted(expected_keys - set(pairs))} "
                f"extra={sorted(set(pairs) - expected_keys)}"
            )


def validate_source_contract(
    records: list[dict],
    taxonomy: list[dict],
    assignments: dict[str, str],
    repositories: dict,
    project_root: Path,
    *,
    expected_total: int | None = None,
) -> None:
    """Cross-check catalog membership and every generated knowledge-base leaf."""
    taxonomy_map = _taxonomy_by_code(taxonomy)
    sorted_records = _sorted_records(records)
    ids = [record["id"] for record in sorted_records]
    if expected_total is not None and len(ids) != expected_total:
        raise ValueError(f"归属台账总成员数错误: expected={expected_total} actual={len(ids)}")
    if set(assignments) != set(ids):
        raise ValueError("通俗目录与归属台账的 Skill ID 不一致")
    assignment_groups: dict[str, list[str]] = defaultdict(list)
    for skill_id, code in assignments.items():
        if code not in taxonomy_map:
            raise ValueError(f"归属台账含未知或跨类小分类: {skill_id}={code}")
        assignment_groups[code].append(skill_id)
    grouped: dict[str, list[dict]] = defaultdict(list)
    for record in sorted_records:
        code = record["subcategory_code"]
        category = taxonomy_map.get(code)
        if category is None:
            raise ValueError(f"未知小分类: {code}")
        if record.get("cat") != code[:2] or assignments.get(record["id"]) != code:
            raise ValueError(f"{record['id']} 的唯一归属不一致")
        if record.get("subcategory_name") != category["name"]:
            raise ValueError(f"{record['id']} 的小分类名称与 taxonomy 不一致")
        if record.get("repo") not in repositories:
            raise ValueError(f"{record['id']} 的仓库不在 repositories.json 中")
        grouped[code].append(record)
    for code, category in taxonomy_map.items():
        members = grouped.get(code, [])
        member_ids = [record["id"] for record in members]
        if sorted(assignment_groups.get(code, [])) != member_ids:
            raise ValueError(
                f"归属台账小分类成员集合不一致: {code} "
                f"expected={member_ids} actual={sorted(assignment_groups.get(code, []))}"
            )
        if not members:
            raise ValueError(f"空小分类: {code}")
        path = (
            project_root
            / "02_知识库"
            / "functional_domains"
            / BIG_CATEGORY_DIRECTORIES[code[:2]]
            / "subcategories"
            / f"{code}_{category['name']}"
            / "INDEX.md"
        )
        if not path.is_file():
            raise ValueError(f"知识库入口不存在: {path}")
        content = path.read_text(encoding="utf-8")
        nonempty_lines = [line.strip() for line in content.splitlines() if line.strip()]
        required_singletons = (
            f"# {code} {category['name']}",
            category["inclusion_focus"],
            f"共 {len(members)} 项 Skill。",
        )
        if any(nonempty_lines.count(value) != 1 for value in required_singletons):
            raise ValueError(f"知识库内容与 taxonomy/成员数量不一致: {code}")
        rows = _parse_knowledge_base_rows(content, code)
        expected_rows = [
            (
                record["id"],
                record["cn"],
                record["plain_purpose"],
                record["priority"],
                f"../../skills/{record['id']}_{record['name']}.md",
            )
            for record in members
        ]
        if rows != expected_rows:
            raise ValueError(
                f"知识库收录表逐行不一致: {code} expected={expected_rows} actual={rows}"
            )
        for row in rows:
            target = (path.parent / Path(*PurePosixPath(row[4]).parts)).resolve()
            expected_target = (
                path.parent / ".." / ".." / "skills" / f"{row[0]}_{next(record['name'] for record in members if record['id'] == row[0])}.md"
            ).resolve()
            if target != expected_target or not target.is_file():
                raise ValueError(f"知识库条目链接不存在或指向错误文件: {code} {row[0]} {row[4]}")


def _markdown_cells(line: str) -> list[str]:
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        raise ValueError(f"知识库 Markdown 表格行格式错误: {line!r}")
    return [cell.strip() for cell in stripped[1:-1].split("|")]


def _parse_knowledge_base_rows(content: str, code: str) -> list[tuple[str, str, str, str, str]]:
    """Parse exactly one five-column leaf ledger and preserve row order/duplicates."""
    lines = content.splitlines()
    expected_header = ["内部编号", "中文名称", "主要用途", "推荐程度", "条目链接"]
    header_indexes = [
        index for index, line in enumerate(lines)
        if line.strip().startswith("|") and _markdown_cells(line) == expected_header
    ]
    if len(header_indexes) != 1:
        raise ValueError(f"知识库收录表表头数量或字段错误: {code}")
    header_index = header_indexes[0]
    if header_index + 1 >= len(lines):
        raise ValueError(f"知识库收录表缺少分隔行: {code}")
    separator = _markdown_cells(lines[header_index + 1])
    if len(separator) != 5 or any(not re.fullmatch(r":?-{3,}:?", cell) for cell in separator):
        raise ValueError(f"知识库收录表分隔行错误: {code}")
    rows: list[tuple[str, str, str, str, str]] = []
    for line in lines[header_index + 2:]:
        if not line.strip().startswith("|"):
            break
        cells = _markdown_cells(line)
        if len(cells) != 5:
            raise ValueError(f"知识库收录表列数错误: {code}")
        link = re.fullmatch(r"\[查看\]\(<([^<>]+)>\)", cells[4])
        if link is None:
            raise ValueError(f"知识库条目链接格式错误: {code} {cells[4]!r}")
        rows.append((cells[0], cells[1], cells[2], cells[3], link.group(1)))
    return rows


def select_manifest_items(manifest: list[dict], only: list[str] | None) -> list[dict]:
    """Validate and select stable DOCX manifest entries by public key."""
    docx_items: list[dict] = []
    seen_keys: set[str] = set()
    for original in manifest:
        if original.get("format") != "docx":
            continue
        item = dict(original)
        _safe_relative_docx_path(item.get("path"))
        key = _manifest_key(item)
        if key in seen_keys:
            raise ValueError(f"重复 DOCX manifest 键: {key}")
        item["key"] = key
        seen_keys.add(key)
        docx_items.append(item)
    if only:
        requested = set(only)
        unknown = sorted(requested - seen_keys)
        if unknown:
            raise ValueError(f"未知 --only 选择项: {unknown}")
        docx_items = [item for item in docx_items if item["key"] in requested]
    return sorted(
        docx_items,
        key=lambda item: (
            item["big_category_code"],
            0 if item["scope"] == "overview" else 1,
            item.get("subcategory_code", ""),
        ),
    )


def _deterministic_save(document: Document, output_path: Path) -> None:
    """Save a DOCX with fixed ZIP timestamps so identical inputs reproduce bytes."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory() as temporary_directory:
        temporary_path = Path(temporary_directory) / "source.docx"
        document.save(temporary_path)
        with zipfile.ZipFile(temporary_path, "r") as source, zipfile.ZipFile(
            output_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
        ) as target:
            for name in sorted(source.namelist()):
                info = zipfile.ZipInfo(name, date_time=(2026, 8, 7, 0, 0, 0))
                info.compress_type = zipfile.ZIP_DEFLATED
                info.external_attr = 0o600 << 16
                target.writestr(info, source.read(name))


def generate_documents(
    records: list[dict],
    taxonomy: list[dict],
    manifest: list[dict],
    project_root: Path,
    *,
    only: list[str] | None = None,
) -> list[Path]:
    """Generate only selected manifest DOCX entries under a validated project root."""
    root = project_root.resolve()
    sorted_records = _sorted_records(records)
    validate_manifest_contract(manifest, taxonomy)
    taxonomy_by_code = _taxonomy_by_code(taxonomy)
    grouped: dict[str, list[dict]] = defaultdict(list)
    for record in sorted_records:
        if record["subcategory_code"] not in taxonomy_by_code:
            raise ValueError(f"未知小分类: {record['subcategory_code']}")
        if record.get("subcategory_name") != taxonomy_by_code[record["subcategory_code"]]["name"]:
            raise ValueError(f"{record['id']} 的小分类名称与 taxonomy 不一致")
        grouped[record["subcategory_code"]].append(record)

    written: list[Path] = []
    for item in select_manifest_items(manifest, only):
        relative_path = _safe_relative_docx_path(item["path"])
        output_path = (root / relative_path).resolve()
        if root not in output_path.parents:
            raise ValueError(f"不安全的输出路径: {item['path']}")
        if item["scope"] == "overview":
            big_code = item["big_category_code"]
            big_categories = [taxonomy_by_code[code] for code in sorted(taxonomy_by_code) if code.startswith(f"{big_code}-")]
            big_records = [record for record in sorted_records if record["subcategory_code"].startswith(f"{big_code}-")]
            document = build_overview_document(
                {"code": big_code, "name": BIG_CATEGORY_NAMES[big_code], "subcategories": big_categories},
                big_records,
            )
        else:
            subcategory_code = item["subcategory_code"]
            category = taxonomy_by_code.get(subcategory_code)
            if category is None:
                raise ValueError(f"未知小分类: {subcategory_code}")
            document = build_subcategory_document(
                {
                    **category,
                    "big_category_name": BIG_CATEGORY_NAMES[subcategory_code[:2]],
                },
                grouped.get(subcategory_code, []),
            )
        _deterministic_save(document, output_path)
        written.append(output_path)
    return written


def load_inputs() -> tuple[list[dict], list[dict], list[dict], dict]:
    records = _load_json(PLAIN_CATALOG_FILE)
    assignment = _load_json(ASSIGNMENT_FILE)
    manifest = _load_json(MANIFEST_FILE)
    repositories = _load_json(REPOSITORIES_FILE)
    if not isinstance(records, list) or not isinstance(assignment, dict) or not isinstance(manifest, list) or not isinstance(repositories, dict):
        raise ValueError("DOCX 输入数据格式错误")
    taxonomy = assignment.get("taxonomy")
    assignments = assignment.get("assignments")
    if not isinstance(taxonomy, list) or not isinstance(assignments, dict):
        raise ValueError("小分类归属数据格式错误")
    validate_manifest_contract(manifest, taxonomy)
    validate_source_contract(
        records, taxonomy, assignments, repositories, PROJECT_ROOT, expected_total=157
    )
    return records, taxonomy, manifest, repositories


def _parse_only(values: list[str] | None) -> list[str] | None:
    if not values:
        return None
    parsed = [part.strip() for value in values for part in value.split(",") if part.strip()]
    return parsed or None


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--only", nargs="*", help="仅生成键，如 05-overview 05-05")
    args = parser.parse_args(argv)
    records, taxonomy, manifest, _repositories = load_inputs()
    only = _parse_only(args.only)
    written = generate_documents(records, taxonomy, manifest, PROJECT_ROOT, only=only)
    print(f"docx={len(written)} keys={','.join(path.stem for path in written)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
