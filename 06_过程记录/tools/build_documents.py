"""为三个功能分类生成独立的详细 DOCX 调研报告。"""

from __future__ import annotations

import json
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[2]
DATA_DIR = PROJECT_ROOT / "03_候选池" / "deduplicated"
OUTPUT_DIR = PROJECT_ROOT / "05_交付物"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

manifest = json.loads((DATA_DIR / "manifest.json").read_text(encoding="utf-8"))
repositories = json.loads((DATA_DIR / "repositories.json").read_text(encoding="utf-8"))

NAVY = "16324F"
TEAL = "1F6F8B"
LIGHT_TEAL = "DCEAF0"
PALE = "F4F7F9"
BLUE_GRAY = "E8EEF5"
INK = "1F2933"
MUTED = "52606D"
GRID = "CBD5E1"
WHITE = "FFFFFF"
GOLD = "B7791F"
RISK = "9B1C1C"


def set_run_font(run, size=11, bold=False, italic=False, color=INK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text, size=9.5, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), GRID)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                margin = margins.find(qn(f"w:{side}"))
                if margin is None:
                    margin = OxmlElement(f"w:{side}")
                    margins.append(margin)
                margin.set(qn("w:w"), str(value))
                margin.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_hyperlink(paragraph, text, url, size=9.5):
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(size * 2)))
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run_properties.extend([fonts, color, underline, size_node])
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    field_run.extend([begin, instr, separate, value, end])
    paragraph._p.append(field_run)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def configure_document(document, category_name):
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    header_run = header.add_run(f"高校 AI 技能库调研  |  {category_name}  |  GitHub")
    set_run_font(header_run, size=9, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.paragraph_format.space_before = Pt(0)
    add_page_number(footer)


def add_cover(document, category_name, record_count):
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(92)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    set_run_font(kicker.add_run("GITHUB SKILL RESEARCH"), size=10.5, bold=True, color=GOLD)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_run_font(title.add_run(category_name), size=28, bold=True, color=NAVY)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    set_run_font(subtitle.add_run("GitHub Skills 初步调研报告"), size=15, color=TEAL)
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(8)
    set_run_font(meta.add_run(f"入选 {record_count} 项  |  数据日期 2026-08-06"), size=11, bold=True, color=NAVY)
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(0)
    set_run_font(note.add_run("仅含入选项；未进行安装或运行验证"), size=9.5, italic=True, color=MUTED)
    document.add_page_break()


def add_label_paragraph(document, label, text, color=INK):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run_font(paragraph.add_run(f"{label}："), size=11, bold=True, color=NAVY)
    set_run_font(paragraph.add_run(text), size=11, color=color)
    return paragraph


def add_skill_section(document, row):
    heading = document.add_heading(f"{row['id']}｜{row['cn']}（{row['name']}）", level=2)
    heading.paragraph_format.page_break_before = False
    summary = document.add_paragraph()
    summary.paragraph_format.space_before = Pt(0)
    summary.paragraph_format.space_after = Pt(7)
    summary.paragraph_format.left_indent = Inches(0.08)
    summary.paragraph_format.right_indent = Inches(0.08)
    summary.paragraph_format.keep_together = True
    summary.paragraph_format.keep_with_next = True
    paragraph_properties = summary._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), PALE)
    paragraph_properties.append(shading)
    set_run_font(summary.add_run(row["summary"]), size=10.5, bold=True, color=NAVY)

    table = document.add_table(rows=4, cols=2)
    entries = [
        ("生态 / 来源形态", f"{row['ecosystem']}；{row['form']}"),
        ("兼容 / 优先级", f"{row['compat']}；{row['priority']}"),
        ("适用角色", row["roles"]),
        ("典型场景", row["scenario"]),
    ]
    for index, (label, value) in enumerate(entries):
        set_cell_text(table.cell(index, 0), label, size=9.5, bold=True, color=NAVY)
        shade_cell(table.cell(index, 0), BLUE_GRAY)
        set_cell_text(table.cell(index, 1), value, size=9.5)
    for metadata_row in table.rows[:-1]:
        for metadata_cell in metadata_row.cells:
            for metadata_paragraph in metadata_cell.paragraphs:
                metadata_paragraph.paragraph_format.keep_with_next = True
    set_table_geometry(table, [1701, 7659])

    add_label_paragraph(document, "详细功能", row["detail"])
    add_label_paragraph(document, "功能标签", row["tags"])
    add_label_paragraph(document, "适配建议", row["adapt"])
    add_label_paragraph(document, "依赖条件", row["deps"])
    add_label_paragraph(document, "风险与边界", row["risk"], color=RISK)
    add_label_paragraph(
        document,
        "功能验证",
        f"{row['verify']}；已读取说明或包内容，未安装、未执行，不能据此推定真实任务效果。",
    )
    source = document.add_paragraph()
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(4)
    set_run_font(source.add_run("GitHub 来源："), size=9.5, bold=True, color=NAVY)
    add_hyperlink(source, row["skill_url"], row["skill_url"], size=9.5)
    if row.get("related"):
        add_label_paragraph(document, "关联说明", row["related"])


requested_categories = tuple(sys.argv[1:])
known_categories = tuple(dict.fromkeys(item["category"] for item in manifest))
build_categories = requested_categories or known_categories
invalid_categories = set(build_categories) - set(known_categories)
if invalid_categories:
    raise SystemExit(f"未知分类：{', '.join(sorted(invalid_categories))}")

for category in build_categories:
    payload = json.loads((DATA_DIR / f"category_{category}.json").read_text(encoding="utf-8"))
    records = payload["records"]
    category_name = payload["category_name"]
    document = Document()
    configure_document(document, category_name)
    add_cover(document, category_name, len(records))

    document.add_heading("报告摘要", level=1)
    add_label_paragraph(document, "范围", "本轮仅调研 GitHub 上公开可读取的 skill、skill 仓库和可移植 agent 工作流。")
    add_label_paragraph(document, "入选原则", "功能与本分类直接相关，来源可定位，说明足以判断输入输出、依赖和边界，并具有直接使用或合理适配价值。")
    add_label_paragraph(document, "排除规则", "镜像副本、仅关键词命中、说明不足、功能重复且来源更弱的候选只在内部归档，不进入本报告。")
    add_label_paragraph(document, "验证边界", "仅读取 SKILL.md、README 和必要的目录结构；未安装、未运行，也未对任何外部系统执行写入。")

    compatibility = Counter(row["compat"] for row in records)
    priority = Counter(row["priority"] for row in records)
    validation = Counter(row["verify"] for row in records)
    summary_table = document.add_table(rows=2, cols=4)
    summary_values = [
        ("入选数量", len(records)),
        ("兼容 A/B/C/D", f"{compatibility.get('A', 0)}/{compatibility.get('B', 0)}/{compatibility.get('C', 0)}/{compatibility.get('D', 0)}"),
        ("优先级 高/中", f"{priority.get('高', 0)}/{priority.get('中', 0)}"),
        ("包内容/说明核验", f"{validation.get('二级包内容验证', 0)}/{validation.get('说明已核验', 0)}"),
    ]
    for column, (label, value) in enumerate(summary_values):
        set_cell_text(summary_table.cell(0, column), label, size=9, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(summary_table.cell(0, column), TEAL)
        set_cell_text(summary_table.cell(1, column), value, size=11, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(summary_table, [2340, 2340, 2340, 2340])

    document.add_heading("阅读与选型指南", level=1)
    guide_table = document.add_table(rows=5, cols=4)
    guide_rows = [
        ("等级", "判断", "建议", "边界"),
        ("A", "skill 结构可直接使用", "优先进入试点", "仍需补充本校制度与工具配置"),
        ("B", "核心流程可移植", "进入适配队列", "通常需要替换工具、路径或相邻 skill"),
        ("C", "需要明显重写或配套部署", "先做场景验证", "关注维护、权限和部署成本"),
        ("D", "主要提供方法或模板参考", "按需吸收", "不建议原样接入"),
    ]
    for row_index, values in enumerate(guide_rows):
        for column, value in enumerate(values):
            set_cell_text(guide_table.cell(row_index, column), value, size=9.2, bold=row_index == 0, color=WHITE if row_index == 0 else INK)
            if row_index == 0:
                shade_cell(guide_table.cell(row_index, column), TEAL)
            elif column == 0:
                shade_cell(guide_table.cell(row_index, column), BLUE_GRAY)
    set_table_geometry(guide_table, [800, 2600, 2700, 3260])
    repeat_table_header(guide_table.rows[0])

    ecosystem_counts = Counter(row["ecosystem"] for row in records)
    source_form_counts = Counter(row["form"] for row in records)
    document.add_heading("生态与来源形态（独立标注）", level=1)
    add_label_paragraph(document, "生态分布", "；".join(f"{name}：{count}" for name, count in sorted(ecosystem_counts.items())))
    add_label_paragraph(document, "来源形态", "；".join(f"{name}：{count}" for name, count in sorted(source_form_counts.items())))
    add_label_paragraph(document, "说明", "生态表示 skill 的原始或兼容环境；来源形态表示其发布方式与可移植性质。两者不等同，故分别保留。")

    document.add_heading("入选 Skill 详细说明", level=1)
    intro = document.add_paragraph("以下条目按稳定 ID 排序。每项均独立给出功能、使用场景、兼容性、适配建议、依赖、风险边界、验证层级和 GitHub 地址。")
    intro.paragraph_format.keep_with_next = True
    for record in records:
        add_skill_section(document, record)

    document.add_page_break()
    document.add_heading("GitHub 仓库来源清单", level=1)
    repo_counts = Counter(row["repo"] for row in records)
    source_table = document.add_table(rows=1, cols=5)
    for column, label in enumerate(("仓库", "入选数", "Stars", "最近推送", "许可证")):
        set_cell_text(source_table.cell(0, column), label, size=9, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(source_table.cell(0, column), TEAL)
    repeat_table_header(source_table.rows[0])
    for repo in sorted(repo_counts):
        meta = repositories[repo]
        row_cells = source_table.add_row().cells
        row_cells[0].text = ""
        repo_paragraph = row_cells[0].paragraphs[0]
        repo_paragraph.paragraph_format.space_after = Pt(2)
        add_hyperlink(repo_paragraph, repo, f"https://github.com/{repo}", size=8.8)
        set_cell_text(row_cells[1], repo_counts[repo], size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row_cells[2], meta["stars"], size=8.8, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(row_cells[3], meta["pushed"], size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row_cells[4], meta["license"], size=8.8)
    set_table_geometry(source_table, [3900, 720, 900, 1260, 2580])
    citation = document.add_paragraph("来源：GitHub 仓库元数据与本轮读取的 skill 路径；数据日期 2026-08-06。Stars 和许可证会变化，采用前应再次核验。")
    citation.paragraph_format.space_before = Pt(4)
    citation.paragraph_format.space_after = Pt(4)
    set_run_font(citation.runs[0], size=9, italic=True, color=MUTED)

    document.add_heading("结论与后续建议", level=1)
    high_priority = [row for row in records if row["priority"] == "高"]
    add_label_paragraph(document, "优先队列", "、".join(f"{row['id']} {row['name']}" for row in high_priority))
    add_label_paragraph(document, "建议顺序", "先核验许可证和依赖，再按高频校园场景挑选少量 A/B 级候选做受控试点；涉及账号、外部写入、未发表材料或版权内容时另设授权与审计边界。")
    add_label_paragraph(document, "运行验证", "本报告不包含运行结论。若用户后续指定某个候选，再针对该候选设计最小、可回滚、无外部写入的验证方案。")

    output_item = next(item for item in manifest if item["category"] == category and item["format"] == "docx")
    output_path = OUTPUT_DIR / output_item["path"]
    document.save(output_path)
    print(f"{category}: {len(records)} skills -> {output_path}")
