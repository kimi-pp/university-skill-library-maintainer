from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
CODES = {
    "0301": "法学类",
    "0305": "马克思主义理论类",
    "1304": "美术学类",
}
EXPECTED_HEADINGS = [
    "一页读懂本报告",
    "适用专业",
    "能力导航",
    "本学科安全闸门",
    "逐项 Skill 说明",
    "覆盖空白与下一步",
    "验证边界",
    "正式来源索引",
]
FORBIDDEN_FORMAL = [
    "court-records",
    "china-lawyer-analyst",
    "ptreezh",
    "youaifuou/Nhj",
    "chinese-calligraphy-recognition",
    "Tibsfox",
]


def docx_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def audit_geometry(document: Document) -> None:
    section = document.sections[0]
    assert section.page_width == 7772400
    assert section.page_height == 10058400
    assert section.left_margin == 914400
    assert section.right_margin == 914400
    assert section.top_margin == 914400
    assert section.bottom_margin == 914400

    normal = document.styles["Normal"]
    assert round(normal.font.size.pt, 1) == 11.0
    assert round(normal.paragraph_format.line_spacing, 2) == 1.25
    assert round(normal.paragraph_format.space_after.pt, 1) == 6.0
    for style_name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)):
        assert round(document.styles[style_name].font.size.pt, 1) == size

    for table in document.tables:
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        assert tbl_w is not None and tbl_w.get(qn("w:w")) == "9360"
        assert tbl_ind is not None and tbl_ind.get(qn("w:w")) == "120"
        grid = [int(item.get(qn("w:w"))) for item in table._tbl.tblGrid]
        assert sum(grid) == 9360
        for row in table.rows:
            widths = []
            for cell in row.cells:
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                assert tc_w is not None
                widths.append(int(tc_w.get(qn("w:w"))))
            assert widths == grid


def audit_docx(code: str, name: str, records: list[dict]) -> dict:
    path = ROOT / "05_交付物" / f"{code}_{name}_学科专属技能调研" / f"{code}_{name}_学科专属技能调研.docx"
    assert path.exists() and path.stat().st_size > 20_000
    document = Document(path)
    text = docx_text(document)
    audit_geometry(document)

    ids = re.findall(r"DISC-\d{4}-\d{4}", text)
    expected_ids = [record["id"] for record in records]
    assert ids == expected_ids
    assert all(record["cn_name"] in text for record in records)
    assert all(record["canonical_url"] not in text for record in records), "超长 URL 应保留为可点击链接而不是正文堆叠"
    assert all(heading in text for heading in EXPECTED_HEADINGS)
    assert "全部通过（未实测）" in text
    assert "没有安装、没有执行脚本" in text
    assert "内部落选项" in text
    assert not any(item.lower() in text.lower() for item in FORBIDDEN_FORMAL)

    with zipfile.ZipFile(path) as archive:
        rels = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        document_xml = archive.read("word/document.xml").decode("utf-8")
        numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
        assert rels.count("TargetMode=\"External\"") >= len({record["canonical_url"] for record in records})
        assert re.search(r"<w:tblW\b[^>]*\bw:w=\"9360\"[^>]*/>", document_xml)
        assert "w:numFmt w:val=\"bullet\"" in numbering_xml
        assert "w:footerReference" in document_xml or any(name.startswith("word/footer") for name in archive.namelist())

    return {
        "path": path.relative_to(ROOT).as_posix(),
        "size": path.stat().st_size,
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "formal_skills": len(records),
        "hyperlinks_minimum": len({record["canonical_url"] for record in records}),
        "structural_audit": "PASS",
        "visual_render_audit": "BLOCKED：LibreOffice 未安装；Word 后台导出在打开阶段停滞，已终止独立后台实例。",
        "spreadsheet_expected_path": f"05_交付物/{code}_{name}_学科专属技能调研/{code}_{name}_学科专属技能调研.xlsx",
        "spreadsheet_expected_sheets": ["使用说明", "AI技能清单", "专业覆盖", "专业映射", "能力分类", "来源清单", "检索覆盖", "安全准入", "规则说明"],
        "spreadsheet_audit": "BLOCKED：工作区提供的 @oai/artifact-tool 组件目录为空；按表格规范未使用其他库替代。",
    }


def audit_data(code: str, name: str, records: list[dict]) -> dict:
    assert [record["id"] for record in records] == [f"DISC-{code}-{index:04d}" for index in range(1, len(records) + 1)]
    assert all(record["security_grade"] in {"SA", "SB"} for record in records)
    assert all(record["verification_status"] == "全部通过（未实测）" for record in records)
    assert all(2 <= record["quality"] <= 5 for record in records)
    assert all(record["license"].strip() for record in records)
    assert all(record["fixed_version"].strip() for record in records)
    assert all(record["canonical_url"].startswith("https://") for record in records)
    assert all(record["discipline"] == f"{code} {name}" for record in records)

    raw_root = ROOT / "03_候选池" / "raw" / code
    raw_rows = 0
    for path in raw_root.glob("*.jsonl"):
        for line in path.read_text(encoding="utf-8").splitlines():
            if line.strip():
                json.loads(line)
                raw_rows += 1
    assert raw_rows > 0

    kb_root = ROOT / "02_知识库" / "discipline_pilots" / f"{code}_{name}"
    kb_skills = sorted((kb_root / "skills").glob("DISC-*.md"))
    assert len(kb_skills) == len(records)
    assert all(record["id"] == path.stem for record, path in zip(records, kb_skills))
    for record in records:
        assert (ROOT / record["evidence_paths"][0]).exists()
    return {"raw_rows": raw_rows, "knowledge_base_skills": len(kb_skills), "data_audit": "PASS"}


def main() -> None:
    results = {}
    for code, name in CODES.items():
        payload_path = ROOT / "03_候选池" / "deduplicated" / f"{code}_{name}.json"
        payload = json.loads(payload_path.read_text(encoding="utf-8"))
        records = payload["records"]
        assert payload["candidate_count"] == len(records)
        results[code] = {
            **audit_data(code, name, records),
            **audit_docx(code, name, records),
        }
        print(code, results[code])
    output = ROOT / "06_过程记录" / "discipline_artifacts" / "verification_summary.json"
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(results, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
