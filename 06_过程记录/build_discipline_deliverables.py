from __future__ import annotations

import json
import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
AUDIT_DATE = "2026-08-12"
DELIVERY_ROOT = ROOT / "05_交付物"

NAVY = "17324D"
TEAL = "187B80"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "22313F"
MUTED = "5B6872"
PALE_BLUE = "E8EEF5"
PALE_TEAL = "DDEEEF"
PALE_GOLD = "FFF3D6"
PALE_GREEN = "E0F1E7"
GRID = "C9D2DA"
WHITE = "FFFFFF"


DISCIPLINES = {
    "0301": {
        "name": "法学类",
        "json": ROOT / "03_候选池" / "deduplicated" / "0301_法学类.json",
        "folder": "0301_法学类_学科专属技能调研",
        "boundary": "这些工具只帮助检索、整理、起草和复核前准备，不能替代律师意见、司法裁判或有权人员的最终决定。法条、案例和行政规则必须回到有效的权威原文核对。",
        "threshold": "只有直接处理法律检索、案例分析、法源引用、法律文书、证据、模拟法庭或合规任务的项目才进入正式名单；通用写作、通用搜索和通用办公工具不收录。",
        "gate": [
            "案件材料先去标识化；未公开案情、个人信息和商业秘密不得发往未经批准的外部服务。",
            "法规、案例和引证必须核对法域、生效状态、裁判层级、全文和检索截止日期。",
            "合同、诉状、合规意见和庭前材料只能作为草案，由有权限的专业人员复核后使用。",
            "涉及数据库账号时采用学校批准的访问方式和最小权限，不绕过订阅或访问限制。",
        ],
        "choice": "先按任务类型选择：找法源看“法律检索”，读判决看“案例裁判”，规范注释看“法源引用”，写材料看“法律文书”，整理证明关系看“证据整理”，课堂演练看“模拟法庭”，制度或产品审查看“合规风控”。",
    },
    "0305": {
        "name": "马克思主义理论类",
        "json": ROOT / "03_候选池" / "deduplicated" / "0305_马克思主义理论类.json",
        "folder": "0305_马克思主义理论类_学科专属技能调研",
        "boundary": "这些工具只用于辅助阅读和搭建分析框架，不能替代原著研读、可靠译本、学术史辨析或教师与研究者的判断。原文、页码、历史事实和争议观点必须另行核对。",
        "threshold": "只有直接服务于马克思主义经典文献定位、理论概念或马克思主义结构分析的项目才进入正式名单；人物模仿、泛哲学角色扮演和通用写作工具不收录。",
        "gate": [
            "区分原著表述、编者概括和现代延伸，不能把模型整理后的句子当成原文引用。",
            "引文必须回到可靠版本核对原文、译文、卷次、章节和页码。",
            "对理论争议、历史事件和现实判断同时保留证据、竞争解释与不确定性。",
            "课堂或公开材料中的观点归纳须由教师或研究者复核，不能自动形成评价结论。",
        ],
        "choice": "需要按卷章寻找《资本论》内容时选择“经典文献”；需要用矛盾、总体、历史化、意识形态批判和实践反馈来分析问题时选择“社会分析”。当前正式池较小，宁缺毋滥。",
    },
    "1304": {
        "name": "美术学类",
        "json": ROOT / "03_候选池" / "deduplicated" / "1304_美术学类.json",
        "folder": "1304_美术学类_学科专属技能调研",
        "boundary": "这些工具只辅助研究、记录和方案草拟，不能据此认定作品真伪、合法所有权、保存修复方案、市场价值或最终展陈决定；这些事项必须由合格人员复核。",
        "threshold": "只有直接处理美术史、艺术品元数据、数字保存、书画研究、策展或艺术图像权利的项目才进入正式名单；通用图像生成、通用网页设计和通用项目管理工具不收录。",
        "gate": [
            "作品图像、馆藏记录和来源链信息先判断是否允许上传或联网处理；未公开材料留在批准环境。",
            "艺术家、年代、材料、尺寸、来源、展览史和修复记录须回到馆藏、档案或权威目录核对。",
            "版权、公共领域和开放许可必须逐件核验，不能用“可下载”代替“可使用”。",
            "真伪、权属、文物保护和修复建议只能由专业人员作最终判断，并保留完整记录。",
        ],
        "choice": "先按对象选择：研究作品与艺术家看“美术史论”，建馆藏条目看“作品元数据”，做保存计划看“数字保存”，书法教学与观察看“书画研究”，做策展方案和展签看“展览作品集”，找可用馆藏图像看“版权溯源”。",
    },
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, size: float = 11, color: str = INK, bold: bool | None = None, italic: bool | None = None, font: str = "Calibri") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_hyperlink(paragraph, text: str, url: str, color: str = BLUE) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "等线")
    r_pr.append(r_fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    r_pr.append(size)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    next_abs = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    def add_definition(abstract_id: int, num_id: int, fmt: str, text: str, font: str | None = None) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        if font:
            r_pr = OxmlElement("w:rPr")
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), font)
            r_fonts.set(qn("w:hAnsi"), font)
            r_pr.append(r_fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(abstract_id))
        num.append(abs_id)
        numbering.append(num)

    add_definition(next_abs, next_num, "bullet", "•", "Symbol")
    add_definition(next_abs + 1, next_num + 1, "decimal", "%1.")
    return next_num, next_num + 1


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_bullet(doc: Document, text: str, bullet_num_id: int) -> None:
    paragraph = doc.add_paragraph()
    apply_num(paragraph, bullet_num_id)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run_font(paragraph.add_run(text))


def add_label_paragraph(doc: Document, label: str, text: str, *, after: float = 5, keep: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.keep_together = keep
    set_run_font(paragraph.add_run(label), bold=True, color=DARK_BLUE)
    set_run_font(paragraph.add_run(text))


def add_callout(doc: Document, title: str, text: str, fill: str = PALE_TEAL) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.25
    set_run_font(p.add_run(f"{title}  "), bold=True, color=NAVY)
    set_run_font(p.add_run(text))
    set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])
    set_run_font(paragraph.add_run(" 页"), size=9, color=MUTED)


def configure_header_footer(doc: Document, discipline: str) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.paragraph_format.space_after = Pt(0)
        header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
        set_run_font(header.add_run(discipline), size=9, color=MUTED, bold=True)
        set_run_font(header.add_run("\t学科专属 Skill 调研"), size=9, color=MUTED)
        add_page_number(section.footer.paragraphs[0])


def add_cover(doc: Document, discipline: str, count: int, grades: Counter) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    set_run_font(p.add_run("高校 AI 技能库调研"), size=11, color=TEAL, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run(discipline), size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(32)
    set_run_font(p.add_run("学科专属 Skill 调研报告"), size=16, color=DARK_BLUE)

    add_callout(doc, "正式结论", f"收录 {count} 项学科专属 Skill，其中 SA {grades.get('SA', 0)} 项、SB {grades.get('SB', 0)} 项。全部通过静态准入，但均未安装、未运行。", PALE_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(f"核验日期：{AUDIT_DATE}"), size=10.5, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("公开来源全网检索｜固定版本留档｜静态安全审查"), size=10, color=MUTED)
    doc.add_page_break()


def add_summary_tables(doc: Document, records: list[dict]) -> None:
    majors = records[0]["majors"]
    major_table = doc.add_table(rows=1, cols=2)
    major_table.rows[0].cells[0].text = "专业代码"
    major_table.rows[0].cells[1].text = "专业名称"
    for item in majors:
        match = re.match(r"(\S+)\s+(.+)", item)
        row = major_table.add_row().cells
        row[0].text = match.group(1) if match else ""
        row[1].text = match.group(2) if match else item
    set_repeat_table_header(major_table.rows[0])
    for cell in major_table.rows[0].cells:
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, bold=True, color=WHITE)
    for row in major_table.rows[1:]:
        for index, cell in enumerate(row.cells):
            for run in cell.paragraphs[0].runs:
                set_run_font(run)
            if index == 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(major_table, [2200, 7160])

    doc.add_heading("能力导航", level=2)
    grouped = defaultdict(list)
    for record in records:
        grouped[(record["category"], record["category_name"])].append(record["cn_name"])
    capability_table = doc.add_table(rows=1, cols=3)
    for cell, value in zip(capability_table.rows[0].cells, ("能力小类", "数量", "正式 Skill")):
        cell.text = value
    for (code, name), items in sorted(grouped.items()):
        row = capability_table.add_row().cells
        row[0].text = f"{code} {name}"
        row[1].text = str(len(items))
        row[2].text = "、".join(items)
    set_repeat_table_header(capability_table.rows[0])
    for cell in capability_table.rows[0].cells:
        set_cell_shading(cell, TEAL)
        for run in cell.paragraphs[0].runs:
            set_run_font(run, bold=True, color=WHITE)
    for row in capability_table.rows[1:]:
        for index, cell in enumerate(row.cells):
            for run in cell.paragraphs[0].runs:
                set_run_font(run)
            if index == 1:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(capability_table, [2800, 900, 5660])


def add_skill_profile(doc: Document, record: dict, bullet_num_id: int) -> None:
    doc.add_heading(f"{record['id']}｜{record['cn_name']}", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    set_run_font(p.add_run(f"能力小类：{record['category']} {record['category_name']}"), size=10.5, color=TEAL, bold=True)
    set_run_font(p.add_run(f"    安全等级：{record['security_grade']}    质量评分：{record['quality']}/5    采用建议：{record['adoption_plain']}"), size=10.5, color=MUTED)

    add_label_paragraph(doc, "它能做什么：", record["purpose"], keep=True)
    add_label_paragraph(doc, "适合在什么时候使用：", f"当任务明确属于“{record['category_name']}”，并且已有必要材料、责任人和人工复核安排时使用。")
    add_label_paragraph(doc, "需要准备：", record["inputs"])
    add_label_paragraph(doc, "可以得到：", record["outputs"])
    add_label_paragraph(doc, "适用专业：", "、".join(record["majors"]))
    add_label_paragraph(doc, "为什么入选：", f"说明内容完整、固定版本可取得、许可条件可确认，并且直接服务于本学科任务；当前质量评分为 {record['quality']}/5。")

    doc.add_heading("使用前必须确认", level=3)
    add_bullet(doc, record["limits"], bullet_num_id)
    add_bullet(doc, f"文件与网络：{record['file_behavior']} {record['network_behavior']}", bullet_num_id)
    add_bullet(doc, f"账号与凭据：{record['credential_behavior']}", bullet_num_id)
    add_bullet(doc, f"安全结论：{record['security_plain']}", bullet_num_id)

    doc.add_heading("本次验证做到哪一步", level=3)
    add_label_paragraph(doc, "验证结果：", record["verification_status"])
    add_label_paragraph(doc, "验证方式：", record["verification_depth"])
    add_label_paragraph(doc, "静态清点：", f"固定版本包共清点 {record['package_files']} 个文件，其中脚本文件 {record['script_files']} 个。没有安装或运行候选。")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    set_run_font(p.add_run("来源与许可："), bold=True, color=DARK_BLUE)
    set_run_font(p.add_run(f"{record['platform']}｜{record.get('repo', record.get('repo_url', ''))}｜固定版本 {record['fixed_version']}｜{record['license']}｜"))
    add_hyperlink(p, "打开固定版本来源", record["canonical_url"])

    add_callout(doc, "采用判断", f"{record['adoption_plain']}。{record['alternative']}", PALE_GREEN if record["security_grade"] == "SA" else PALE_GOLD)


def add_gaps(doc: Document, records: list[dict], cfg: dict, bullet_num_id: int) -> None:
    doc.add_heading("覆盖空白与下一步", level=1)
    if len(records) <= 3:
        add_callout(doc, "重要结论", "本学科目前可确认的高质量开源专属 Skill 很少。正式池小不是遗漏落选项，而是坚持许可、完整性、固定版本和学科专属性后的结果。", PALE_GOLD)
    add_bullet(doc, "优先补齐目前没有正式 Skill 的能力小类，但继续坚持学科专属，不用通用工具凑数。", bullet_num_id)
    add_bullet(doc, "后续如出现新版本，先重新固定版本并复做静态审查，再更新正式名单。", bullet_num_id)
    add_bullet(doc, "只有在用户另行明确指令后，才可在隔离环境中做最小运行验证；当前报告不包含运行结论。", bullet_num_id)
    add_bullet(doc, "正式接入学校系统前，还需完成数据分类、账号权限、日志留存、人工审批和退出机制设计。", bullet_num_id)

    doc.add_heading("验证边界", level=1)
    add_label_paragraph(doc, "本次做了：", "全网公开来源检索、版本固定、许可证确认、说明阅读、包内容清点、静态安全扫描和数据流推演。")
    add_label_paragraph(doc, "本次没有做：", "没有安装、没有执行脚本、没有调用外部接口、没有提交真实师生或业务数据，也没有检验实际输出质量。")
    add_label_paragraph(doc, "怎样理解“全部通过（未实测）”：", "表示它通过了本轮资料与静态安全准入，可以进入后续人工评估；不表示已经证明能在学校环境稳定、安全或准确运行。")
    add_callout(doc, "正式使用原则", cfg["boundary"], PALE_BLUE)


def add_sources(doc: Document, records: list[dict]) -> None:
    doc.add_heading("正式来源索引", level=1)
    seen = set()
    for record in records:
        key = (record["platform"], record.get("repo_url", ""), record["fixed_version"])
        if key in seen:
            continue
        seen.add(key)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        set_run_font(p.add_run(f"{record['platform']}｜{record.get('repo', record.get('repo_url', ''))}｜{record['fixed_version']}｜{record['license']}｜"))
        add_hyperlink(p, "查看来源", record.get("repo_url", record["canonical_url"]))


def build_document(code: str, cfg: dict) -> Path:
    payload = json.loads(cfg["json"].read_text(encoding="utf-8"))
    records = payload["records"]
    grades = Counter(record["security_grade"] for record in records)

    doc = Document()
    setup_styles(doc)
    bullet_num_id, _ = configure_numbering(doc)
    add_cover(doc, payload["discipline"], len(records), grades)

    doc.add_heading("一页读懂本报告", level=1)
    add_callout(doc, "收录结果", f"正式收录 {len(records)} 项：SA {grades.get('SA', 0)} 项、SB {grades.get('SB', 0)} 项。正式名单没有 SB-A 或 X，也没有内部落选项。", PALE_TEAL)
    add_label_paragraph(doc, "调研对象：", "面向本学科学生、教师、科研人员及相关专业支持岗位的学科专属 Skill。")
    add_label_paragraph(doc, "调研范围：", "公开网络来源，包括 ClawHub、GitHub、Hugging Face Spaces 及公开代码与项目平台；按既定顺序检索并做补充轮次。")
    add_label_paragraph(doc, "准入门槛：", cfg["threshold"])
    add_label_paragraph(doc, "状态说明：", "全部通过（未实测）。这表示已完成静态准入，不表示已经运行验证。")
    add_label_paragraph(doc, "如何选择：", cfg["choice"])

    doc.add_heading("适用专业", level=2)
    add_summary_tables(doc, records)

    doc.add_heading("本学科安全闸门", level=1)
    add_callout(doc, "先看这一页", cfg["boundary"], PALE_GOLD)
    for item in cfg["gate"]:
        add_bullet(doc, item, bullet_num_id)

    doc.add_heading("逐项 Skill 说明", level=1)
    add_label_paragraph(doc, "阅读方法：", "先看能力小类和采用建议，再看“使用前必须确认”。SA 表示静态风险较低；SB 表示存在受控文件读写、联网或脚本等条件，使用前需要按说明做限制。")
    for index, record in enumerate(records):
        if index and index % 3 == 0:
            doc.add_page_break()
        add_skill_profile(doc, record, bullet_num_id)

    add_gaps(doc, records, cfg, bullet_num_id)
    add_sources(doc, records)
    configure_header_footer(doc, payload["discipline"])

    output_dir = DELIVERY_ROOT / cfg["folder"]
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{code}_{cfg['name']}_学科专属技能调研.docx"
    doc.core_properties.title = f"{payload['discipline']}学科专属 Skill 调研报告"
    doc.core_properties.subject = "高校 AI 技能库调研"
    doc.core_properties.author = "高校 AI 技能库调研项目"
    doc.core_properties.comments = "公开来源固定版本静态审查；候选未安装、未运行。"
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    for discipline_code, config in DISCIPLINES.items():
        result = build_document(discipline_code, config)
        print(result)
