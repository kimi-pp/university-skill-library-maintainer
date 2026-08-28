"""Chinese Word/Excel report builders and affected-scope selection."""

from __future__ import annotations

from collections import Counter
from dataclasses import asdict, is_dataclass
from datetime import date, datetime
from hashlib import sha256
import json
import os
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
from typing import Any, Iterable, Mapping

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .ledger import LedgerStore
from .paths import ProjectPaths, assert_ordinary_path, contained_child


DAILY_WORD_SECTIONS = (
    "本轮执行摘要",
    "专业目录变化",
    "四平台覆盖情况",
    "新增正式推荐",
    "已有 Skill 版本更新",
    "发现新版本但未升级",
    "条件候选",
    "需适配候选",
    "全局去重和来源别名",
    "受影响专业类",
    "排除原因汇总",
    "失败、覆盖降级和人工复核事项",
    "核验边界和未运行声明",
)

DAILY_SHEETS = (
    "使用说明",
    "执行概览",
    "目录变化",
    "新增正式推荐",
    "版本更新",
    "发现更新未升级",
    "条件候选",
    "需适配候选",
    "去重与来源别名",
    "受影响专业类",
    "排除原因汇总",
    "来源请求审计",
)

_REPORT_KEYS = {
    "catalog_changes": ("catalog_changes", "目录变化"),
    "formal_additions": ("formal_additions", "新增正式推荐", "正式推荐"),
    "version_updates": ("version_updates", "版本更新"),
    "updates_not_applied": ("updates_not_applied", "发现更新未升级"),
    "conditional_candidates": ("conditional_candidates", "条件候选"),
    "adaptation_candidates": ("adaptation_candidates", "需适配候选"),
    "aliases": ("aliases", "来源别名", "去重与来源别名"),
    "affected_scopes": ("affected_scopes", "受影响专业类"),
    "exclusions": ("exclusions", "排除项", "排除原因"),
    "manual_reviews": ("manual_reviews", "人工复核事项", "失败与人工复核"),
    "source_requests": ("source_requests", "来源请求审计"),
}

_MATERIAL_SKILL_FIELDS = (
    "入库层级",
    "固定版本",
    "固定版本内容指纹",
    "许可证",
    "安全等级",
    "安全限制条件",
    "外部依赖",
    "外部联网/API 调用",
    "远程服务端点",
    "本地专业软件或运行时依赖",
    "本地脚本/插件接口",
    "维护状态",
    "推荐优先级",
    "适配建议",
    "验证级别",
    "验证状态",
    "风险提示",
    "可执行行为",
    "网络与数据行为",
    "凭据行为",
    "文件行为",
    "质量评分",
    "实施准备度",
)

_EXCLUSION_REASON_ALIASES = {
    "许可证不明确": "许可证不明确",
    "license_unknown": "许可证不明确",
    "license-unverified": "许可证不明确",
    "安全要求未满足": "安全要求未满足",
    "安全不符合": "安全要求未满足",
    "security_rejected": "安全要求未满足",
    "证据不足": "证据不足",
    "insufficient_evidence": "证据不足",
    "不在研究范围": "不在研究范围",
    "out_of_scope": "不在研究范围",
    "重复条目": "重复条目",
    "duplicate": "重复条目",
    "来源不可核验": "来源不可核验",
    "source_unverifiable": "来源不可核验",
    "质量门槛未满足": "质量门槛未满足",
    "quality_below_threshold": "质量门槛未满足",
    "其他合规原因": "其他合规原因",
}

_APPROVED_CATEGORY_CODES = frozenset({
    "01", "02", "03", "04", "05", "06", "07", "08", "09", "10", "12", "13", "14",
})


class ReportBuildError(RuntimeError):
    """A report could not be built with the approved portable runtime."""


def _plain_mapping(value: object) -> dict[str, Any]:
    if value is None:
        return {}
    if isinstance(value, Mapping):
        return {str(key): item for key, item in value.items()}
    if is_dataclass(value):
        return asdict(value)
    if hasattr(value, "__dict__"):
        return {str(key): item for key, item in vars(value).items() if not key.startswith("_")}
    raise TypeError("报告摘要必须是 Mapping、dataclass 或带公开属性的对象")


def _json_value(value: object) -> object:
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    if isinstance(value, Path):
        return str(value)
    if is_dataclass(value):
        return {key: _json_value(item) for key, item in asdict(value).items()}
    if isinstance(value, Mapping):
        return {str(key): _json_value(item) for key, item in value.items()}
    if isinstance(value, (list, tuple, set, frozenset)):
        return [_json_value(item) for item in value]
    return value


def _rows(summary: Mapping[str, Any], canonical: str) -> list[dict[str, Any]]:
    for key in _REPORT_KEYS[canonical]:
        if key not in summary:
            continue
        value = summary[key]
        if value is None:
            return []
        if isinstance(value, Mapping):
            return [dict(value)]
        if isinstance(value, (str, Path)):
            return [{"名称": str(value)}]
        result = []
        for item in value:
            result.append(_plain_mapping(item) if not isinstance(item, str) else {"名称": item})
        return result
    return []


def _summary_payload(summary: object) -> dict[str, Any]:
    raw = _plain_mapping(summary)
    payload = {key: _rows(raw, key) for key in _REPORT_KEYS}
    scopes = []
    for item in payload["affected_scopes"]:
        scopes.append(str(item.get("专业类") or item.get("名称") or item.get("scope") or "").strip())
    payload["affected_scopes"] = [item for item in scopes if item]
    payload.update({
        "run_id": str(raw.get("run_id") or raw.get("运行标识") or "未提供"),
        "generated_at": _json_value(raw.get("generated_at") or raw.get("生成时间") or datetime.now()),
        "blocked": bool(raw.get("blocked", False)),
        "warnings": list(raw.get("warnings") or ()),
        "source_statuses": dict(raw.get("source_statuses") or raw.get("来源状态") or {}),
    })
    # Exclusion candidate names are deliberately never serialized into a deliverable.
    payload["exclusions"] = [{"原因": _standard_exclusion_reason(row)} for row in payload["exclusions"]]
    return _json_value(payload)  # type: ignore[return-value]


def _standard_exclusion_reason(row: Mapping[str, Any]) -> str:
    raw = str(
        row.get("原因代码") or row.get("reason_code") or row.get("原因类别")
        or row.get("原因") or row.get("reason") or ""
    ).strip()
    return _EXCLUSION_REASON_ALIASES.get(raw.casefold(), _EXCLUSION_REASON_ALIASES.get(raw, "其他合规原因"))


def _set_run_font(run, *, size: float = 11, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _add_heading_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(item.get(qn("w:abstractNumId"))) for item in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "space")
    paragraph_props = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "360")
    indent.set(qn("w:hanging"), "360")
    paragraph_props.extend((tabs, indent))
    level.extend((start, num_fmt, lvl_text, suffix, paragraph_props))
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _number_heading(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num_ref))


def _set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: tuple[int, ...]) -> None:
    if sum(widths) != 9360:
        raise ValueError("Word 表格列宽必须合计 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    width = tbl_pr.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), "9360")
    width.set(qn("w:type"), "dxa")
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for item in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(item))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_width.set(qn("w:w"), str(widths[index]))
            tc_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.10
                for run in paragraph.runs:
                    _set_run_font(run)


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _add_field(paragraph, instruction: str, placeholder: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run_text = OxmlElement("w:t")
    run_text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    run.extend((begin, text, separate, run_text, end))


def _metadata_table(document: Document, payload: Mapping[str, Any]) -> None:
    table = document.add_table(rows=4, cols=2)
    rows = (
        ("运行标识", str(payload["run_id"])),
        ("生成时间", str(payload["generated_at"])),
        ("报告口径", "正式推荐、条件候选、需适配候选分别统计"),
        ("核验方式", "静态证据审阅；候选未安装、未运行"),
    )
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        _shade(row.cells[0], "F2F4F7")
        for run in row.cells[0].paragraphs[0].runs:
            _set_run_font(run, bold=True)
    _set_table_geometry(table, (2700, 6660))


def _add_detail_table(document: Document, rows: Iterable[tuple[str, object]]) -> None:
    values = tuple(rows)
    table = document.add_table(rows=max(1, len(values)), cols=2)
    for row, (label, value) in zip(table.rows, values):
        row.cells[0].text = str(label)
        row.cells[1].text = str(value or "未提供")
        _shade(row.cells[0], "F2F4F7")
        for run in row.cells[0].paragraphs[0].runs:
            _set_run_font(run, bold=True)
    _set_table_geometry(table, (2700, 6660))


def _display_name(row: Mapping[str, Any]) -> str:
    return str(row.get("规范名称") or row.get("Skill名称") or row.get("候选名称") or row.get("内部标识") or "未命名项")


def _add_skill_items(document: Document, rows: list[dict[str, Any]]) -> None:
    if not rows:
        document.add_paragraph("本轮无相关记录。")
        return
    for row in rows:
        heading = document.add_paragraph(style="Heading 2")
        heading.add_run(_display_name(row))
        _add_detail_table(document, (
            ("稳定 ID", row.get("内部标识") or "未分配"),
            ("英文原名", row.get("规范名称") or row.get("Skill名称") or "未提供"),
            ("用途", row.get("用途") or row.get("简要功能") or row.get("详细功能摘要") or "未提供"),
            ("适用人员", row.get("适用人员") or row.get("适用用户角色") or "未提供"),
            ("输入", row.get("输入") or "未提供"),
            ("输出", row.get("输出") or "未提供"),
            ("限制", row.get("使用限制") or row.get("安全限制条件") or row.get("适配建议") or "未提供"),
            ("固定版本", row.get("固定版本") or row.get("新版本") or row.get("发现版本") or "未提供"),
            ("许可证", row.get("许可证") or "未提供"),
            ("URL", row.get("Canonical source") or row.get("来源地址") or row.get("发现地址") or "未提供"),
            ("层级", row.get("入库层级") or row.get("观察状态") or "未提供"),
            ("原因/结论", row.get("原因") or row.get("结论") or "未提供"),
        ))


def build_daily_docx(summary: object, output: str | Path) -> Path:
    """Build the fixed-order Chinese daily Word report with deterministic geometry."""

    payload = _summary_payload(summary)
    destination = Path(output).resolve()
    destination.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    _configure_styles(document)
    heading_num_id = _add_heading_numbering(document)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    _set_run_font(header.add_run("高校专业 Skill 库｜自动维护日报"), size=9, color="667085")
    _set_run_font(header.add_run(f"    {payload['run_id']}"), size=9, color="667085")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(footer.add_run("第 "), size=9, color="667085")
    _add_field(footer, " PAGE ", "1")
    _set_run_font(footer.add_run(" 页"), size=9, color="667085")

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)
    _set_run_font(title.add_run("高校专业 Skill 库自动查验报告"), size=23, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    _set_run_font(subtitle.add_run("每日维护摘要与人工复核边界"), size=14, color="373737")
    _metadata_table(document, payload)

    formal = payload["formal_additions"]
    conditional = payload["conditional_candidates"]
    adaptation = payload["adaptation_candidates"]
    sections: dict[str, Any] = {
        "本轮执行摘要": lambda: _add_detail_table(document, (
            ("运行状态", "阻断" if payload["blocked"] else "完成"),
            ("新增正式推荐", len(formal)),
            ("条件候选", len(conditional)),
            ("需适配候选", len(adaptation)),
            ("受影响专业类", len(payload["affected_scopes"])),
        )),
        "专业目录变化": lambda: _add_generic_rows(document, payload["catalog_changes"]),
        "四平台覆盖情况": lambda: _add_source_statuses(document, payload["source_statuses"]),
        "新增正式推荐": lambda: _add_skill_items(document, formal),
        "已有 Skill 版本更新": lambda: _add_skill_items(document, payload["version_updates"]),
        "发现新版本但未升级": lambda: _add_skill_items(document, payload["updates_not_applied"]),
        "条件候选": lambda: _add_skill_items(document, conditional),
        "需适配候选": lambda: _add_skill_items(document, adaptation),
        "全局去重和来源别名": lambda: _add_generic_rows(document, payload["aliases"]),
        "受影响专业类": lambda: _add_scope_rows(document, payload["affected_scopes"]),
        "排除原因汇总": lambda: _add_exclusion_summary(document, payload["exclusions"]),
        "失败、覆盖降级和人工复核事项": lambda: _add_manual_items(document, payload),
        "核验边界和未运行声明": lambda: document.add_paragraph(
            "本报告仅依据固定版本、来源快照、许可证与静态安全证据进行查验。候选 Skill 未安装、未运行；"
            "未调用候选自身外部服务，也未上传真实教学或科研数据。正式采用前仍须由责任人复核许可证、依赖和数据边界。"
        ),
    }
    for name in DAILY_WORD_SECTIONS:
        heading = document.add_paragraph(name, style="Heading 1")
        _number_heading(heading, heading_num_id)
        sections[name]()
    document.save(destination)
    return destination


def _add_generic_rows(document: Document, rows: list[dict[str, Any]]) -> None:
    if not rows:
        document.add_paragraph("本轮无相关记录。")
        return
    for row in rows:
        fields = [(key, value) for key, value in row.items() if value not in (None, "")]
        _add_detail_table(document, fields or (("说明", "已记录"),))


def _add_source_statuses(document: Document, statuses: Mapping[str, Any]) -> None:
    ordered = ("SkillHub", "ClawHub", "GitHub", "Hugging Face Spaces")
    _add_detail_table(document, ((platform, statuses.get(platform, "未请求")) for platform in ordered))


def _add_scope_rows(document: Document, scopes: list[str]) -> None:
    if scopes:
        _add_detail_table(document, (("专业类", scope) for scope in scopes))
    else:
        document.add_paragraph("本轮无受影响专业类；既有专业类交付不重写。")


def _add_exclusion_summary(document: Document, exclusions: list[dict[str, Any]]) -> None:
    counts = Counter(str(row.get("原因") or "未分类原因") for row in exclusions)
    if counts:
        _add_detail_table(document, ((reason, count) for reason, count in sorted(counts.items())))
    else:
        document.add_paragraph("本轮无排除原因记录。为保护审查边界，本节只汇总原因和数量，不列排除项名称。")


def _add_manual_items(document: Document, payload: Mapping[str, Any]) -> None:
    items = list(payload["manual_reviews"])
    items.extend({"事项": warning} for warning in payload.get("warnings", ()))
    for platform, status in payload["source_statuses"].items():
        if str(status).lower() not in {"complete", "success", "ok", "成功", "完整"}:
            items.append({"事项": f"{platform} 覆盖状态为 {status}，需保留降级说明。"})
    _add_generic_rows(document, items)


def _require_runtime_path(variable: str, *, directory: bool) -> Path:
    raw = os.environ.get(variable, "").strip()
    if not raw:
        raise ReportBuildError(f"缺少环境变量 {variable}；无法使用批准的可移植 Node 运行时")
    path = Path(raw).absolute()
    valid = path.is_dir() if directory else path.is_file()
    if not valid:
        kind = "目录" if directory else "文件"
        raise ReportBuildError(f"环境变量 {variable} 指向的{kind}不存在：{path}")
    try:
        assert_ordinary_path(path, require_directory=directory)
    except ValueError as exc:
        raise ReportBuildError(f"环境变量 {variable} 不得指向链接或重解析点：{path}") from exc
    return path


def _verify_node_identity(node: Path) -> None:
    expected = "node.exe" if os.name == "nt" else "node"
    if node.name.casefold() != expected.casefold():
        raise ReportBuildError(f"SKILL_MAINTAINER_NODE 必须指向名为 {expected} 的 Node 可执行文件")
    try:
        result = subprocess.run(
            [str(node), "--version"], capture_output=True, text=True, encoding="utf-8",
            errors="replace", check=False, timeout=15,
        )
    except (OSError, subprocess.SubprocessError) as exc:
        raise ReportBuildError("SKILL_MAINTAINER_NODE 不是可执行的 Node 运行时") from exc
    version = (result.stdout or "").strip()
    if result.returncode or re.fullmatch(r"v\d+\.\d+\.\d+(?:[-+][0-9A-Za-z.-]+)?", version) is None:
        raise ReportBuildError("SKILL_MAINTAINER_NODE 的可执行身份不是受支持的 Node 运行时")


def _make_node_modules_link(link: Path, target: Path) -> None:
    try:
        link.symlink_to(target, target_is_directory=True)
    except (NotImplementedError, OSError) as exc:
        raise ReportBuildError(f"无法以无 shell 方式创建临时 node_modules 链接：{exc}") from exc


def _remove_node_modules_link(link: Path) -> None:
    if not link.exists() and not link.is_symlink():
        return
    if link.is_symlink():
        link.unlink()
    else:
        os.rmdir(link)


def _run_xlsx_builder(payload: Mapping[str, Any], destination: Path, *, verify_dir: Path | None = None) -> None:
    node = _require_runtime_path("SKILL_MAINTAINER_NODE", directory=False)
    modules = _require_runtime_path("SKILL_MAINTAINER_NODE_MODULES", directory=True)
    _verify_node_identity(node)
    artifact_tool = modules / "@oai" / "artifact-tool"
    if not artifact_tool.is_dir():
        raise ReportBuildError("SKILL_MAINTAINER_NODE_MODULES 中缺少 @oai/artifact-tool")
    try:
        assert_ordinary_path(artifact_tool, require_directory=True)
    except ValueError as exc:
        raise ReportBuildError("@oai/artifact-tool 必须位于普通、无重解析点的模块目录") from exc
    source_builder = Path(__file__).with_name("daily_xlsx_builder.mjs")
    if not source_builder.is_file():
        raise ReportBuildError(f"缺少 Excel 生成器：{source_builder.name}")
    with tempfile.TemporaryDirectory(prefix="skill-maintainer-report-") as raw_temp:
        work = Path(raw_temp)
        link = work / "node_modules"
        _make_node_modules_link(link, modules)
        try:
            builder = work / source_builder.name
            request = work / "report-input.json"
            shutil.copyfile(source_builder, builder)
            request.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
            command = [str(node), str(builder), str(request), str(destination)]
            if verify_dir is not None:
                command.extend(("--verify-dir", str(verify_dir)))
            result = subprocess.run(
                command,
                cwd=work,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                check=False,
            )
            if result.returncode:
                detail = (result.stderr or result.stdout or "无错误输出").strip()
                raise ReportBuildError(f"artifact-tool Excel 生成失败（退出码 {result.returncode}）：{detail}")
        finally:
            _remove_node_modules_link(link)


def build_daily_xlsx(summary: object, output: str | Path) -> Path:
    """Build the 12-sheet daily workbook exclusively through artifact-tool."""

    destination = Path(output).resolve()
    destination.parent.mkdir(parents=True, exist_ok=True)
    _run_xlsx_builder(_summary_payload(summary), destination)
    return destination


def _ledger_rows(ledger: object, sheet: str) -> list[dict[str, Any]]:
    if hasattr(ledger, "rows") and callable(getattr(ledger, "rows")):
        return [dict(row) for row in ledger.rows(sheet)]
    mapping = _plain_mapping(ledger)
    value = mapping.get(sheet, ())
    return [_plain_mapping(row) for row in value]


def _scope_name(row: Mapping[str, Any]) -> str:
    direct = str(row.get("专业类") or row.get("scope") or "").strip()
    if direct:
        return direct
    code = str(row.get("专业代码") or "").strip()
    name = str(row.get("专业名称") or "").strip()
    return " ".join(item for item in (code, name) if item)


def _scope_code(row: Mapping[str, Any]) -> str:
    direct = str(row.get("专业代码") or "").strip()
    if direct:
        return direct
    match = re.match(r"^(\d{2,6})(?:\s|$)", _scope_name(row))
    return match.group(1) if match else ""


def _active_catalog_rows(catalog_snapshot: object | None) -> tuple[object, ...]:
    """Read the exact active rows from this run's captured catalog."""

    if catalog_snapshot is None:
        return ()
    if isinstance(catalog_snapshot, Mapping):
        staged = catalog_snapshot.get("staged_snapshot")
        if staged is not None:
            staged_mapping = _plain_mapping(staged)
            rows = staged_mapping.get("rows", ())
        else:
            rows = catalog_snapshot.get("rows", ())
    else:
        staged = getattr(catalog_snapshot, "staged_snapshot", None)
        rows = getattr(staged, "rows", ()) if staged is not None else getattr(catalog_snapshot, "rows", ())
    try:
        return tuple(rows)
    except TypeError:
        return ()


def _approved_catalog_scopes(catalog_snapshot: object | None) -> dict[str, str]:
    names: dict[str, set[str]] = {}
    for value in _active_catalog_rows(catalog_snapshot):
        row = _plain_mapping(value)
        category = str(row.get("category_code") or "").strip()
        if category in _APPROVED_CATEGORY_CODES - {"14"}:
            code = str(row.get("class_code") or "").strip()
            name = str(row.get("class_name") or "").strip()
            valid = re.fullmatch(r"\d{4}", code) is not None and code[:2] == category
        elif category == "14":
            code = str(row.get("major_code") or "").strip()
            name = str(row.get("major_name") or "").strip()
            valid = re.fullmatch(r"14\d{4}", code) is not None
        else:
            continue
        if valid and name:
            names.setdefault(code, set()).add(name)
    return {
        code: f"{code} {next(iter(scope_names))}"
        for code, scope_names in names.items()
        if len(scope_names) == 1
    }


def _approved_scope_name(row: Mapping[str, Any], catalog_snapshot: object | None) -> str:
    code = _scope_code(row)
    if code == "99":
        return "99 跨学科通用"
    return _approved_catalog_scopes(catalog_snapshot).get(code, "")


def _is_approved_scope_mapping(row: Mapping[str, Any], catalog_snapshot: object | None) -> bool:
    return bool(_approved_scope_name(row, catalog_snapshot))


def _scope_index(
    snapshot: object,
    catalog_snapshot: object | None,
) -> tuple[dict[str, set[str]], dict[str, dict[str, Any]], dict[str, tuple[tuple[str, str], ...]]]:
    skills = {
        str(row.get("内部标识") or "").strip(): row
        for row in _ledger_rows(snapshot, "当前Skill")
        if str(row.get("内部标识") or "").strip()
    }
    scopes: dict[str, set[str]] = {}
    mapping_fingerprints: dict[str, list[tuple[str, str]]] = {}
    for row in _ledger_rows(snapshot, "专业任务映射"):
        stable_id = str(row.get("内部标识") or "").strip()
        scope = _approved_scope_name(row, catalog_snapshot)
        if not stable_id or not scope:
            continue
        scopes.setdefault(stable_id, set()).add(scope)
        material = tuple(sorted(
            (str(key), str(value)) for key, value in row.items()
            if key not in {"映射标识", "mapping_id"} and value not in (None, "")
        ))
        mapping_fingerprints.setdefault(stable_id, []).append((scope, repr(material)))
    frozen = {key: tuple(sorted(value)) for key, value in mapping_fingerprints.items()}
    return scopes, skills, frozen


def _catalog_scope(row: object) -> str:
    mapping = _plain_mapping(row)
    category_code = str(mapping.get("category_code") or "").strip()
    if category_code not in _APPROVED_CATEGORY_CODES:
        return ""
    class_code = str(mapping.get("class_code") or "").strip()
    class_name = str(mapping.get("class_name") or "").strip()
    if category_code == "14":
        code = str(mapping.get("major_code") or "").strip()
        name = str(mapping.get("major_name") or "").strip()
        if re.fullmatch(r"14\d{4}", code) is None:
            return ""
    else:
        code, name = class_code, class_name
        if re.fullmatch(r"\d{4}", code) is None or code[:2] != category_code:
            return ""
    return " ".join(value for value in (code, name) if value)


def _catalog_changed_scopes(catalog_snapshot: object | None) -> tuple[str, ...]:
    if catalog_snapshot is None:
        return ()
    if isinstance(catalog_snapshot, Mapping):
        explicit = catalog_snapshot.get("affected_scopes") or catalog_snapshot.get("changed_scopes") or ()
        return tuple(sorted({str(value).strip() for value in explicit if str(value).strip()}))
    diff = getattr(catalog_snapshot, "staged_diff", None)
    if diff is None:
        return ()
    scopes: set[str] = set()
    for field in ("added", "removed"):
        for row in getattr(diff, field, ()):
            scope = _catalog_scope(row)
            if scope:
                scopes.add(scope)
    for field in ("renamed", "major_code_changes", "class_moves", "category_moves"):
        for change in getattr(diff, field, ()):
            for row in (getattr(change, "old", None), getattr(change, "new", None)):
                if row is not None:
                    scope = _catalog_scope(row)
                    if scope:
                        scopes.add(scope)
    return tuple(sorted(scopes))


def _catalog_baseline_fingerprint(rows: Iterable[Mapping[str, Any]]) -> tuple[tuple[tuple[str, str], ...], ...]:
    return tuple(sorted(
        tuple(sorted((str(key), str(value)) for key, value in row.items() if key != "访问日期" and value not in (None, "")))
        for row in rows
    ))


def affected_scopes(before: object, after: object, *, catalog_snapshot: object | None = None) -> tuple[str, ...]:
    """Return only scopes whose formal content or catalog boundary materially changed."""

    before_scopes, before_skills, before_maps = _scope_index(before, catalog_snapshot)
    after_scopes, after_skills, after_maps = _scope_index(after, catalog_snapshot)
    affected: set[str] = set()
    for stable_id in sorted(set(before_skills) | set(after_skills)):
        old = before_skills.get(stable_id)
        new = after_skills.get(stable_id)
        old_scopes = before_scopes.get(stable_id, set())
        new_scopes = after_scopes.get(stable_id, set())
        if old is None or new is None:
            affected.update(old_scopes | new_scopes)
            continue
        old_material = tuple(str(old.get(field) or "") for field in _MATERIAL_SKILL_FIELDS)
        new_material = tuple(str(new.get(field) or "") for field in _MATERIAL_SKILL_FIELDS)
        if old_material != new_material or before_maps.get(stable_id, ()) != after_maps.get(stable_id, ()):
            affected.update(old_scopes | new_scopes)
    before_catalog = _ledger_rows(before, "目录基线")
    after_catalog = _ledger_rows(after, "目录基线")
    affected.update(_catalog_changed_scopes(catalog_snapshot))
    if _catalog_baseline_fingerprint(before_catalog) != _catalog_baseline_fingerprint(after_catalog):
        catalog_scopes = {
            scope for row in before_catalog + after_catalog
            if (scope := _approved_scope_name(row, catalog_snapshot))
        }
        affected.update(catalog_scopes or {scope for values in before_scopes.values() for scope in values} | {scope for values in after_scopes.values() for scope in values})
    return tuple(sorted(scope for scope in affected if scope))


def _safe_scope_name(scope: str) -> str:
    value = re.sub(r"[<>:\"/\\|?*\x00-\x1f]", "_", scope).strip(" .")
    if not value or value in {".", ".."}:
        raise ValueError("专业类名称不能安全地用作交付目录")
    return value


def _scope_directory_name(scope: str) -> str:
    return f"{_safe_scope_name(scope)}-{sha256(scope.encode('utf-8')).hexdigest()[:10]}"


def _stable_join(values: Iterable[object]) -> str:
    normalized = {str(value).strip() for value in values if str(value or "").strip()}
    return "；".join(sorted(normalized))


def _ordinary_output_root(output_root: str | Path) -> Path:
    root = Path(output_root).absolute()
    try:
        assert_ordinary_path(root)
        if root.exists():
            assert_ordinary_path(root, require_directory=True)
        else:
            root.mkdir(parents=True, exist_ok=False)
            assert_ordinary_path(root, require_directory=True)
    except (OSError, ValueError) as exc:
        raise ValueError(f"专业类交付根目录必须是普通、无重解析父链的目录：{root}") from exc
    return root


def build_scope_deliveries(
    scopes: Iterable[str],
    ledger: object,
    output_root: str | Path,
) -> tuple[Path, ...]:
    """Build one Word and one Excel copy for each materially affected scope."""

    master_rows = _ledger_rows(ledger, "当前Skill")
    mappings = _ledger_rows(ledger, "专业任务映射")
    by_id = {str(row.get("内部标识") or "").strip(): row for row in master_rows if str(row.get("内部标识") or "").strip()}
    root = _ordinary_output_root(output_root)
    outputs: list[Path] = []
    seen_scopes: set[str] = set()
    for raw_scope in scopes:
        scope = str(raw_scope).strip()
        if not scope or scope in seen_scopes:
            continue
        seen_scopes.add(scope)
        scope_maps = [row for row in mappings if _scope_name(row) == scope]
        ids = sorted({str(row.get("内部标识") or "").strip() for row in scope_maps if str(row.get("内部标识") or "").strip()})
        formal: list[dict[str, Any]] = []
        for stable_id in ids:
            if stable_id not in by_id:
                continue
            row = dict(by_id[stable_id])
            stable_maps = [item for item in scope_maps if str(item.get("内部标识") or "").strip() == stable_id]
            row.update({
                "专业类": scope,
                "用途": _stable_join(item.get("专业任务") for item in stable_maps) or row.get("简要功能"),
                "输入": _stable_join(item.get("输入") for item in stable_maps) or row.get("输入"),
                "输出": _stable_join(item.get("输出") for item in stable_maps) or row.get("输出"),
                "使用限制": _stable_join(item.get("使用限制") for item in stable_maps) or row.get("安全限制条件"),
            })
            formal.append(row)
        payload = {
            "run_id": f"scope-{_scope_directory_name(scope)}",
            "generated_at": datetime.now(),
            "formal_additions": formal,
            "affected_scopes": [scope],
            "source_statuses": {},
        }
        directory = contained_child(root, _scope_directory_name(scope))
        directory.mkdir(parents=False, exist_ok=True)
        assert_ordinary_path(directory, require_directory=True)
        docx_path = contained_child(directory, "专业类Skill清单.docx")
        xlsx_path = contained_child(directory, "专业类Skill清单.xlsx")
        build_daily_docx(payload, docx_path)
        build_daily_xlsx(payload, xlsx_path)
        outputs.extend((docx_path, xlsx_path))
    return tuple(outputs)


def _row_fingerprint(row: Mapping[str, Any]) -> str:
    return json.dumps(_json_value(row), ensure_ascii=False, sort_keys=True, separators=(",", ":"))


def _new_ledger_rows(before: object, after: object, sheet: str) -> list[dict[str, Any]]:
    old = {_row_fingerprint(row) for row in _ledger_rows(before, sheet)}
    return [row for row in _ledger_rows(after, sheet) if _row_fingerprint(row) not in old]


def _catalog_change_rows(catalog_snapshot: object) -> list[dict[str, str]]:
    scopes = _catalog_changed_scopes(catalog_snapshot)
    return [{"专业类": scope, "变化": "目录逐记录差异影响本专业类"} for scope in scopes]


def _normalize_candidate_observation(row: Mapping[str, Any]) -> dict[str, Any]:
    stable_id = str(row.get("观察标识") or "").strip()
    name = str(row.get("候选名称") or "").strip()
    tier = str(row.get("观察状态") or "").strip()
    reason = str(row.get("原因") or "").strip()
    return {
        **row,
        "内部标识": stable_id,
        "Skill名称": name,
        "规范名称": name,
        "入库层级": tier,
        "结论": tier,
        "用途": reason or tier,
        "使用限制": reason or "须按候选层级完成人工复核",
    }


def _report_input_from_run(prepared: object, before: object, after: object) -> dict[str, Any]:
    before_skills = {
        str(row.get("内部标识") or "").strip(): row for row in _ledger_rows(before, "当前Skill")
        if str(row.get("内部标识") or "").strip()
    }
    after_skills = {
        str(row.get("内部标识") or "").strip(): row for row in _ledger_rows(after, "当前Skill")
        if str(row.get("内部标识") or "").strip()
    }
    formal_additions = [after_skills[key] for key in sorted(set(after_skills) - set(before_skills))]
    version_updates = [
        {**after_skills[key], "原版本": before_skills[key].get("固定版本"), "新版本": after_skills[key].get("固定版本")}
        for key in sorted(set(before_skills) & set(after_skills))
        if str(before_skills[key].get("固定版本") or "") != str(after_skills[key].get("固定版本") or "")
    ]
    observations = _new_ledger_rows(before, after, "候选观察")

    def observation_status(row: Mapping[str, Any]) -> str:
        return str(row.get("观察状态") or "").strip()

    conditional = [_normalize_candidate_observation(row) for row in observations if observation_status(row) == "条件候选"]
    adaptation = [_normalize_candidate_observation(row) for row in observations if observation_status(row) == "需适配候选"]
    updates_not_applied = [row for row in observations if observation_status(row) in {"发现更新未升级", "更新未升级"}]
    exclusions = [
        row for row in observations
        if observation_status(row) in {"排除", "已排除", "不推荐", "拒绝"}
    ]
    source_runs = tuple(getattr(prepared, "source_runs", ()))
    source_statuses = {str(run.platform): str(run.status) for run in source_runs}
    source_requests = []
    for run in source_runs:
        for event in tuple(getattr(run, "request_events", ())):
            source_requests.append({
                "来源平台": str(event.platform),
                "请求地址": str(event.url),
                "查询标识": str(event.query_id),
                "页码": event.page,
                "状态码": event.status_code if event.status_code is not None else "未记录",
                "尝试次数": event.attempts,
                "响应SHA-256": str(event.response_sha256 or "未记录"),
                "证据位置": str(event.evidence_path) if event.evidence_path is not None else "未记录",
                "完成": "是" if event.completed else "否",
                "请求时间": "未记录",
            })
    manual_reviews = [
        {"事项": f"{run.platform} 覆盖状态为 {run.status}，需人工复核。"}
        for run in source_runs if str(run.status).casefold() not in {"complete", "success", "ok"}
    ]
    catalog_snapshot = getattr(prepared, "catalog_snapshot", None)
    return {
        "run_id": str(getattr(prepared, "run_id", "未提供")),
        "generated_at": datetime.now(),
        "blocked": False,
        "source_statuses": source_statuses,
        "catalog_changes": _catalog_change_rows(catalog_snapshot),
        "formal_additions": formal_additions,
        "version_updates": version_updates,
        "updates_not_applied": updates_not_applied,
        "conditional_candidates": conditional,
        "adaptation_candidates": adaptation,
        "aliases": _new_ledger_rows(before, after, "来源别名"),
        "affected_scopes": affected_scopes(before, after, catalog_snapshot=catalog_snapshot),
        "exclusions": exclusions,
        "manual_reviews": manual_reviews,
        "source_requests": source_requests,
    }


def _validate_new_formal_mappings(before: object, after: object, catalog_snapshot: object | None) -> None:
    before_ids = {
        str(row.get("内部标识") or "").strip() for row in _ledger_rows(before, "当前Skill")
        if str(row.get("内部标识") or "").strip()
    }
    after_ids = {
        str(row.get("内部标识") or "").strip() for row in _ledger_rows(after, "当前Skill")
        if str(row.get("内部标识") or "").strip()
    }
    mappings = _ledger_rows(after, "专业任务映射")
    for stable_id in sorted(after_ids - before_ids):
        approved = [
            row for row in mappings
            if str(row.get("内部标识") or "").strip() == stable_id
            and _is_approved_scope_mapping(row, catalog_snapshot)
        ]
        if not approved:
            raise ReportBuildError(f"新增正式项 {stable_id} 缺少非军事学、批准范围内的专业任务映射")


def make_project_report_builder(root: str | Path):
    """Bind the runner's pre-publication callback to one project's trusted paths."""

    project = ProjectPaths.from_root(root)

    def build(prepared: object, staging_root: Path) -> tuple[Path, ...]:
        staging = Path(staging_root).absolute()
        prepared_staging = Path(getattr(prepared, "staging_dir")).absolute()
        staged_ledger = Path(getattr(prepared, "staging_ledger")).absolute()
        try:
            if staging != prepared_staging:
                raise ValueError("runner callback 的 staging root 与 PreparedRun 不一致")
            staging.relative_to(project.staging_root.absolute())
            staged_ledger.relative_to(staging)
            assert_ordinary_path(staging, require_directory=True)
            assert_ordinary_path(staged_ledger)
            assert_ordinary_path(project.ledger)
        except (OSError, ValueError) as exc:
            raise ReportBuildError("报告回调只能读取项目台账并写入本轮普通暂存目录") from exc
        before = LedgerStore.load(project.ledger)
        after = LedgerStore.load(staged_ledger)
        try:
            _validate_new_formal_mappings(before, after, getattr(prepared, "catalog_snapshot", None))
            payload = _report_input_from_run(prepared, before, after)
            delivery = contained_child(staging, "deliveries")
            delivery.mkdir(parents=False, exist_ok=False)
            assert_ordinary_path(delivery, require_directory=True)
            daily_docx = contained_child(delivery, "维护日报.docx")
            daily_xlsx = contained_child(delivery, "维护日报.xlsx")
            build_daily_docx(payload, daily_docx)
            build_daily_xlsx(payload, daily_xlsx)
            scope_root = contained_child(delivery, "受影响专业类")
            scope_outputs = build_scope_deliveries(payload["affected_scopes"], after, scope_root)
            return (daily_docx, daily_xlsx, *scope_outputs)
        finally:
            before.workbook.close()
            after.workbook.close()

    return build
