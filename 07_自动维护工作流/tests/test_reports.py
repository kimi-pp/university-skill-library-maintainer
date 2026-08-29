"""Task 10: Chinese daily reports and affected-scope deliveries."""

from __future__ import annotations

from datetime import datetime
import os
from pathlib import Path
import tempfile
from types import SimpleNamespace
import unittest
from unittest.mock import patch

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from openpyxl import load_workbook

from skill_maintainer.catalog import Catalog, CatalogRow, CatalogSnapshot, diff_catalog
from skill_maintainer.ledger import LedgerStore
from skill_maintainer.reports import (
    DAILY_SHEETS,
    DAILY_WORD_SECTIONS,
    affected_scopes,
    build_daily_docx,
    build_daily_xlsx,
    build_scope_deliveries,
    ReportBuildError,
    _make_node_modules_link,
    _remove_node_modules_link,
    _report_input_from_run,
)
from skill_maintainer.runner import SourceRun


def formal_row(index: int, scope: str = "0809 计算机类") -> dict[str, object]:
    return {
        "内部标识": f"GH-05-{index:04d}",
        "Skill名称": f"campus-data-{index}",
        "规范名称": f"Campus Data Skill {index}",
        "固定版本": f"v1.{index}",
        "许可证": "Apache-2.0",
        "Canonical source": f"https://github.com/example/campus-data-{index}",
        "简要功能": "面向高校数据治理任务提供结构化分析。",
        "详细功能摘要": "读取课程数据并输出可审计的中文分析结果。",
        "适用用户角色": "教师、科研人员和数据管理员",
        "典型高校场景": "课程质量分析",
        "外部依赖": "Python",
        "安全限制条件": "仅处理已授权数据，不上传敏感教学数据。",
        "适配建议": "先在隔离环境复核输入字段。",
        "专业类": scope,
        "用途": "课程数据质量检查",
        "输入": "脱敏后的课程数据表",
        "输出": "中文质量报告和问题清单",
        "使用限制": "不得输入学生个人敏感信息",
        "专业任务": "课程数据治理",
        "收集日期": datetime(2026, 8, 28),
    }


def approved_scope_catalog() -> Catalog:
    """A captured catalog whose exact codes define the report scope boundary."""

    return Catalog((
        CatalogRow("08", "工学", "0801", "力学类", "080101", "理论与应用力学"),
        CatalogRow("08", "工学", "0809", "计算机类", "080901", "计算机科学与技术"),
        CatalogRow("11", "军事学", "1101", "军事类", "110101", "军事专业"),
        CatalogRow("14", "交叉学科", None, None, "140101", "集成电路科学与工程"),
    ))


def report_summary(count: int = 2) -> dict[str, object]:
    rows = [formal_row(index) for index in range(1, count + 1)]
    return {
        "run_id": "run-20260828-220000",
        "generated_at": datetime(2026, 8, 28, 22, 0),
        "source_statuses": {
            "SkillHub": "complete",
            "ClawHub": "partial",
            "GitHub": "complete",
            "Hugging Face Spaces": "failed",
        },
        "catalog_changes": [{"专业类": "0809 计算机类", "变化": "专业任务范围已复核"}],
        "formal_additions": rows,
        "version_updates": [{**rows[0], "原版本": "v1.0", "新版本": "v1.1"}],
        "updates_not_applied": [{**rows[0], "发现版本": "v2.0", "原因": "证据不足"}],
        "conditional_candidates": [{**rows[0], "内部标识": "COND-0001", "结论": "条件候选"}],
        "adaptation_candidates": [{**rows[0], "内部标识": "ADAPT-0001", "结论": "需适配候选"}],
        "aliases": [{"内部标识": rows[0]["内部标识"], "来源平台": "SkillHub", "来源地址": "https://skillhub.example/item/1"}],
        "affected_scopes": ["0809 计算机类"],
        "exclusions": [{"候选名称": "绝不能泄露的落选项目名", "原因": "许可证不明确"}],
        "manual_reviews": [{"事项": "ClawHub 覆盖降级，需人工复核"}],
        "source_requests": [{"来源平台": "GitHub", "请求地址": "https://api.github.com/search/repositories?q=campus", "状态": "200", "请求时间": datetime(2026, 8, 28, 21, 50)}],
    }


class ReportContentTestCase(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory()
        self.addCleanup(self.temporary.cleanup)
        self.root = Path(self.temporary.name)

    def require_runtime(self) -> tuple[Path, Path]:
        node = os.environ.get("SKILL_MAINTAINER_NODE")
        modules = os.environ.get("SKILL_MAINTAINER_NODE_MODULES")
        if not node or not modules:
            self.skipTest("report integration requires caller-supplied SKILL_MAINTAINER_NODE and SKILL_MAINTAINER_NODE_MODULES")
        return Path(node), Path(modules)

    def test_word_contains_fixed_sections_chinese_fields_and_non_execution_boundary(self):
        output = self.root / "日报.docx"
        build_daily_docx(report_summary(), output)
        document = Document(output)
        headings = [p.text for p in document.paragraphs if p.style.name == "Heading 1"]
        self.assertEqual(headings, list(DAILY_WORD_SECTIONS))
        text = "\n".join(p.text for p in document.paragraphs)
        text += "\n" + "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        for required in (
            "用途", "适用人员", "输入", "输出", "限制", "Campus Data Skill 1",
            "Apache-2.0", "https://github.com/example/campus-data-1", "未安装、未运行",
        ):
            self.assertIn(required, text)
        self.assertNotIn("绝不能泄露的落选项目名", text)
        section = document.sections[0]
        self.assertEqual((section.page_width, section.page_height), (Inches(8.5), Inches(11)))
        self.assertEqual(
            (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin),
            (Inches(1), Inches(1), Inches(1), Inches(1)),
        )
        normal = document.styles["Normal"]
        self.assertEqual((normal.font.name, normal.font.size), ("Calibri", Pt(11)))
        self.assertEqual((normal.paragraph_format.space_after, normal.paragraph_format.line_spacing), (Pt(6), 1.1))
        expected_styles = {
            "Heading 1": (Pt(16), "2E74B5", Pt(16), Pt(8)),
            "Heading 2": (Pt(13), "2E74B5", Pt(12), Pt(6)),
            "Heading 3": (Pt(12), "1F4D78", Pt(8), Pt(4)),
        }
        for name, expected in expected_styles.items():
            style = document.styles[name]
            actual = (style.font.size, str(style.font.color.rgb), style.paragraph_format.space_before, style.paragraph_format.space_after)
            self.assertEqual(actual, expected)
        for paragraph in (p for p in document.paragraphs if p.style.name == "Heading 1"):
            self.assertIsNotNone(paragraph._p.pPr.find(qn("w:numPr")))
        self.assertNotIn("w:pBdr", section.header._element.xml)
        for table in document.tables:
            self.assertIn('w:w="9360"', table._tbl.tblPr.xml)
            self.assertIn('w:w="120"', table._tbl.tblPr.xml)

    def test_excel_has_exact_sheets_dynamic_520_row_tables_and_operational_features(self):
        self.require_runtime()
        output = self.root / "日报.xlsx"
        summary = report_summary(520)
        build_daily_xlsx(summary, output)
        workbook = load_workbook(output, data_only=False)
        self.addCleanup(workbook.close)
        self.assertEqual(workbook.sheetnames, list(DAILY_SHEETS))
        for sheet in workbook:
            self.assertEqual(sheet.freeze_panes, "A2", sheet.title)
            self.assertEqual(len(sheet.tables), 1, sheet.title)
        formal = workbook["新增正式推荐"]
        self.assertEqual(formal.max_row, 521)
        self.assertEqual(formal["A521"].value, "GH-05-0520")
        self.assertEqual(len({formal.cell(row, 1).value for row in range(2, 522)}), 520)
        self.assertEqual(formal.freeze_panes, "A2")
        self.assertEqual(len(formal.tables), 1)
        self.assertTrue(next(iter(formal.tables.values())).ref.endswith("521"))
        self.assertEqual(formal["G2"].value, "https://github.com/example/campus-data-1")
        self.assertEqual(formal["G2"].hyperlink.target, "https://github.com/example/campus-data-1")
        self.assertTrue(formal["H2"].alignment.wrap_text)
        self.assertEqual(formal["N2"].number_format, "yyyy-mm-dd")
        overview = workbook["执行概览"]
        self.assertEqual(overview["B3"].value, datetime(2026, 8, 28, 22, 0))
        formulas = [cell.value for row in overview.iter_rows() for cell in row if isinstance(cell.value, str) and cell.value.startswith("=")]
        self.assertTrue(any("521" in formula for formula in formulas), formulas)
        self.assertEqual(overview["B5"].value, 520)
        self.assertNotIn("绝不能泄露的落选项目名", "\n".join(str(cell.value or "") for sheet in workbook for row in sheet.iter_rows() for cell in row))

    def test_exclusion_free_text_is_normalized_without_leaking_candidate_name(self):
        self.require_runtime()
        summary = report_summary()
        summary["exclusions"] = [{"原因": "候选项目 secret-candidate 因维护者失联而排除"}]
        docx = self.root / "脱敏.docx"
        xlsx = self.root / "脱敏.xlsx"
        build_daily_docx(summary, docx)
        build_daily_xlsx(summary, xlsx)
        document = Document(docx)
        word_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        workbook = load_workbook(xlsx, data_only=False)
        self.addCleanup(workbook.close)
        excel_text = "\n".join(str(cell.value or "") for sheet in workbook for row in sheet.iter_rows() for cell in row)
        for text in (word_text, excel_text):
            self.assertNotIn("secret-candidate", text)
            self.assertNotIn("维护者失联", text)
            self.assertIn("其他合规原因", text)

    def test_affected_scope_refresh_is_material_and_alias_only_is_not(self):
        base_skill = formal_row(1, "0801 力学类")
        before = {
            "当前Skill": [base_skill],
            "专业任务映射": [{
                "内部标识": base_skill["内部标识"], "专业代码": "0801",
                "专业名称": "力学类", "专业任务": "建模",
            }],
            "来源别名": [],
            "目录基线": [{"目录版本": "2025", "专业类": "0801 力学类"}],
        }
        alias_after = {**before, "来源别名": [{"内部标识": base_skill["内部标识"], "来源地址": "https://skillhub.example/alias"}]}
        catalog = approved_scope_catalog()
        self.assertEqual(affected_scopes(before, alias_after, catalog_snapshot=catalog), ())

        cases = []
        added = formal_row(2, "0809 计算机类")
        cases.append({**before, "当前Skill": [base_skill, added], "专业任务映射": before["专业任务映射"] + [{
            "内部标识": added["内部标识"], "专业代码": "0809",
            "专业名称": "计算机类", "专业任务": "数据分析",
        }]})
        cases.append({**before, "当前Skill": [{**base_skill, "固定版本": "v2.0"}]})
        cases.append({**before, "当前Skill": [{**base_skill, "许可证": "BSD-3-Clause"}]})
        cases.append({**before, "当前Skill": [{**base_skill, "安全等级": "SB"}]})
        cases.append({**before, "专业任务映射": [{**before["专业任务映射"][0], "专业任务": "实验建模"}]})
        cases.append({**before, "目录基线": [{"目录版本": "2026", "专业类": "0801 力学类"}]})
        for after in cases:
            with self.subTest(after=after):
                self.assertTrue(affected_scopes(before, after, catalog_snapshot=catalog))

    def test_affected_scope_includes_formal_validation_and_risk_conclusions(self):
        skill = formal_row(1, "0801 力学类")
        skill.update({"验证状态": "全部通过（未实测）", "风险提示": "低风险"})
        before = {
            "当前Skill": [skill],
            "专业任务映射": [{
                "内部标识": skill["内部标识"], "专业代码": "0801",
                "专业名称": "力学类", "专业任务": "建模",
            }],
            "目录基线": [],
        }
        for field, value in (("验证状态", "部分通过"), ("风险提示", "需隔离运行")):
            after = {**before, "当前Skill": [{**skill, field: value}]}
            with self.subTest(field=field):
                self.assertEqual(
                    affected_scopes(before, after, catalog_snapshot=approved_scope_catalog()),
                    ("0801 力学类",),
                )

    def test_catalog_access_date_only_does_not_refresh_every_scope(self):
        skill = formal_row(1, "0801 力学类")
        before = {
            "当前Skill": [skill], "专业任务映射": [],
            "目录基线": [{"目录版本": "2025", "SHA-256": "a" * 64, "访问日期": "2026-08-27"}],
        }
        after = {**before, "目录基线": [{**before["目录基线"][0], "访问日期": "2026-08-28"}]}
        self.assertEqual(affected_scopes(before, after), ())

    def test_catalog_snapshot_refreshes_only_exact_changed_professional_class(self):
        old = (
            CatalogRow("08", "工学", "0801", "力学类", "080101", "理论与应用力学"),
            CatalogRow("08", "工学", "0809", "计算机类", "080901", "计算机科学与技术"),
        )
        new = (*old, CatalogRow("08", "工学", "0809", "计算机类", "080902", "软件工程"))
        catalog = Catalog(old, staged_snapshot=CatalogSnapshot(new, "b" * 64), staged_diff=diff_catalog(old, new))
        ledger = {
            "当前Skill": [formal_row(1, "0801 力学类"), formal_row(2, "0809 计算机类")],
            "专业任务映射": [], "目录基线": [],
        }
        self.assertEqual(affected_scopes(ledger, ledger, catalog_snapshot=catalog), ("0809 计算机类",))

    def test_catalog_snapshot_excludes_military_and_unapproved_categories(self):
        rows = (
            CatalogRow("08", "工学", "0809", "计算机类", "080901", "计算机科学与技术"),
            CatalogRow("08", "工学", "08evil", "伪造类", "080999", "伪造专业"),
            CatalogRow("11", "军事学", "1101", "军事类", "110101", "军事专业"),
            CatalogRow("14", "交叉学科", None, None, "1401", "伪交叉学科专业"),
            CatalogRow("15", "未批准门类", "1501", "未批准类", "150101", "未批准专业"),
        )
        catalog = Catalog((), staged_snapshot=CatalogSnapshot(rows, "c" * 64), staged_diff=diff_catalog((), rows))
        ledger = {"当前Skill": [], "专业任务映射": [], "目录基线": []}
        self.assertEqual(affected_scopes(ledger, ledger, catalog_snapshot=catalog), ("0809 计算机类",))

    def test_new_formal_with_mapping_refreshes_only_its_exact_scope(self):
        before = {"当前Skill": [], "专业任务映射": [], "目录基线": []}
        skill = formal_row(1, "0809 计算机类")
        after = {
            "当前Skill": [skill],
            "专业任务映射": [{"内部标识": skill["内部标识"], "专业代码": "0809", "专业名称": "计算机类", "专业任务": "课程分析"}],
            "目录基线": [],
        }
        self.assertEqual(
            affected_scopes(before, after, catalog_snapshot=approved_scope_catalog()),
            ("0809 计算机类",),
        )

    def test_candidate_only_change_refreshes_every_human_mapped_scope_without_name_keyword(self):
        candidate = {
            "观察标识": "OBS-SK-SHARED-条件候选", "内部标识": "SK-SHARED",
            "候选名称": "generic-table-helper", "Canonical source": "https://example.test/shared",
            "观察状态": "条件候选", "许可证": "MIT", "原因": "需人工复核输入",
            "固定版本": "a" * 40, "固定版本内容指纹": "b" * 64,
            "验证证据位置": "evidence/shared/SKILL.md", "显示层级": "条件候选",
        }
        mappings = [
            {"映射标识": "MAP-SK-SHARED-0801", "内部标识": "SK-SHARED", "专业代码": "0801", "专业名称": "力学类", "专业任务": "整理实验表", "输入": "测量表", "输出": "核验表", "适用理由": "支持实验数据准备", "使用限制": "人工复核", "相关度": 4},
            {"映射标识": "MAP-SK-SHARED-0809", "内部标识": "SK-SHARED", "专业代码": "0809", "专业名称": "计算机类", "专业任务": "清洗评测表", "输入": "评测表", "输出": "核验表", "适用理由": "支持软件评测", "使用限制": "人工复核", "相关度": 4},
        ]
        before = {"当前Skill": [], "候选观察": [candidate], "专业任务映射": mappings, "目录基线": []}
        after = {**before, "候选观察": [{**candidate, "固定版本": "c" * 40, "原因": "新版本仍需人工复核"}]}
        self.assertEqual(
            affected_scopes(before, after, catalog_snapshot=approved_scope_catalog()),
            ("0801 力学类", "0809 计算机类"),
        )

        self.require_runtime()
        paths = build_scope_deliveries(("0801 力学类", "0809 计算机类"), after, self.root / "候选跨专业")
        self.assertEqual(len(paths), 4)
        for workbook_path in (path for path in paths if path.suffix == ".xlsx"):
            workbook = load_workbook(workbook_path, data_only=False)
            try:
                self.assertEqual(workbook["条件候选"]["A2"].value, "SK-SHARED")
                for sheet_name in ("新增正式推荐", "需适配候选"):
                    self.assertNotIn(
                        "SK-SHARED",
                        [cell.value for row in workbook[sheet_name].iter_rows() for cell in row],
                    )
            finally:
                workbook.close()

    def test_scope_mappings_must_match_exact_codes_in_the_captured_catalog(self):
        catalog = approved_scope_catalog()
        cases = (
            ("malformed-prefix", "08evil", "伪造类", ()),
            ("fake-interdisciplinary-class", "1401", "伪交叉学科类", ()),
            ("empty-code-military-text", "", "军事学自由文本", ()),
            ("empty-code-leading-valid-ordinary-text", "", "0809 任意自由文本", ()),
            ("empty-code-leading-valid-military-text", "", "0809 军事自由文本", ()),
            ("well-formed-but-absent", "0808", "自动化类", ()),
            ("military-code", "1101", "军事类", ()),
            ("real-class", "0809", "计算机类", ("0809 计算机类",)),
            ("real-interdisciplinary-major", "140101", "集成电路科学与工程", ("140101 集成电路科学与工程",)),
            ("generic", "99", "跨学科通用", ("99 跨学科通用",)),
        )
        for index, (label, code, name, expected) in enumerate(cases, start=1):
            stable_id = f"SCOPE-{index}"
            before = {"当前Skill": [], "专业任务映射": [], "目录基线": []}
            after = {
                "当前Skill": [{**formal_row(index, ""), "内部标识": stable_id}],
                "专业任务映射": [{
                    "内部标识": stable_id,
                    "专业代码": code,
                    "专业名称": name,
                    "专业类": name if not code else "",
                    "专业任务": "课程分析",
                }],
                "目录基线": [],
            }
            with self.subTest(label=label):
                self.assertEqual(affected_scopes(before, after, catalog_snapshot=catalog), expected)

    def test_real_candidate_observations_are_normalized_for_both_candidate_sheets(self):
        self.require_runtime()
        before_path = self.root / "before.xlsx"
        after_path = self.root / "after.xlsx"
        before = LedgerStore.create(before_path)
        after = LedgerStore.create(self.root / "after-source.xlsx")
        after.append_rows("候选观察", [
            {"观察标识": "OBS-COND-1", "内部标识": "SK-COND-1", "候选名称": "Conditional Candidate", "Canonical source": "https://example.test/conditional", "观察状态": "条件候选", "许可证": "MIT", "记录日期": datetime(2026, 8, 28), "原因": "仅限脱敏数据"},
            {"观察标识": "OBS-ADAPT-1", "内部标识": "SK-ADAPT-1", "候选名称": "Adaptation Candidate", "Canonical source": "https://example.test/adaptation", "观察状态": "需适配候选", "许可证": "Apache-2.0", "记录日期": datetime(2026, 8, 28), "原因": "需本地适配"},
        ])
        after.save_staged(after_path)
        prepared = SimpleNamespace(
            run_id="candidate-report", catalog_snapshot=None,
            source_runs=(SourceRun("SkillHub", "partial", query="not-a-url-query"),),
        )
        try:
            payload = _report_input_from_run(prepared, before, after)
            output = self.root / "candidate-report.xlsx"
            word_output = self.root / "candidate-report.docx"
            build_daily_docx(payload, word_output)
            build_daily_xlsx(payload, output)
            workbook = load_workbook(output, data_only=False)
            self.addCleanup(workbook.close)
            for sheet, stable_id, name, reason in (
                ("条件候选", "SK-COND-1", "Conditional Candidate", "仅限脱敏数据"),
                ("需适配候选", "SK-ADAPT-1", "Adaptation Candidate", "需本地适配"),
            ):
                self.assertEqual(workbook[sheet]["A2"].value, stable_id)
                self.assertEqual(workbook[sheet]["B2"].value, name)
                self.assertEqual(workbook[sheet]["L2"].value, reason)
                self.assertEqual(workbook[sheet]["O2"].value, sheet)
                self.assertEqual(workbook[sheet]["P2"].value, reason)
                self.assertIsNotNone(workbook[sheet]["G2"].hyperlink)
            document = Document(word_output)
            word_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
            self.assertIn("层级", word_text)
            self.assertIn("原因/结论", word_text)
        finally:
            before.workbook.close()
            after.workbook.close()

    def test_source_query_without_request_events_creates_no_audit_row(self):
        prepared = SimpleNamespace(
            run_id="no-events", catalog_snapshot=None,
            source_runs=(SourceRun("GitHub", "complete", query="campus skill query"),),
        )
        empty = {"当前Skill": [], "专业任务映射": [], "候选观察": [], "来源别名": [], "目录基线": []}
        payload = _report_input_from_run(prepared, empty, empty)
        self.assertEqual(payload["source_requests"], [])

    def test_existing_formal_candidate_observation_is_only_an_unapplied_update(self):
        current = formal_row(1, "0809 计算机类")
        observation = {
            "观察标识": "review-v2", "内部标识": current["内部标识"], "候选名称": current["Skill名称"],
            "Canonical source": current["Canonical source"], "Skill入口路径": "SKILL.md",
            "观察状态": "条件候选", "许可证": "MIT", "记录日期": "2026-08-29",
            "原因": "新版本许可证仍待确认", "固定版本": "b" * 40,
            "固定版本内容指纹": "c" * 64, "验证证据位置": "evidence/new-version.json",
            "显示层级": "条件候选",
        }
        before = {"当前Skill": [current], "候选观察": [], "来源别名": [], "专业任务映射": [], "目录基线": []}
        after = {**before, "候选观察": [observation]}
        prepared = SimpleNamespace(run_id="update-report", catalog_snapshot=None, source_runs=())

        payload = _report_input_from_run(prepared, before, after)

        self.assertEqual(payload["conditional_candidates"], [])
        self.assertEqual(payload["adaptation_candidates"], [])
        self.assertEqual(len(payload["updates_not_applied"]), 1)
        self.assertEqual(payload["updates_not_applied"][0]["新版本"], "b" * 40)
        self.assertIn("许可证", payload["updates_not_applied"][0]["结论"])

    def test_upstream_deletion_is_an_unapplied_update_with_old_version_retained(self):
        current = formal_row(1, "0809 计算机类")
        attention = {
            "观察标识": "deleted-v1", "内部标识": current["内部标识"], "候选名称": current["Skill名称"],
            "Canonical source": current["Canonical source"], "Skill入口路径": "SKILL.md",
            "观察状态": "attention_required", "许可证": "MIT", "记录日期": "2026-08-29",
            "原因": "上游入口删除；旧版本保留", "固定版本": current["固定版本"],
            "固定版本内容指纹": "d" * 64, "验证证据位置": "evidence/deleted.json",
            "原因代码": "upstream-entry-deleted", "显示层级": "不展示",
        }
        mapping = {"内部标识": current["内部标识"], "专业代码": "0809", "专业名称": "计算机类", "专业任务": "课程分析"}
        before = {"当前Skill": [current], "候选观察": [], "来源别名": [], "专业任务映射": [mapping], "目录基线": []}
        after = {**before, "候选观察": [attention]}
        prepared = SimpleNamespace(run_id="delete-report", catalog_snapshot=None, source_runs=())

        payload = _report_input_from_run(prepared, before, after)

        self.assertEqual(len(payload["updates_not_applied"]), 1)
        self.assertIn("上游入口删除", payload["updates_not_applied"][0]["结论"])
        self.assertIn("旧版本保留", payload["updates_not_applied"][0]["使用限制"])

    def test_scope_delivery_surfaces_attention_alongside_retained_formal_row(self):
        current = formal_row(1, "0809 计算机类")
        attention = {
            "观察标识": "deleted-scope", "内部标识": current["内部标识"], "候选名称": current["Skill名称"],
            "Canonical source": current["Canonical source"], "Skill入口路径": "SKILL.md",
            "观察状态": "attention_required", "许可证": "MIT", "记录日期": "2026-08-29",
            "原因": "上游入口删除；旧版本保留", "固定版本": current["固定版本"],
            "固定版本内容指纹": "d" * 64, "验证证据位置": "evidence/deleted.json",
            "原因代码": "upstream-entry-deleted", "显示层级": "不展示",
        }
        ledger = {
            "当前Skill": [current], "候选观察": [attention],
            "专业任务映射": [{
                "内部标识": current["内部标识"], "专业代码": "0809", "专业名称": "计算机类",
                "专业任务": "课程分析", "输入": "课程表", "输出": "报告", "使用限制": "人工复核",
            }],
        }
        captured = []
        with patch("skill_maintainer.reports.build_daily_docx", side_effect=lambda payload, path: captured.append(payload)), \
             patch("skill_maintainer.reports.build_daily_xlsx", side_effect=lambda payload, path: None):
            build_scope_deliveries(("0809 计算机类",), ledger, self.root / "attention-scope")

        self.assertEqual(len(captured[0]["formal_additions"]), 1)
        self.assertEqual(len(captured[0]["updates_not_applied"]), 1)
        self.assertIn("上游入口删除", captured[0]["updates_not_applied"][0]["结论"])

    def test_source_audit_hyperlinks_only_http_or_https_urls(self):
        self.require_runtime()
        summary = report_summary()
        summary["source_requests"] = [
            {"来源平台": "GitHub", "请求地址": "campus skill query", "状态": "未记录"},
            {"来源平台": "SkillHub", "请求地址": "https://api.example.test/items", "状态": 200},
        ]
        output = self.root / "audit-links.xlsx"
        build_daily_xlsx(summary, output)
        workbook = load_workbook(output, data_only=False)
        self.addCleanup(workbook.close)
        audit = workbook["来源请求审计"]
        self.assertIsNone(audit["B2"].hyperlink)
        self.assertEqual(audit["B3"].hyperlink.target, "https://api.example.test/items")

    def test_scope_deliveries_reference_same_stable_ids_without_master_duplication(self):
        self.require_runtime()
        shared = formal_row(1, "0809 计算机类")
        ledger = {
            "当前Skill": [shared],
            "专业任务映射": [
                {"内部标识": shared["内部标识"], "专业类": "0809 计算机类", "专业任务": "课程分析", "输入": "课程表", "输出": "报告", "使用限制": "脱敏"},
                {"内部标识": shared["内部标识"], "专业类": "0201 经济学类", "专业任务": "计量分析", "输入": "统计表", "输出": "报告", "使用限制": "脱敏"},
            ],
        }
        paths = build_scope_deliveries(("0809 计算机类", "0201 经济学类"), ledger, self.root / "受影响专业类")
        self.assertEqual(len(paths), 4)
        for path in paths:
            self.assertTrue(path.exists())
        for docx_path in (path for path in paths if path.suffix == ".docx"):
            document = Document(docx_path)
            text = "\n".join(p.text for p in document.paragraphs)
            text += "\n" + "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
            self.assertEqual(text.count("GH-05-0001"), 1)

    def test_scope_delivery_aggregates_multiple_task_mappings_for_one_stable_id(self):
        self.require_runtime()
        scope = "0809 计算机类"
        shared = formal_row(1, scope)
        ledger = {
            "当前Skill": [shared],
            "专业任务映射": [
                {"映射标识": "M-2", "内部标识": shared["内部标识"], "专业类": scope, "专业任务": "实验排课", "输入": "实验安排", "输出": "冲突清单", "使用限制": "教师复核"},
                {"映射标识": "M-1", "内部标识": shared["内部标识"], "专业类": scope, "专业任务": "课程分析", "输入": "课程表", "输出": "质量报告", "使用限制": "数据脱敏"},
            ],
        }
        paths = build_scope_deliveries((scope,), ledger, self.root / "聚合交付")
        document = Document(next(path for path in paths if path.suffix == ".docx"))
        text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        self.assertEqual(text.count("GH-05-0001"), 1)
        for expected in ("实验排课", "课程分析", "实验安排", "课程表", "冲突清单", "质量报告", "教师复核", "数据脱敏"):
            self.assertIn(expected, text)

    def test_scope_names_that_sanitize_alike_have_distinct_deterministic_directories(self):
        self.require_runtime()
        ledger = {
            "当前Skill": [formal_row(1, "A/B"), formal_row(2, "A:B")],
            "专业任务映射": [
                {"内部标识": "GH-05-0001", "专业类": "A/B", "专业任务": "甲"},
                {"内部标识": "GH-05-0002", "专业类": "A:B", "专业任务": "乙"},
            ],
        }
        first = build_scope_deliveries(("A/B", "A:B"), ledger, self.root / "碰撞")
        self.assertEqual(len({path.parent.name for path in first}), 2)
        second = build_scope_deliveries(("A:B", "A/B"), ledger, self.root / "碰撞复验")
        self.assertEqual({path.parent.name for path in first}, {path.parent.name for path in second})

    def test_scope_delivery_rejects_reparse_output_root_without_outside_side_effects(self):
        outside = self.root / "outside"
        outside.mkdir()
        linked = self.root / "linked-output"
        try:
            linked.symlink_to(outside, target_is_directory=True)
        except (NotImplementedError, OSError) as exc:
            self.skipTest(f"test filesystem cannot create a directory link: {exc}")
        with self.assertRaises(ValueError):
            build_scope_deliveries(("0809 计算机类",), {"当前Skill": [], "专业任务映射": []}, linked)
        self.assertEqual(list(outside.iterdir()), [])

    def test_runtime_paths_reject_missing_reparse_and_wrong_node_identity(self):
        real_node, real_modules = self.require_runtime()
        with patch.dict(os.environ, {}, clear=False):
            os.environ.pop("SKILL_MAINTAINER_NODE", None)
            with self.assertRaisesRegex(ReportBuildError, "SKILL_MAINTAINER_NODE"):
                build_daily_xlsx(report_summary(), self.root / "missing.xlsx")

        fake = self.root / "node.exe"
        fake.write_text("not node", encoding="utf-8")
        with patch.dict(os.environ, {"SKILL_MAINTAINER_NODE": str(fake), "SKILL_MAINTAINER_NODE_MODULES": str(real_modules)}):
            with self.assertRaisesRegex(ReportBuildError, "Node|node"):
                build_daily_xlsx(report_summary(), self.root / "wrong.xlsx")

        node_link = self.root / "linked-node.exe"
        modules_link = self.root / "linked-modules"
        try:
            node_link.symlink_to(real_node)
            modules_link.symlink_to(real_modules, target_is_directory=True)
        except (NotImplementedError, OSError) as exc:
            self.skipTest(f"test filesystem cannot create runtime links: {exc}")
        for node, modules in ((node_link, real_modules), (real_node, modules_link)):
            with self.subTest(node=node, modules=modules), patch.dict(os.environ, {"SKILL_MAINTAINER_NODE": str(node), "SKILL_MAINTAINER_NODE_MODULES": str(modules)}):
                with self.assertRaisesRegex(ReportBuildError, "链接|重解析点"):
                    build_daily_xlsx(report_summary(), self.root / "linked.xlsx")

    def test_node_modules_link_handles_ampersand_path_without_shell_parsing(self):
        target = self.root / "modules & trusted"
        target.mkdir()
        link = self.root / "node_modules"
        try:
            _make_node_modules_link(link, target)
        except (NotImplementedError, OSError) as exc:
            self.skipTest(f"test filesystem cannot create a directory link: {exc}")
        self.addCleanup(lambda: _remove_node_modules_link(link))
        self.assertTrue(link.is_dir())
        marker = target / "proof.txt"
        marker.write_text("safe", encoding="utf-8")
        self.assertEqual((link / "proof.txt").read_text(encoding="utf-8"), "safe")

    def test_committed_xlsx_template_reopens_with_all_required_sheets(self):
        template = Path(__file__).parents[1] / "templates" / "daily_review.xlsx"
        workbook = load_workbook(template, data_only=False)
        self.addCleanup(workbook.close)
        self.assertEqual(workbook.sheetnames, list(DAILY_SHEETS))
        self.assertEqual(workbook["使用说明"]["A2"].value, "报告用途")
        self.assertEqual(workbook["条件候选"]["O1"].value, "层级")
        self.assertEqual(workbook["条件候选"]["P1"].value, "原因/结论")
        self.assertEqual(
            [workbook["来源请求审计"].cell(1, column).value for column in range(1, 11)],
            ["来源平台", "请求地址", "查询标识", "页码", "状态码", "尝试次数", "响应SHA-256", "证据位置", "完成", "请求时间"],
        )


if __name__ == "__main__":
    unittest.main()
