"""Task 11: real Office verification and single-authority publication."""

from __future__ import annotations

import copy
from concurrent.futures import ThreadPoolExecutor
from dataclasses import replace
from hashlib import sha256
import json
import os
from pathlib import Path
import sys
import tempfile
import threading
import time
import unittest
from unittest.mock import patch
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from openpyxl import Workbook, load_workbook
from PIL import Image

from skill_maintainer.office import (
    OfficeCheck,
    OfficeEvidenceBundle,
    OfficeVerificationError,
    WordRenderDecision,
    bind_word_visual_decision,
    clear_office_run_state,
    verify_excel,
    verify_word,
)
from skill_maintainer.publish import (
    PublishFile,
    PublishError,
    build_publish_plan,
    commit_prepared_generation,
    publish_atomically,
)


def file_sha256(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def write_excel(path: Path, *, rows: int) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "当前Skill"
    sheet.append(["内部标识", "Skill名称"])
    for index in range(1, rows + 1):
        sheet.append([f"GH-01-{index:04d}", f"skill-{index}"])
    workbook.save(path)
    workbook.close()


def write_corrupt_xlsx(path: Path) -> None:
    """A ZIP/OOXML-shaped input whose workbook part cannot be parsed by Excel."""
    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Override PartName="/xl/workbook.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/></Types>',
        )
        archive.writestr("xl/workbook.xml", "<workbook><broken>")


def write_word(path: Path) -> None:
    document = Document()
    document.add_heading("高校专业 Skill 自动查验", level=1)
    document.add_paragraph("这是一份用于 Office 只读复读和逐页渲染验收的中文测试报告。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "结果"
    table.cell(1, 0).text = "Office 复读"
    table.cell(1, 1).text = "待验证"
    document.save(path)


def trusted_excel_check(path: Path) -> OfficeCheck:
    """Deterministic unit evidence; the real COM path is exercised separately below."""
    return OfficeCheck(
        kind="excel", source_path=path.absolute(), source_sha256=file_sha256(path),
        passed=True, office_passed=True, office_opened=True, read_only=True,
        key_sheet="当前Skill", last_row=2, last_column=2, last_value="ok",
        process_count_before=0, process_count_after=0,
    )


def write_renderer_command(root: Path, *, body_pixels: int | None = None):
    """Return an explicit renderer command that is safe for Chinese/space paths."""
    from skill_maintainer.office import RendererCommand

    script = root / "显式 PDF 渲染器.py"
    pixel_statement = (
        "pixels=0 if ImageChops.difference(body,white).getbbox() is None else 1"
        if body_pixels is None else f"pixels={body_pixels}"
    )
    script.write_text(
        "import argparse, json\n"
        "from pathlib import Path\n"
        "from pdf2image import convert_from_path\n"
        "from PIL import ImageChops\n"
        "p=argparse.ArgumentParser(); p.add_argument('--pdf'); p.add_argument('--output-dir'); a=p.parse_args()\n"
        "out=Path(a.output_dir); pages=[]\n"
        "for i,image in enumerate(convert_from_path(a.pdf,dpi=110,fmt='png',thread_count=1),1):\n"
        " target=out/f'page-{i}.png'; image.save(target,'PNG'); rgb=image.convert('RGB')\n"
        " top=max(1,int(rgb.height*0.10)); bottom=min(rgb.height,int(rgb.height*0.90))\n"
        " body=rgb.crop((0,top,rgb.width,bottom)); white=body.copy(); white.paste((255,255,255),(0,0,white.width,white.height))\n"
        f" {pixel_statement}\n"
        " pages.append({'path':target.name,'body_nonwhite_pixels':pixels})\n"
        "print(json.dumps({'pages':pages},ensure_ascii=False,separators=(',',':')))\n",
        encoding="utf-8",
    )
    return RendererCommand((sys.executable, str(script)))


@unittest.skipUnless(os.name == "nt", "Microsoft Office COM acceptance requires Windows")
class OfficeVerificationTestCase(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory()
        self.addCleanup(self.temporary.cleanup)
        self.addCleanup(clear_office_run_state)
        self.root = Path(self.temporary.name)

    def test_excel_opens_read_only_reopens_and_keeps_520th_data_row(self):
        workbook_path = self.root / "ledger-520.xlsx"
        write_excel(workbook_path, rows=520)

        first = verify_excel(workbook_path)
        second = verify_excel(workbook_path)

        for check in (first, second):
            self.assertTrue(check.passed, check.error)
            self.assertTrue(check.office_opened)
            self.assertTrue(check.read_only)
            self.assertEqual(check.key_sheet, "当前Skill")
            self.assertEqual(check.last_row, 521)
            self.assertEqual(check.last_value, "skill-520")
            self.assertEqual(check.process_count_before, check.process_count_after)
            self.assertEqual(check.source_sha256, file_sha256(workbook_path))
        workbook = load_workbook(workbook_path, read_only=True, data_only=False)
        self.addCleanup(workbook.close)
        self.assertEqual(workbook["当前Skill"]["A521"].value, "GH-01-0520")

    def test_excel_role_is_fail_closed_and_daily_overview_cannot_hide_behind_instructions(self):
        daily = self.root / "维护日报.xlsx"
        workbook = Workbook()
        instructions = workbook.active
        instructions.title = "使用说明"
        instructions.append(["说明", "这张表有内容"])
        overview = workbook.create_sheet("执行概览")
        overview.append(["项目", "结果"])
        workbook.save(daily)
        workbook.close()

        rejected = verify_excel(daily, run_id="run-daily-empty", role="daily")

        self.assertFalse(rejected.passed)
        self.assertEqual(rejected.key_sheet, "执行概览")
        self.assertIn("数据行", rejected.error or "")

        workbook = load_workbook(daily)
        workbook["执行概览"].append(["运行状态", "完成"])
        workbook.save(daily)
        workbook.close()
        accepted = verify_excel(daily, run_id="run-daily-filled", role="daily")
        self.assertTrue(accepted.passed, accepted.error)
        self.assertEqual(accepted.key_sheet, "执行概览")
        self.assertEqual((accepted.last_row, accepted.last_column, accepted.last_value), (2, 2, "完成"))

        arbitrary = self.root / "任意首表.xlsx"
        workbook = Workbook()
        workbook.active.title = "任意首表"
        workbook.active.append(["项目", "结果"])
        workbook.active.append(["伪装", "有数据"])
        workbook.save(arbitrary)
        workbook.close()
        missing = verify_excel(arbitrary, run_id="run-ledger-missing", role="ledger")
        self.assertFalse(missing.passed)
        self.assertIn("关键工作表", missing.error or "")

    def test_ledger_uses_run_record_when_current_skill_has_no_formal_id(self):
        ledger = self.root / "ledger-with-nonformal-residue.xlsx"
        workbook = Workbook()
        current = workbook.active
        current.title = "当前Skill"
        current.append(["内部标识", "Skill名称"])
        current["B2"] = "非正式残留说明"
        runs = workbook.create_sheet("运行记录")
        runs.append(["运行标识", "状态"])
        runs.append(["run-ledger-fallback", "成功"])
        workbook.save(ledger)
        workbook.close()

        check = verify_excel(ledger, run_id="run-ledger-fallback", role="ledger")

        self.assertTrue(check.passed, check.error)
        self.assertEqual(check.key_sheet, "运行记录")
        self.assertEqual((check.last_row, check.last_column, check.last_value), (2, 2, "成功"))

        workbook = load_workbook(ledger)
        del workbook["运行记录"]
        workbook.save(ledger)
        workbook.close()
        rejected = verify_excel(ledger, run_id="run-ledger-no-fallback", role="ledger")
        self.assertFalse(rejected.passed)
        self.assertEqual(rejected.key_sheet, "当前Skill")
        self.assertIn("数据行", rejected.error or "")

    def test_excel_rejects_header_only_and_corrupt_inputs_without_leaking_processes(self):
        empty = self.root / "empty.xlsx"
        corrupt = self.root / "corrupt.xlsx"
        write_excel(empty, rows=0)
        write_corrupt_xlsx(corrupt)

        empty_check = verify_excel(empty)
        corrupt_check = verify_excel(corrupt)

        self.assertFalse(empty_check.passed)
        self.assertIn("数据行", empty_check.error or "")
        self.assertEqual(empty_check.process_count_before, empty_check.process_count_after)
        self.assertFalse(corrupt_check.passed)
        self.assertTrue(corrupt_check.error)
        self.assertEqual(corrupt_check.process_count_before, corrupt_check.process_count_after)

    def test_word_exports_pdf_renders_pages_and_requires_hash_bound_visual_decision(self):
        word_path = self.root / "report.docx"
        write_word(word_path)

        renderer = write_renderer_command(self.root)
        first = verify_word(word_path, self.root / "render-first", renderer=renderer, run_id="run-word-first")
        second = verify_word(word_path, self.root / "render-second", renderer=renderer, run_id="run-word-second")

        for check in (first, second):
            self.assertTrue(check.office_passed, check.error)
            self.assertFalse(check.passed, "未绑定外部逐页视觉判定时不得发布")
            self.assertTrue(check.read_only)
            self.assertEqual(check.process_count_before, check.process_count_after)
            self.assertTrue(check.pdf_path and check.pdf_path.is_file())
            self.assertTrue(check.pdf_sha256)
            self.assertGreaterEqual(len(check.page_paths), 1)
            self.assertEqual(len(check.page_paths), len(check.page_sha256))
            self.assertFalse(check.blank_pages)
            self.assertTrue(all(path.is_file() and path.stat().st_size > 0 for path in check.page_paths))

        approved = bind_word_visual_decision(
            first,
            WordRenderDecision.from_check(first, approved=True, reviewer="Task 11 external visual review"),
        )
        self.assertTrue(approved.passed, approved.error)
        rejected = bind_word_visual_decision(
            second,
            WordRenderDecision.from_check(second, approved=False, reviewer="Task 11 external visual review", rejected_pages=(1,)),
        )
        self.assertFalse(rejected.passed)
        self.assertIn("视觉", rejected.error or "")

    def test_word_rejects_nonempty_render_directory_without_mutating_it(self):
        word_path = self.root / "report.docx"
        render = self.root / "existing-render"
        render.mkdir()
        sentinel = render / "keep.txt"
        sentinel.write_text("owned by caller", encoding="utf-8")
        write_word(word_path)

        with self.assertRaisesRegex(OfficeVerificationError, "空目录"):
            verify_word(word_path, render, renderer=write_renderer_command(self.root))

        self.assertEqual(sentinel.read_text(encoding="utf-8"), "owned by caller")
        self.assertEqual(tuple(render.iterdir()), (sentinel,))

    def test_word_visual_binding_rejects_hand_constructed_check(self):
        page = self.root / "page-1.png"
        Image.new("RGB", (64, 64), "white").save(page)
        source = self.root / "sample.docx"
        pdf = self.root / "sample.pdf"
        source.write_bytes(b"docx")
        pdf.write_bytes(b"pdf")
        check = OfficeCheck(
            kind="word", source_path=source, source_sha256=file_sha256(source),
            passed=False, office_passed=True, office_opened=True, read_only=True,
            pdf_path=pdf, pdf_sha256=file_sha256(pdf), page_paths=(page,),
            page_sha256=(file_sha256(page),), blank_pages=(1,),
            process_count_before=0, process_count_after=0,
        )
        decision = WordRenderDecision.from_check(check, approved=True, reviewer="external")
        with self.assertRaisesRegex(OfficeVerificationError, "签发|capability"):
            bind_word_visual_decision(check, decision)

    def test_office_result_requires_exact_process_baseline_not_merely_no_increase(self):
        workbook_path = self.root / "baseline.xlsx"
        write_excel(workbook_path, rows=1)
        result = {
            "passed": True, "office_opened": True, "read_only": True,
            "key_sheet": "当前Skill", "last_row": 2, "last_column": 2, "last_value": "skill-1",
            "process_count_before": 2, "process_count_after": 1, "error": None,
        }
        with patch("skill_maintainer.office._run_office", return_value=result):
            check = verify_excel(workbook_path)
        self.assertFalse(check.passed)
        self.assertIn("基线", check.error or "")

    def test_corrupt_docx_open_failure_returns_to_exact_word_process_baseline(self):
        corrupt = self.root / "损坏 报告.docx"
        corrupt.write_bytes(b"not-an-ooxml-package")
        check = verify_word(corrupt, self.root / "损坏 渲染", renderer=write_renderer_command(self.root))
        self.assertFalse(check.passed)
        self.assertFalse(check.office_passed)
        self.assertEqual(check.process_count_before, check.process_count_after)

    def test_word_export_failure_after_open_returns_to_exact_process_baseline(self):
        from skill_maintainer.office import _run_office

        source = self.root / "export-failure.docx"
        render = self.root / "export-failure-render"
        render.mkdir()
        write_word(source)
        # A directory at the exact PDF target makes ExportAsFixedFormat fail only
        # after Word has opened the document, exercising the COM cleanup path.
        (render / "export-failure.office.pdf").mkdir()
        result = _run_office("-Word", str(source), "-RenderDirectory", str(render))
        self.assertFalse(result["passed"])
        self.assertTrue(result["office_opened"])
        self.assertTrue(result["error"])
        self.assertEqual(result["process_count_before"], result["process_count_after"])

    def test_powershell_pdf_stability_probe_waits_for_two_equal_nonzero_sizes(self):
        from skill_maintainer.office import _run_office

        pdf = self.root / "slow-export.pdf"

        def slow_writer():
            time.sleep(0.10)
            pdf.write_bytes(b"first")
            time.sleep(0.04)
            with pdf.open("ab") as stream:
                stream.write(b"-final")

        writer = threading.Thread(target=slow_writer)
        writer.start()
        try:
            result = _run_office(
                "-StableFileProbe", str(pdf),
                "-StableTimeoutMilliseconds", "3000",
                "-StablePollMilliseconds", "80",
            )
        finally:
            writer.join(timeout=3)

        self.assertTrue(result["passed"], result.get("error"))
        self.assertEqual(result["stable_observations"], 2)
        self.assertEqual(result["observed_size"], len(b"first-final"))
        timeout = _run_office(
            "-StableFileProbe", str(self.root / "never-created.pdf"),
            "-StableTimeoutMilliseconds", "180",
            "-StablePollMilliseconds", "50",
        )
        self.assertFalse(timeout["passed"])
        self.assertEqual(timeout["error"], "word-pdf-stability-timeout")

    def test_renderer_contract_supports_clean_chinese_space_path_and_detects_header_footer_only_page(self):
        source = self.root / "中文 路径报告.docx"
        write_word(source)
        check = verify_word(
            source,
            self.root / "中文 渲染目录",
            renderer=write_renderer_command(self.root, body_pixels=0),
        )
        self.assertTrue(check.office_passed, check.error)
        self.assertEqual(check.blank_pages, tuple(range(1, len(check.page_paths) + 1)))
        self.assertFalse(check.passed)

    def test_visual_decision_must_enumerate_and_approve_each_exact_page_hash(self):
        source = self.root / "decision-pages.docx"
        write_word(source)
        check = verify_word(
            source, self.root / "decision-render", renderer=write_renderer_command(self.root),
            run_id="run-incomplete-decision",
        )
        incomplete = WordRenderDecision(
            source_sha256=check.source_sha256,
            pdf_sha256=check.pdf_sha256 or "",
            pages=(),
            reviewer="external",
        )
        with self.assertRaisesRegex(OfficeVerificationError, "每一页|逐页"):
            bind_word_visual_decision(check, incomplete)
        complete = WordRenderDecision.from_check(check, approved=True, reviewer="external")
        page = check.page_paths[0]
        original = page.read_bytes()
        page.write_bytes(original + b"tampered")
        with self.assertRaisesRegex(OfficeVerificationError, "哈希"):
            bind_word_visual_decision(check, complete)
        page.write_bytes(original)
        page.unlink()
        with self.assertRaisesRegex(OfficeVerificationError, "缺失"):
            bind_word_visual_decision(check, complete)


class PublicationTestCase(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory()
        self.addCleanup(self.temporary.cleanup)
        self.addCleanup(clear_office_run_state)
        self.root = Path(self.temporary.name)

    def make_tree(self, name: str = "run-20260828-220000") -> tuple[Path, Path]:
        staging = self.root / "staging" / name
        production = self.root / f"production-{name}"
        deliveries = staging / "deliveries"
        deliveries.mkdir(parents=True)
        (deliveries / "受影响专业类").mkdir()
        (deliveries / "检查摘要.txt").write_text("approved", encoding="utf-8")
        (deliveries / "受影响专业类" / "0809.bin").write_bytes(b"scope")
        write_excel(staging / "Skills主台账.xlsx", rows=2)
        (production / "ledger" / "archive").mkdir(parents=True)
        (production / "output" / "generations").mkdir(parents=True)
        write_excel(production / "ledger" / "Skills主台账.xlsx", rows=1)
        return staging, production

    def make_scoped_tree(self, project: str, run_id: str) -> tuple[Path, Path]:
        staging = self.root / project / "staging" / run_id
        production = self.root / project / "production"
        deliveries = staging / "deliveries"
        deliveries.mkdir(parents=True)
        (deliveries / "检查摘要.txt").write_text("approved", encoding="utf-8")
        write_excel(staging / "Skills主台账.xlsx", rows=2)
        (production / "ledger" / "archive").mkdir(parents=True)
        (production / "output" / "generations").mkdir(parents=True)
        write_excel(production / "ledger" / "Skills主台账.xlsx", rows=1)
        return staging, production

    @staticmethod
    def tree_snapshot(root: Path) -> tuple[tuple[str, str, bytes | None], ...]:
        records: list[tuple[str, str, bytes | None]] = []
        for path in sorted(root.rglob("*"), key=lambda item: item.relative_to(root).as_posix()):
            relative = path.relative_to(root).as_posix()
            records.append((relative, "directory", None) if path.is_dir() else (relative, "file", path.read_bytes()))
        return tuple(records)

    @staticmethod
    def evidence_for(staging: Path):
        ledger = staging / "Skills主台账.xlsx"
        result = {
            "passed": True, "office_opened": True, "read_only": True,
            "key_sheet": "当前Skill", "last_row": 3, "last_column": 2,
            "last_value": "skill-2", "process_count_before": 0,
            "process_count_after": 0, "error": None,
        }
        with patch("skill_maintainer.office._run_office", return_value=result):
            check = verify_excel(ledger, run_id=staging.name, role="ledger")
        return OfficeEvidenceBundle.from_checks((check,), run_id=staging.name)

    def plan_for(self, staging: Path, production: Path):
        return build_publish_plan(staging, production, office_evidence=self.evidence_for(staging))

    def test_only_real_verifier_capability_can_issue_run_bound_one_time_evidence(self):
        staging, production = self.make_tree("run-office-capability")
        ledger = staging / "Skills主台账.xlsx"
        office_result = {
            "passed": True, "office_opened": True, "read_only": True,
            "key_sheet": "当前Skill", "last_row": 3, "last_column": 2,
            "last_value": "skill-2", "process_count_before": 0,
            "process_count_after": 0, "error": None,
        }
        with patch("skill_maintainer.office._run_office", return_value=office_result):
            issued = verify_excel(ledger, run_id=staging.name, role="ledger")
        bundle = OfficeEvidenceBundle.from_checks((issued,), run_id=staging.name)
        bundle.assert_covers((ledger,))
        bundle.assert_covers((ledger,))  # pre-publication validation remains retryable

        for forged in (
            trusted_excel_check(ledger), copy.copy(issued), replace(issued),
            replace(issued, last_value="tampered"),
        ):
            with self.subTest(forged=forged.last_value):
                with self.assertRaisesRegex(OfficeVerificationError, "签发|受信|capability"):
                    OfficeEvidenceBundle.from_checks((forged,), run_id=staging.name)
        with self.assertRaisesRegex(OfficeVerificationError, "运行|run"):
            OfficeEvidenceBundle.from_checks((issued,), run_id="another-run")
        for forged_bundle in (copy.copy(bundle), replace(bundle)):
            with self.assertRaisesRegex(OfficeVerificationError, "签发|受信|capability"):
                forged_bundle.assert_covers((ledger,))

        plan = build_publish_plan(staging, production, office_evidence=bundle)
        publish_atomically(plan)
        with self.assertRaisesRegex(OfficeVerificationError, "消费|受信|签发"):
            bundle.assert_covers((ledger,))

    def test_registry_cleanup_is_scoped_to_its_own_run(self):
        first_staging, _ = self.make_tree("run-registry-first")
        second_staging, _ = self.make_tree("run-registry-second")
        first = self.evidence_for(first_staging)
        second = self.evidence_for(second_staging)

        clear_office_run_state(first.run_id)

        with self.assertRaisesRegex(OfficeVerificationError, "消费|受信|签发"):
            first.assert_covers((first_staging / "Skills主台账.xlsx",))
        second.assert_covers((second_staging / "Skills主台账.xlsx",))

    def test_same_requested_run_id_cleanup_is_isolated_by_staging_scope(self):
        run_id = "run-shared-across-projects"
        for mode in ("success", "failure", "base-exception", "explicit-clear"):
            with self.subTest(mode=mode):
                first_staging, first_production = self.make_scoped_tree(f"{mode}-project-a", run_id)
                second_staging, _ = self.make_scoped_tree(f"{mode}-project-b", run_id)
                first = self.evidence_for(first_staging)
                second = self.evidence_for(second_staging)

                if mode == "explicit-clear":
                    clear_office_run_state(run_id, scope_root=first_staging)
                else:
                    plan = build_publish_plan(first_staging, first_production, office_evidence=first)
                    if mode == "success":
                        publish_atomically(plan)
                    elif mode == "failure":
                        with self.assertRaisesRegex(PublishError, "注入失败"):
                            publish_atomically(plan, fail_at="generation-replace")
                    else:
                        def stop_before_backup() -> None:
                            raise SystemExit("stop before backup replace")

                        with self.assertRaises(SystemExit):
                            publish_atomically(plan, before_backup_replace=stop_before_backup)

                with self.assertRaisesRegex(OfficeVerificationError, "消费|受信|签发"):
                    first.assert_covers((first_staging / "Skills主台账.xlsx",))
                second.assert_covers((second_staging / "Skills主台账.xlsx",))

    def test_shared_commit_rejects_every_invalid_preflight_input_before_any_production_side_effect(self):
        cases = (
            "wrong-role", "forged", "wrong-run", "wrong-generation-path",
            "wrong-authority-hash", "office-set-mismatch",
        )
        for case in cases:
            with self.subTest(case=case):
                run_id = f"run-preflight-{case}"
                staging = self.root / case / "staging" / run_id
                production = self.root / case / "production"
                staging.mkdir(parents=True)
                write_excel(staging / "Skills主台账.xlsx", rows=2)
                (production / "ledger").mkdir(parents=True)
                generations = production / "output" / "generations"
                generation = generations / run_id
                generation.mkdir(parents=True)
                write_excel(production / "ledger" / "Skills主台账.xlsx", rows=1)
                (generation / "payload.txt").write_text("prepared", encoding="utf-8")
                manifest = generation / "generation-manifest.json"
                manifest.write_text('{"prepared":true}', encoding="utf-8")

                if case == "office-set-mismatch":
                    write_excel(generation / "unexpected.xlsx", rows=1)
                if case == "wrong-role":
                    result = {
                        "passed": True, "office_opened": True, "read_only": True,
                        "key_sheet": "执行概览", "last_row": 2, "last_column": 2,
                        "last_value": "完成", "process_count_before": 0,
                        "process_count_after": 0, "error": None,
                    }
                    with patch("skill_maintainer.office._run_office", return_value=result):
                        check = verify_excel(
                            staging / "Skills主台账.xlsx", run_id=run_id, role="daily",
                        )
                    evidence = OfficeEvidenceBundle.from_checks((check,), run_id=run_id)
                elif case == "wrong-run":
                    evidence_run_id = f"{run_id}-evidence"
                    result = {
                        "passed": True, "office_opened": True, "read_only": True,
                        "key_sheet": "当前Skill", "last_row": 3, "last_column": 2,
                        "last_value": "skill-2", "process_count_before": 0,
                        "process_count_after": 0, "error": None,
                    }
                    with patch("skill_maintainer.office._run_office", return_value=result):
                        check = verify_excel(
                            staging / "Skills主台账.xlsx", run_id=evidence_run_id, role="ledger",
                        )
                    evidence = OfficeEvidenceBundle.from_checks((check,), run_id=evidence_run_id)
                else:
                    evidence = self.evidence_for(staging)
                    if case == "forged":
                        evidence = replace(evidence)

                commit_generation = (
                    generations / "wrong-generation-name"
                    if case == "wrong-generation-path" else generation
                )
                expected_authority = (
                    "0" * 64
                    if case == "wrong-authority-hash"
                    else file_sha256(production / "ledger" / "Skills主台账.xlsx")
                )

                before = self.tree_snapshot(production)
                with self.assertRaisesRegex(PublishError, "Office|证据|ledger|主台账|发布代次|generation"):
                    commit_prepared_generation(
                        production_root=production,
                        run_id=run_id,
                        staged_ledger=staging / "Skills主台账.xlsx",
                        expected_authority_sha256=expected_authority,
                        generation_path=commit_generation,
                        generation_manifest_sha256=file_sha256(manifest),
                        office_evidence=evidence,
                        office_paths=(staging / "Skills主台账.xlsx",),
                    )

                self.assertEqual(self.tree_snapshot(production), before)
                self.assertFalse((production / "ledger" / "archive").exists())

    def test_checks_issue_one_bundle_atomically_and_concurrently(self):
        staging, _ = self.make_tree("run-check-issuance")
        first_path = staging / "Skills主台账.xlsx"
        second_path = staging / "second.xlsx"
        write_excel(second_path, rows=1)

        def issue_check(path):
            result = {
                "passed": True, "office_opened": True, "read_only": True,
                "key_sheet": "当前Skill", "last_row": 2, "last_column": 2,
                "last_value": "verified", "process_count_before": 0,
                "process_count_after": 0, "error": None,
            }
            with patch("skill_maintainer.office._run_office", return_value=result):
                return verify_excel(path, run_id=staging.name, role="ledger")

        first = issue_check(first_path)
        second = issue_check(second_path)
        with self.assertRaisesRegex(OfficeVerificationError, "签发|受信|capability"):
            OfficeEvidenceBundle.from_checks((first, replace(second)), run_id=staging.name)
        first_bundle = OfficeEvidenceBundle.from_checks((first,), run_id=staging.name)
        first_bundle.assert_covers((first_path,))
        with self.assertRaisesRegex(OfficeVerificationError, "已签发|单次|capability"):
            OfficeEvidenceBundle.from_checks((first,), run_id=staging.name)

        concurrent_path = staging / "concurrent.xlsx"
        write_excel(concurrent_path, rows=1)
        concurrent_check = issue_check(concurrent_path)
        barrier = threading.Barrier(2)

        def sign_once():
            barrier.wait(timeout=3)
            try:
                return OfficeEvidenceBundle.from_checks((concurrent_check,), run_id=staging.name)
            except OfficeVerificationError as exc:
                return exc

        with ThreadPoolExecutor(max_workers=2) as executor:
            results = tuple(executor.map(lambda _: sign_once(), range(2)))
        self.assertEqual(sum(isinstance(item, OfficeEvidenceBundle) for item in results), 1)
        self.assertEqual(sum(isinstance(item, OfficeVerificationError) for item in results), 1)

    def test_delivery_excel_verified_as_ledger_is_rejected_before_side_effects(self):
        staging, production = self.make_tree("run-role-path-binding")
        daily = staging / "deliveries" / "维护日报.xlsx"
        workbook = Workbook()
        current = workbook.active
        current.title = "当前Skill"
        current.append(["内部标识", "Skill名称"])
        current.append(["WRONG-ROLE-1", "wrong-role"])
        overview = workbook.create_sheet("执行概览")
        overview.append(["项目", "结果"])
        workbook.save(daily)
        workbook.close()
        wrong_daily = verify_excel(daily, run_id=staging.name, role="ledger")
        self.assertTrue(wrong_daily.passed, wrong_daily.error)
        ledger_result = {
            "passed": True, "office_opened": True, "read_only": True,
            "key_sheet": "当前Skill", "last_row": 3, "last_column": 2,
            "last_value": "skill-2", "process_count_before": 0,
            "process_count_after": 0, "error": None,
        }
        with patch("skill_maintainer.office._run_office", return_value=ledger_result):
            ledger_check = verify_excel(
                staging / "Skills主台账.xlsx", run_id=staging.name, role="ledger",
            )
        evidence = OfficeEvidenceBundle.from_checks(
            (ledger_check, wrong_daily), run_id=staging.name,
        )
        authority = production / "ledger" / "Skills主台账.xlsx"
        authority_before = authority.read_bytes()

        with self.assertRaisesRegex(PublishError, "角色|daily|交付.*Excel"):
            build_publish_plan(staging, production, office_evidence=evidence)

        self.assertEqual(authority.read_bytes(), authority_before)
        self.assertEqual(tuple((production / "ledger" / "archive").iterdir()), ())
        self.assertEqual(tuple((production / "output" / "generations").iterdir()), ())

    def test_publish_revalidates_role_path_binding_before_backup_or_generation(self):
        staging, production = self.make_tree("run-role-boundary")
        daily = staging / "deliveries" / "维护日报.xlsx"
        workbook = Workbook()
        current = workbook.active
        current.title = "当前Skill"
        current.append(["内部标识", "Skill名称"])
        current.append(["ROLE-1", "role"])
        overview = workbook.create_sheet("执行概览")
        overview.append(["项目", "结果"])
        overview.append(["运行状态", "完成"])
        workbook.save(daily)
        workbook.close()

        def issued(path, role, key_sheet, last_value):
            result = {
                "passed": True, "office_opened": True, "read_only": True,
                "key_sheet": key_sheet, "last_row": 2, "last_column": 2,
                "last_value": last_value, "process_count_before": 0,
                "process_count_after": 0, "error": None,
            }
            with patch("skill_maintainer.office._run_office", return_value=result):
                return verify_excel(path, run_id=staging.name, role=role)

        ledger = staging / "Skills主台账.xlsx"
        correct = OfficeEvidenceBundle.from_checks((
            issued(ledger, "ledger", "当前Skill", "skill-2"),
            issued(daily, "daily", "执行概览", "完成"),
        ), run_id=staging.name)
        plan = build_publish_plan(staging, production, office_evidence=correct)
        wrong = OfficeEvidenceBundle.from_checks((
            issued(ledger, "ledger", "当前Skill", "skill-2"),
            issued(daily, "ledger", "当前Skill", "role"),
        ), run_id=staging.name)
        forged = replace(plan, office_evidence=wrong, office_evidence_sha256=wrong.sha256)
        authority_before = forged.authority_path.read_bytes()

        with self.assertRaisesRegex(PublishError, "角色|daily|交付.*Excel"):
            publish_atomically(forged)

        self.assertEqual(forged.authority_path.read_bytes(), authority_before)
        self.assertEqual(tuple(forged.backup_path.parent.iterdir()), ())
        self.assertEqual(tuple(forged.generation_path.parent.iterdir()), ())

    def test_build_plan_binds_every_input_and_single_authority(self):
        staging, production = self.make_tree()
        plan = self.plan_for(staging, production)

        self.assertEqual(plan.run_id, staging.name)
        self.assertEqual(plan.authority_path, production / "ledger" / "Skills主台账.xlsx")
        self.assertEqual(plan.expected_authority_sha256, file_sha256(plan.authority_path))
        self.assertEqual(plan.staged_ledger_sha256, file_sha256(staging / "Skills主台账.xlsx"))
        self.assertEqual(
            tuple(item.relative_path for item in plan.delivery_files),
            ("受影响专业类/0809.bin", "检查摘要.txt"),
        )
        self.assertTrue(all(item.sha256 for item in plan.delivery_files))
        self.assertEqual(plan.generation_path, production / "output" / "generations" / staging.name)
        self.assertEqual(plan.backup_path.parent, production / "ledger" / "archive")
        self.assertRegex(plan.backup_path.name, r"^Skills主台账_\d{8}_\d{6}(?:_\d+)?\.xlsx$")

    def test_publish_copies_backup_and_generation_before_authority_replace(self):
        staging, production = self.make_tree()
        plan = self.plan_for(staging, production)

        receipt = publish_atomically(plan)

        self.assertEqual(file_sha256(plan.authority_path), plan.staged_ledger_sha256)
        self.assertTrue(receipt.backup_path.is_file())
        self.assertEqual(receipt.backup_sha256, plan.expected_authority_sha256)
        self.assertEqual(receipt.office_evidence_sha256, plan.office_evidence_sha256)
        self.assertEqual(receipt.generation_path, plan.generation_path)
        self.assertTrue(receipt.generation_path.is_dir())
        for item in plan.delivery_files:
            published = receipt.generation_path / Path(item.relative_path)
            self.assertEqual(file_sha256(published), item.sha256)
        self.assertTrue((receipt.generation_path / "generation-manifest.json").is_file())
        self.assertFalse(any(production.rglob("*receipt*")), "发布回执只在内存返回，不落业务文件")

    def test_publish_fsyncs_every_staged_input_before_copying_or_committing(self):
        staging, production = self.make_tree()
        plan = self.plan_for(staging, production)
        from skill_maintainer import publish as publish_module
        real_fsync_file = publish_module._fsync_file
        fsynced: list[Path] = []

        def record_and_fsync(path: Path) -> None:
            fsynced.append(path)
            real_fsync_file(path)

        with patch("skill_maintainer.publish._fsync_file", side_effect=record_and_fsync):
            publish_atomically(plan)

        required = {plan.staged_ledger}
        required.update(plan.deliveries_root / Path(item.relative_path) for item in plan.delivery_files)
        self.assertTrue(required.issubset(set(fsynced)), (required, fsynced))

    def test_changed_authority_is_rejected_immediately_before_replace(self):
        staging, production = self.make_tree()
        plan = self.plan_for(staging, production)
        plan.authority_path.write_bytes(b"concurrent-change")

        with self.assertRaisesRegex(PublishError, "生产主台账.*变化"):
            publish_atomically(plan)

        self.assertEqual(plan.authority_path.read_bytes(), b"concurrent-change")
        self.assertFalse(plan.generation_path.exists())

    def test_failure_at_each_destination_preserves_old_authority_and_prior_backups(self):
        failure_points = (
            "delivery:受影响专业类/0809.bin",
            "delivery:检查摘要.txt",
            "generation-replace",
            "authority-temp",
            "authority-replace",
        )
        for index, failure_point in enumerate(failure_points, start=1):
            with self.subTest(failure_point=failure_point):
                staging, production = self.make_tree(f"run-{index}")
                prior = production / "ledger" / "archive" / "Skills主台账_20260101_000000.xlsx"
                prior.write_bytes(b"prior-backup")
                plan = self.plan_for(staging, production)
                old = plan.authority_path.read_bytes()

                with self.assertRaisesRegex(PublishError, "注入失败"):
                    publish_atomically(plan, fail_at=failure_point)

                self.assertEqual(plan.authority_path.read_bytes(), old)
                self.assertEqual(prior.read_bytes(), b"prior-backup")
                self.assertTrue(plan.backup_path.is_file())
                self.assertEqual(plan.backup_path.read_bytes(), old)
                self.assertFalse(plan.generation_path.exists())
                self.assertFalse(any(path.name.startswith(f".{plan.run_id}") for path in plan.generation_path.parent.iterdir()))
                self.assertFalse(any(production.rglob("*receipt*")))

    def test_staged_input_change_is_rejected_without_partial_publication(self):
        staging, production = self.make_tree()
        plan = self.plan_for(staging, production)
        (staging / "deliveries" / "检查摘要.txt").write_bytes(b"changed-after-plan")

        with self.assertRaisesRegex(PublishError, "暂存.*变化"):
            publish_atomically(plan)

        self.assertEqual(file_sha256(plan.authority_path), plan.expected_authority_sha256)
        self.assertFalse(plan.generation_path.exists())
        with self.assertRaisesRegex(OfficeVerificationError, "消费|受信|签发"):
            plan.office_evidence.assert_covers((plan.staged_ledger,))
        retry_plan = self.plan_for(staging, production)
        receipt = publish_atomically(retry_plan)
        self.assertTrue(receipt.generation_path.is_dir())

    def test_forged_delivery_traversal_is_rejected_before_any_publication(self):
        staging, production = self.make_tree()
        escaped = staging / "escaped.xlsx"
        escaped.write_bytes(b"outside-deliveries")
        plan = self.plan_for(staging, production)
        forged = replace(
            plan,
            delivery_files=(PublishFile("../escaped.xlsx", file_sha256(escaped)),),
        )

        with self.assertRaisesRegex(PublishError, "相对路径"):
            publish_atomically(forged)

        self.assertEqual(file_sha256(plan.authority_path), plan.expected_authority_sha256)
        self.assertFalse(plan.generation_path.exists())

    def test_build_plan_requires_exact_passed_office_evidence_for_every_office_artifact(self):
        staging, production = self.make_tree("run-office-evidence")
        report = staging / "deliveries" / "报告.docx"
        write_word(report)
        evidence = self.evidence_for(staging)
        with self.assertRaisesRegex(PublishError, "Office.*证据|证据.*Office"):
            build_publish_plan(staging, production, office_evidence=evidence)
        evidence.assert_covers((staging / "Skills主台账.xlsx",))
        report.unlink()
        plan = build_publish_plan(staging, production, office_evidence=evidence)
        self.assertEqual(plan.office_evidence, evidence)

    def test_build_plan_rejects_forged_office_evidence_digest(self):
        staging, production = self.make_tree("run-forged-evidence")
        evidence = replace(self.evidence_for(staging), sha256="0" * 64)
        with self.assertRaisesRegex(PublishError, "Office.*证据|证据.*摘要"):
            build_publish_plan(staging, production, office_evidence=evidence)

    def test_backup_collision_after_plan_fails_without_replacing_existing_backup(self):
        staging, production = self.make_tree("run-backup-collision")
        plan = self.plan_for(staging, production)
        plan.backup_path.write_bytes(b"pre-existing-backup")

        with self.assertRaisesRegex(PublishError, "备份.*存在|覆盖.*备份"):
            publish_atomically(plan)

        self.assertEqual(plan.backup_path.read_bytes(), b"pre-existing-backup")
        self.assertEqual(file_sha256(plan.authority_path), plan.expected_authority_sha256)
        self.assertFalse(plan.generation_path.exists())

    def test_authority_replace_then_base_exception_keeps_committed_generation_and_backup(self):
        staging, production = self.make_tree("run-linearized")
        plan = self.plan_for(staging, production)
        original_replace = os.replace

        def replace_then_crash(source, destination):
            result = original_replace(source, destination)
            if Path(destination) == plan.authority_path:
                raise SystemExit("crash after real authority replace")
            return result

        with patch("skill_maintainer.publish.os.replace", side_effect=replace_then_crash):
            with self.assertRaises(SystemExit):
                publish_atomically(plan)

        self.assertEqual(file_sha256(plan.authority_path), plan.staged_ledger_sha256)
        self.assertTrue(plan.generation_path.is_dir())
        self.assertTrue(plan.backup_path.is_file())
        with self.assertRaisesRegex(OfficeVerificationError, "消费|受信|签发"):
            plan.office_evidence.assert_covers((plan.staged_ledger,))

    @unittest.skipUnless(os.name == "nt", "Windows directory handle pinning acceptance")
    def test_generation_and_parent_directories_are_handle_pinned_through_authority_replace(self):
        staging, production = self.make_tree("run-parent-pins")
        plan = self.plan_for(staging, production)
        attempts: list[str] = []

        def attack_boundary() -> None:
            generation_file = plan.generation_path / "检查摘要.txt"
            for label, operation in (
                ("write", lambda: generation_file.write_text("tamper", encoding="utf-8")),
                ("file-rename", lambda: generation_file.rename(generation_file.with_suffix(".moved"))),
                ("file-delete", generation_file.unlink),
                ("generation-parent-rename", lambda: plan.generation_path.parent.rename(plan.generation_path.parent.with_name("generations-moved"))),
                ("archive-parent-rename", lambda: plan.backup_path.parent.rename(plan.backup_path.parent.with_name("archive-moved"))),
            ):
                with self.assertRaises(PermissionError, msg=label):
                    operation()
                attempts.append(label)

        receipt = publish_atomically(plan, before_authority_replace=attack_boundary)
        self.assertEqual(attempts, ["write", "file-rename", "file-delete", "generation-parent-rename", "archive-parent-rename"])
        self.assertTrue(receipt.generation_path.is_dir())

    @unittest.skipUnless(os.name == "nt", "Windows directory handle pinning acceptance")
    def test_archive_parent_is_pinned_before_backup_and_cannot_be_swapped_to_outside(self):
        staging, production = self.make_tree("run-archive-parent-pin")
        plan = self.plan_for(staging, production)
        outside = self.root / "outside-archive"
        outside.mkdir()

        def attack_backup_boundary() -> None:
            with self.assertRaises(PermissionError):
                plan.backup_path.parent.rename(production / "ledger" / "archive-moved")
            self.assertEqual(tuple(outside.iterdir()), ())

        receipt = publish_atomically(plan, before_backup_replace=attack_backup_boundary)
        self.assertTrue(receipt.backup_path.is_file())
        self.assertEqual(tuple(outside.iterdir()), ())


if __name__ == "__main__":
    unittest.main()
