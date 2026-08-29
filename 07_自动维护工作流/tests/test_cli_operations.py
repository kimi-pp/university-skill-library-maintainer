"""Task 13: deployment operations and the long-lived two-gate CLI protocol."""

from __future__ import annotations

from contextlib import redirect_stdout
from dataclasses import asdict, replace
from hashlib import sha256
import importlib.util
import io
import json
import os
from pathlib import Path
import shutil
import subprocess
import sys
import tempfile
import tomllib
import unittest
from unittest.mock import patch

from docx import Document
from openpyxl import Workbook

from skill_maintainer import cli
from skill_maintainer.ledger import LedgerStore
from skill_maintainer.locking import SingleWriterLock
from skill_maintainer.office import RendererCommand
from skill_maintainer.review import build_review_packet
from skill_maintainer.runner import RunCoordinator, RunRequest, SourceRun
from skill_maintainer.settings import load_settings, settings_sha256
from skill_maintainer.snapshots import SnapshotCandidate, build_snapshot


WORKFLOW = "07_自动维护工作流"
RULE_FILES = (
    "SKILL_RESEARCH_WORKFLOW.md",
    "SECURITY_REVIEW_PROTOCOL.md",
    "DATA_DICTIONARY.md",
    "REPORTING_STANDARD.md",
)


class _ProtocolAbort(BaseException):
    pass


class _DoctorEnvironment:
    def __init__(
        self,
        *,
        windows: bool = True,
        python: bool = True,
        gh: bool = True,
        word: bool = True,
        excel: bool = True,
        loader_output: str | None = None,
    ) -> None:
        self.is_windows = windows
        self.python_available = python
        self.gh_available = gh
        self.word_available = word
        self.excel_available = excel
        self.loader_output = loader_output


class _RecordingOutput:
    def __init__(self, *, exception_on_type: tuple[str, BaseException] | None = None) -> None:
        self.text = ""
        self.flush_count = 0
        self.exception_on_type = exception_on_type
        self.exception_raised = False

    def write(self, value: str) -> int:
        if self.exception_on_type is not None and not self.exception_raised:
            frame_type, exception = self.exception_on_type
            try:
                payload = json.loads(value)
            except json.JSONDecodeError:
                payload = {}
            if payload.get("type") == frame_type:
                self.exception_raised = True
                raise exception
        self.text += value
        return len(value)

    def flush(self) -> None:
        self.flush_count += 1

    @property
    def frames(self) -> list[dict[str, object]]:
        return [json.loads(line) for line in self.text.splitlines() if line.strip()]


class _ReactiveInput:
    """Respond to the latest flushed request without precomputing evidence hashes."""

    def __init__(
        self,
        output: _RecordingOutput,
        *,
        visual_approved: bool = True,
        eof_at: str | None = None,
        exception_at: tuple[str, BaseException] | None = None,
    ) -> None:
        self.output = output
        self.visual_approved = visual_approved
        self.eof_at = eof_at
        self.exception_at = exception_at
        self.seen: set[str] = set()

    def readline(self) -> str:
        request = self.output.frames[-1]
        frame_type = str(request["type"])
        if frame_type in self.seen:
            raise AssertionError(f"protocol requested {frame_type} twice")
        self.seen.add(frame_type)
        if self.exception_at is not None and self.exception_at[0] == frame_type:
            raise self.exception_at[1]
        if self.eof_at == frame_type:
            return ""
        if frame_type == "review_required":
            return json.dumps(
                {"type": "review_decisions", "run_id": request["run_id"], "decisions": []},
                ensure_ascii=False,
            ) + "\n"
        if frame_type == "word_visual_review_required":
            documents = []
            for document in request["documents"]:
                documents.append(
                    {
                        "source_sha256": document["source_sha256"],
                        "pdf_sha256": document["pdf_sha256"],
                        "reviewer": "Task13 离线复核",
                        "pages": [
                            {
                                "page_number": page["page_number"],
                                "sha256": page["sha256"],
                                "approved": self.visual_approved,
                            }
                            for page in document["pages"]
                        ],
                    }
                )
            return json.dumps(
                {"type": "word_visual_decisions", "run_id": request["run_id"], "documents": documents},
                ensure_ascii=False,
            ) + "\n"
        raise AssertionError(f"unexpected protocol request: {frame_type}")


class _UnitRenderer:
    def render(self, pdf: Path, output_dir: Path):
        page = output_dir / "page-1.png"
        page.write_bytes(b"visible-page-body")
        return (page,), (1,)


class _AttackingReader:
    def __init__(self, handle, attack, *, on_exit: bool) -> None:
        self.handle = handle
        self.attack = attack
        self.on_exit = on_exit
        self.attacked = False

    def __enter__(self):
        self.handle.__enter__()
        return self

    def __exit__(self, exception_type, exception, traceback):
        result = self.handle.__exit__(exception_type, exception, traceback)
        if self.on_exit and not self.attacked:
            self.attacked = True
            self.attack()
        return result

    def read(self, *args, **kwargs):
        value = self.handle.read(*args, **kwargs)
        if not self.on_exit and not self.attacked:
            self.attacked = True
            self.attack()
        return value

    def __getattr__(self, name):
        return getattr(self.handle, name)


class CliOperationsTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory()
        self.addCleanup(self.temporary.cleanup)
        self.root = Path(self.temporary.name) / "中文 项目"
        self.root.mkdir()
        self.source_workflow = Path(__file__).resolve().parents[1]

    def _bundled_python(self) -> Path:
        candidate = (
            Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime"
            / "dependencies" / "python" / "python.exe"
        )
        self.assertTrue(candidate.is_file(), f"bundled dependency runtime missing: {candidate}")
        return candidate

    def _authority_fixture(self, *, missing: str | None = None) -> None:
        (self.root / "AGENTS.md").write_text("# authority\n", encoding="utf-8")
        rule_root = self.root / "01_规则"
        rule_root.mkdir()
        for name in RULE_FILES:
            if name != missing:
                (rule_root / name).write_text(f"# {name}\n", encoding="utf-8")

    def _workflow_fixture(self) -> Path:
        workflow = self.root / WORKFLOW
        workflow.mkdir()
        shutil.copyfile(self.source_workflow / "workflow-settings.example.toml", workflow / "workflow-settings.example.toml")
        shutil.copytree(self.source_workflow / "skill", workflow / "skill")
        source_dir = workflow / "src" / "skill_maintainer"
        source_dir.mkdir(parents=True)
        for name in ("pdf_renderer.py", "workspace_renderer.py", "daily_xlsx_builder.mjs"):
            shutil.copyfile(self.source_workflow / "src" / "skill_maintainer" / name, source_dir / name)
        return workflow

    def _setup_project(self) -> Path:
        self._authority_fixture()
        workflow = self._workflow_fixture()
        result = cli.setup_project(self.root, source_workflow=workflow)
        self.assertEqual(result.exit_code, 0)
        return workflow

    @staticmethod
    def _enabled_settings(path: Path) -> None:
        text = path.read_text(encoding="utf-8")
        path.write_text(
            text.replace("enabled = false", "enabled = true").replace('mode = "manual"', 'mode = "daily"'),
            encoding="utf-8",
        )

    def test_parser_keeps_exact_public_commands_and_requires_explicit_project_root(self):
        parser = cli.build_parser()
        commands = parser._subparsers._group_actions[0].choices
        self.assertEqual(
            set(commands),
            {
                "setup", "import-existing", "doctor", "edit-settings", "apply-settings",
                "run-now", "scheduled-run", "status", "repair-ledger", "rebuild-report",
                "prepare", "apply-reviews", "finalize",
            },
        )
        with self.assertRaises(SystemExit) as caught:
            parser.parse_args(["doctor"])
        self.assertEqual(caught.exception.code, 2)

    def test_setup_is_idempotent_creates_disabled_manual_state_and_never_overwrites_settings(self):
        self._authority_fixture()
        workflow = self._workflow_fixture()
        first = cli.setup_project(self.root, source_workflow=workflow)
        self.assertEqual(first.exit_code, 0, first.message)
        settings_path = workflow / "workflow-settings.toml"
        settings = load_settings(settings_path)
        self.assertFalse(settings.workflow.enabled)
        self.assertEqual(settings.schedule.mode, "manual")
        self.assertTrue((workflow / "ledger" / "Skills主台账.xlsx").is_file())

        custom = settings_path.read_text(encoding="utf-8").replace("include_generic_skills = false", "include_generic_skills = true")
        settings_path.write_text(custom, encoding="utf-8")
        second = cli.setup_project(self.root, source_workflow=workflow)
        self.assertEqual(second.exit_code, 0)
        self.assertEqual(settings_path.read_text(encoding="utf-8"), custom)

    def test_setup_fails_closed_on_missing_authority_without_creating_state(self):
        self._authority_fixture(missing="SKILL_RESEARCH_WORKFLOW.md")
        workflow = self._workflow_fixture()
        result = cli.setup_project(self.root, source_workflow=workflow)
        self.assertEqual(result.exit_code, 1)
        self.assertFalse((workflow / "workflow-settings.toml").exists())
        self.assertFalse((workflow / "ledger").exists())

    def test_setup_installs_and_updates_skill_in_supplied_codex_root(self):
        self._authority_fixture()
        workflow = self._workflow_fixture()
        skills_root = Path(self.temporary.name) / "Codex Skills"
        first = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(first.exit_code, 0, first.message)
        installed = skills_root / "university-skill-library-maintainer" / "SKILL.md"
        self.assertEqual(installed.read_bytes(), (workflow / "skill" / "university-skill-library-maintainer" / "SKILL.md").read_bytes())
        installed.write_text("stale", encoding="utf-8")
        second = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(second.exit_code, 0)
        self.assertNotEqual(installed.read_text(encoding="utf-8"), "stale")

    def test_skill_directory_update_removes_files_absent_from_new_source(self):
        self._authority_fixture()
        workflow = self._workflow_fixture()
        skills_root = Path(self.temporary.name) / "Codex Skills"
        first = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(first.exit_code, 0, first.message)
        installed = skills_root / "university-skill-library-maintainer"
        stale = installed / "obsolete-from-old-version.txt"
        stale.write_text("must disappear", encoding="utf-8")

        second = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(second.exit_code, 0)
        self.assertFalse(stale.exists())

    def test_skill_directory_switch_failure_rolls_back_complete_previous_tree(self):
        self._authority_fixture()
        workflow = self._workflow_fixture()
        skills_root = Path(self.temporary.name) / "Codex Skills"
        first = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(first.exit_code, 0, first.message)
        installed = skills_root / "university-skill-library-maintainer"
        marker = installed / "SKILL.md"
        marker.write_text("previous complete skill", encoding="utf-8")
        stale = installed / "previous-only.txt"
        stale.write_text("previous", encoding="utf-8")
        original_rename = os.rename
        calls = 0

        def fail_install_switch(source, destination):
            nonlocal calls
            calls += 1
            if calls == 2:
                raise OSError("injected directory switch failure")
            return original_rename(source, destination)

        with patch("skill_maintainer.cli.os.rename", side_effect=fail_install_switch):
            result = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(result.exit_code, 1)
        self.assertEqual(marker.read_text(encoding="utf-8"), "previous complete skill")
        self.assertEqual(stale.read_text(encoding="utf-8"), "previous")
        pending = tuple(path for path in skills_root.iterdir() if path.name.endswith(".pending"))
        self.assertEqual(len(pending), 1)
        self.assertIn(str(pending[0]), result.message)
        self.assertFalse(any(path.name.endswith(".previous") for path in skills_root.iterdir()))

    def test_skill_failed_switch_never_deletes_unknown_destination_to_force_rollback(self):
        self._authority_fixture()
        workflow = self._workflow_fixture()
        skills_root = Path(self.temporary.name) / "Codex Skills"
        first = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(first.exit_code, 0, first.message)
        installed = skills_root / cli.SKILL_NAME
        (installed / "old-version.txt").write_text("old", encoding="utf-8")
        real_rename = os.rename
        calls = 0
        previous_path: Path | None = None

        def occupy_destination_then_fail(source, destination):
            nonlocal calls, previous_path
            calls += 1
            if calls == 1:
                result = real_rename(source, destination)
                previous_path = Path(destination)
                return result
            if calls == 2:
                Path(destination).mkdir()
                (Path(destination) / "user-data.txt").write_text("must remain", encoding="utf-8")
                raise OSError("injected switch failure after unknown destination appeared")
            return real_rename(source, destination)

        with patch("skill_maintainer.cli.os.rename", side_effect=occupy_destination_then_fail):
            result = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(result.exit_code, 1)
        self.assertEqual((installed / "user-data.txt").read_text(encoding="utf-8"), "must remain")
        self.assertIsNotNone(previous_path)
        self.assertEqual((previous_path / "old-version.txt").read_text(encoding="utf-8"), "old")
        self.assertIn(str(installed), result.message)
        self.assertIn(str(previous_path), result.message)

    def test_skill_forged_previous_sidecar_never_authorizes_deletion(self):
        self._authority_fixture()
        workflow = self._workflow_fixture()
        skills_root = Path(self.temporary.name) / "Codex Skills"
        first = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(first.exit_code, 0, first.message)
        operation = "a" * 32
        fake = skills_root / f".{cli.SKILL_NAME}.{operation}.previous"
        fake.mkdir()
        (fake / "user-data.txt").write_text("must remain", encoding="utf-8")
        forged = skills_root / f".{cli.SKILL_NAME}.{operation}.previous.owner"
        forged.write_bytes(
            ("university-skill-library-maintainer previous tree\n" + operation + "\n").encode("ascii")
        )

        result = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(result.exit_code, 0, result.message)
        self.assertEqual((fake / "user-data.txt").read_text(encoding="utf-8"), "must remain")
        self.assertTrue(forged.is_file())
        self.assertTrue(result.details and any(str(fake) in item for item in result.details.get("warnings", ())))

    def test_skill_update_preserves_previous_for_manual_cleanup(self):
        self._authority_fixture()
        workflow = self._workflow_fixture()
        skills_root = Path(self.temporary.name) / "Codex Skills"
        first = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(first.exit_code, 0, first.message)
        installed = skills_root / "university-skill-library-maintainer"
        (installed / "SKILL.md").write_text("previous incomplete version", encoding="utf-8")
        old_only = installed / "previous-only.txt"
        old_only.write_text("old", encoding="utf-8")
        result = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(result.exit_code, 0, result.message)
        self.assertEqual(
            cli._skill_tree_manifest(installed),
            cli._skill_tree_manifest(workflow / "skill" / "university-skill-library-maintainer"),
        )
        self.assertFalse(old_only.exists())
        self.assertTrue(result.details and result.details.get("warnings"))
        leftovers = tuple(path for path in skills_root.iterdir() if path.name.endswith(".previous"))
        self.assertEqual(len(leftovers), 1)
        self.assertEqual((leftovers[0] / "previous-only.txt").read_text(encoding="utf-8"), "old")
        self.assertEqual((leftovers[0] / "SKILL.md").read_text(encoding="utf-8"), "previous incomplete version")
        self.assertTrue(any(str(leftovers[0]) in item for item in result.details.get("warnings", ())))

        unknown = skills_root / f".{cli.SKILL_NAME}.not-owned.previous"
        unknown.mkdir()
        outside = Path(self.temporary.name) / "outside-previous-link"
        outside.mkdir()
        (outside / "must-remain.txt").write_text("outside", encoding="utf-8")
        linked = skills_root / f".{cli.SKILL_NAME}.{'f' * 32}.previous"
        linked.symlink_to(outside, target_is_directory=True)

        third = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(third.exit_code, 0, third.message)
        self.assertTrue(leftovers[0].is_dir(), "cross-run cleanup must not delete an unknown tree")
        self.assertTrue(unknown.is_dir(), "unknown similarly named tree must not be deleted")
        self.assertTrue(linked.is_symlink(), "reparse previous entry must not be followed or deleted")
        self.assertEqual((outside / "must-remain.txt").read_text(encoding="utf-8"), "outside")
        warnings = tuple(third.details.get("warnings", ())) if third.details else ()
        self.assertTrue(any(str(leftovers[0]) in item and "普通目录" in item for item in warnings))
        self.assertTrue(any(str(linked) in item and "链接或重解析点" in item for item in warnings))

    def test_skill_previous_replaced_after_commit_validation_is_never_deleted(self):
        self._authority_fixture()
        workflow = self._workflow_fixture()
        skills_root = Path(self.temporary.name) / "Codex Skills"
        first = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(first.exit_code, 0, first.message)
        installed = skills_root / cli.SKILL_NAME
        real_manifest = cli._skill_tree_manifest
        destination_checks = 0
        previous_path: Path | None = None
        held = skills_root / "held-committed-previous"

        def replace_after_commit_validation(root):
            nonlocal destination_checks, previous_path
            result = real_manifest(root)
            if Path(root) == installed:
                destination_checks += 1
                if destination_checks == 2:
                    previous_path = next(
                        path for path in skills_root.iterdir()
                        if path.name.endswith(".previous") and not path.is_symlink()
                    )
                    os.rename(previous_path, held)
                    previous_path.mkdir()
                    (previous_path / "user-data.txt").write_text("must remain", encoding="utf-8")
            return result

        with patch("skill_maintainer.cli._skill_tree_manifest", side_effect=replace_after_commit_validation):
            result = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(destination_checks, 2)
        self.assertEqual(result.exit_code, 0, result.message)
        self.assertIsNotNone(previous_path)
        self.assertEqual((previous_path / "user-data.txt").read_text(encoding="utf-8"), "must remain")
        self.assertTrue((held / "SKILL.md").is_file())
        self.assertEqual(
            cli._skill_tree_manifest(installed),
            cli._skill_tree_manifest(workflow / "skill" / cli.SKILL_NAME),
        )

    def test_skill_unknown_pending_created_after_commit_is_never_deleted(self):
        self._authority_fixture()
        workflow = self._workflow_fixture()
        skills_root = Path(self.temporary.name) / "Codex Skills"
        first = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(first.exit_code, 0, first.message)
        installed = skills_root / cli.SKILL_NAME
        real_manifest = cli._skill_tree_manifest
        destination_checks = 0
        unknown_pending: Path | None = None

        def create_unknown_pending_after_commit(root):
            nonlocal destination_checks, unknown_pending
            result = real_manifest(root)
            if Path(root) == installed:
                destination_checks += 1
                if destination_checks == 2:
                    previous = next(path for path in skills_root.iterdir() if path.name.endswith(".previous"))
                    operation = previous.name.removesuffix(".previous").rsplit(".", 1)[-1]
                    unknown_pending = skills_root / f".{cli.SKILL_NAME}.{operation}.pending"
                    unknown_pending.mkdir()
                    (unknown_pending / "user-data.txt").write_text("must remain", encoding="utf-8")
            return result

        with patch("skill_maintainer.cli._skill_tree_manifest", side_effect=create_unknown_pending_after_commit):
            result = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(result.exit_code, 0, result.message)
        self.assertIsNotNone(unknown_pending)
        self.assertEqual((unknown_pending / "user-data.txt").read_text(encoding="utf-8"), "must remain")
        self.assertEqual(
            cli._skill_tree_manifest(installed),
            cli._skill_tree_manifest(workflow / "skill" / cli.SKILL_NAME),
        )
        self.assertTrue(result.details and any(str(unknown_pending) in item for item in result.details.get("warnings", ())))

    def test_skill_stale_previous_replaced_after_status_is_never_deleted(self):
        self._authority_fixture()
        workflow = self._workflow_fixture()
        skills_root = Path(self.temporary.name) / "Codex Skills"
        first = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertEqual(first.exit_code, 0, first.message)
        stale = skills_root / f".{cli.SKILL_NAME}.{'b' * 32}.previous"
        stale.mkdir()
        (stale / "old-owned.txt").write_text("old", encoding="utf-8")
        held = skills_root / "held-old-previous"
        real_status = cli._stale_previous_status
        replaced = False

        def replace_after_status(candidate):
            nonlocal replaced
            result = real_status(candidate)
            if Path(candidate) == stale and not replaced:
                replaced = True
                os.rename(stale, held)
                stale.mkdir()
                (stale / "user-data.txt").write_text("must remain", encoding="utf-8")
            return result

        with patch("skill_maintainer.cli._stale_previous_status", side_effect=replace_after_status):
            result = cli.setup_project(self.root, source_workflow=workflow, codex_skills_root=skills_root)
        self.assertTrue(replaced)
        self.assertEqual(result.exit_code, 0, result.message)
        self.assertEqual((stale / "user-data.txt").read_text(encoding="utf-8"), "must remain")
        self.assertEqual((held / "old-owned.txt").read_text(encoding="utf-8"), "old")

    def test_skill_root_reparse_with_missing_nested_target_causes_no_external_write(self):
        self._authority_fixture()
        workflow = self._workflow_fixture()
        outside = Path(self.temporary.name) / "outside-skills"
        outside.mkdir()
        linked = Path(self.temporary.name) / "linked-skills"
        linked.symlink_to(outside, target_is_directory=True)

        result = cli.setup_project(
            self.root,
            source_workflow=workflow,
            codex_skills_root=linked / "missing" / "nested",
        )
        self.assertEqual(result.exit_code, 1)
        self.assertEqual(list(outside.iterdir()), [])

    def test_rebuild_reparse_output_with_missing_nested_target_causes_no_external_write(self):
        workflow = self._setup_project()
        output_root = workflow / "output"
        output_root.rmdir()
        outside = Path(self.temporary.name) / "outside-rebuild"
        outside.mkdir()
        output_root.symlink_to(outside, target_is_directory=True)

        result = cli.rebuild_reports(self.root, output=output_root / "missing" / "nested")
        self.assertEqual(result.exit_code, 1)
        self.assertEqual(list(outside.iterdir()), [])

    def test_import_reparse_staging_with_missing_nested_target_causes_no_external_write(self):
        workflow = self._setup_project()
        (self.root / "05_交付物").mkdir()
        outside = Path(self.temporary.name) / "outside-import"
        outside.mkdir()
        staging = workflow / "ledger" / "staging"
        staging.symlink_to(outside, target_is_directory=True)
        output = staging / "missing" / "nested" / "首次导入.xlsx"

        stream = io.StringIO()
        with redirect_stdout(stream):
            code = cli.main(["import-existing", "--project-root", str(self.root), "--output", str(output)])
        self.assertEqual(code, 1)
        self.assertEqual(list(outside.iterdir()), [])

    def test_doctor_checks_dependencies_rules_renderer_and_blocks_non_windows_production(self):
        workflow = self._setup_project()
        environment = _DoctorEnvironment(windows=False, python=True, gh=True, word=True, excel=True)
        diagnostic = cli.doctor_project(self.root, environment=environment)
        self.assertEqual(diagnostic.exit_code, 0)
        self.assertFalse(diagnostic.production_ready)
        self.assertEqual(
            set(diagnostic.checks),
            {
                "AGENTS.md", *RULE_FILES, "workflow root", "settings", "ledger",
                "Python", "GitHub CLI", "Microsoft Word", "Microsoft Excel",
                "renderer", "production_driver",
            },
        )
        self.assertEqual(diagnostic.checks["renderer"], "等待 Codex 工作区依赖加载器")
        self._enabled_settings(workflow / "workflow-settings.toml")
        production = cli.doctor_project(self.root, environment=environment)
        self.assertEqual(production.exit_code, 1)
        self.assertIn("Windows", " ".join(production.errors))
        self.assertEqual(production.checks["production_driver"], "已配置")

    def test_doctor_fails_when_any_required_dependency_or_rule_is_missing(self):
        workflow = self._setup_project()
        (self.root / "01_规则" / "REPORTING_STANDARD.md").unlink()
        environment = _DoctorEnvironment(python=False, gh=False, word=False, excel=False)
        report = cli.doctor_project(self.root, environment=environment)
        self.assertEqual(report.exit_code, 1)
        self.assertEqual(report.checks["REPORTING_STANDARD.md"], "缺失")
        self.assertEqual(report.checks["Python"], "不可用")
        self.assertEqual(report.checks["GitHub CLI"], "不可用")
        self.assertEqual(report.checks["Microsoft Word"], "不可用")
        self.assertEqual(report.checks["Microsoft Excel"], "不可用")
        self.assertTrue((workflow / "workflow-settings.toml").is_file())

    def test_apply_settings_returns_validated_prompt_schedule_and_hash_without_mutating_automation(self):
        workflow = self._setup_project()
        settings_path = workflow / "workflow-settings.toml"
        before = {path.relative_to(self.root): sha256(path.read_bytes()).hexdigest() for path in self.root.rglob("*") if path.is_file()}
        plan = cli.build_apply_settings_plan(self.root)
        after = {path.relative_to(self.root): sha256(path.read_bytes()).hexdigest() for path in self.root.rglob("*") if path.is_file()}
        self.assertEqual(before, after)
        self.assertEqual(plan["config_sha256"], settings_sha256(settings_path))
        self.assertEqual(plan["automation_action"], "ensure_absent")
        self.assertEqual(plan["schedule"]["mode"], "manual")
        self.assertEqual(plan["project_root"], str(self.root.absolute()))
        self.assertEqual(plan["toml_path"], str(settings_path.absolute()))
        self.assertIn("scheduled-run", plan["prompt"])
        self.assertNotIn("::automation", plan["prompt"])
        self.assertFalse(plan["production_ready"])

    def test_invalid_configuration_returns_exit_two(self):
        workflow = self._setup_project()
        (workflow / "workflow-settings.toml").write_text("not = [valid", encoding="utf-8")
        output = io.StringIO()
        with redirect_stdout(output):
            code = cli.main(["apply-settings", "--project-root", str(self.root)])
        self.assertEqual(code, 2)
        self.assertEqual(json.loads(output.getvalue())["exit_code"], 2)

    def test_doctor_invalid_configuration_returns_exit_two(self):
        workflow = self._setup_project()
        (workflow / "workflow-settings.toml").write_text("not = [valid", encoding="utf-8")
        report = cli.doctor_project(self.root, environment=_DoctorEnvironment())
        self.assertEqual(report.exit_code, 2)

    def test_corrupt_authority_returns_machine_readable_operational_failure(self):
        workflow = self._setup_project()
        (workflow / "ledger" / "Skills主台账.xlsx").write_bytes(b"not-a-zip-workbook")
        for command in ("doctor", "status", "apply-settings"):
            with self.subTest(command=command):
                output = io.StringIO()
                with redirect_stdout(output):
                    code = cli.main([command, "--project-root", str(self.root)])
                self.assertEqual(code, 1, output.getvalue())
                payload = json.loads(output.getvalue().splitlines()[-1])
                self.assertEqual(payload.get("exit_code"), 1)

    def test_user_supplied_output_and_backup_validation_return_exit_two(self):
        workflow = self._setup_project()
        (self.root / "05_交付物").mkdir()
        cases = (
            ["import-existing", "--project-root", str(self.root), "--output", str(self.root / "outside.xlsx")],
            ["rebuild-report", "--project-root", str(self.root), "--output", str(self.root / "outside-rebuild")],
            ["repair-ledger", "--project-root", str(self.root), "--backup", str(workflow / "ledger" / "not-listed.xlsx")],
        )
        for arguments in cases:
            with self.subTest(command=arguments[0]):
                output = io.StringIO()
                with redirect_stdout(output):
                    code = cli.main(arguments)
                self.assertEqual(code, 2, output.getvalue())

    def test_status_reports_driver_configured_but_environment_not_ready_and_import_defaults_to_safe_noop(self):
        self._setup_project()
        (self.root / "05_交付物").mkdir()
        status = cli.status_project(self.root)
        self.assertFalse(status["production_ready"])
        self.assertEqual(status["production_driver"], "已配置")
        output = io.StringIO()
        with redirect_stdout(output):
            code = cli.main(["import-existing", "--project-root", str(self.root)])
        self.assertEqual(code, 3)
        self.assertEqual(json.loads(output.getvalue())["excel_files"], 0)

    def test_status_ignores_unrecorded_generation_directory(self):
        workflow = self._setup_project()
        fake = workflow / "output" / "generations" / "run-unrecorded"
        fake.mkdir(parents=True)
        (fake / "plausible.docx").write_bytes(b"not authoritative")

        status = cli.status_project(self.root)
        self.assertIsNone(status["latest_output"])
        self.assertIn("成功运行记录", status["output_error"])

    def test_status_rejects_tampered_generation_bound_by_last_success_record(self):
        workflow = self._complete_protocol_run()
        first = cli.status_project(self.root)
        self.assertIsNotNone(first["latest_output"])
        generation = Path(first["latest_output"])
        delivery = next(path for path in generation.rglob("*.docx"))
        delivery.write_bytes(delivery.read_bytes() + b"tampered")

        tampered = cli.status_project(self.root)
        self.assertIsNone(tampered["latest_output"])
        self.assertIn("发布代次", tampered["output_error"])

    def test_status_and_apply_plan_do_not_treat_driver_factory_as_full_production_readiness(self):
        self._setup_project()
        with patch("skill_maintainer.cli.PRODUCTION_DRIVER_FACTORY", object()):
            status = cli.status_project(self.root)
            plan = cli.build_apply_settings_plan(self.root)
        self.assertFalse(status["production_ready"])
        self.assertFalse(plan["production_ready"])
        self.assertIn("doctor", status)
        self.assertIn("doctor", plan)
        self.assertEqual(status["doctor"].checks["production_driver"], "生产发现驱动配置不可调用")

    def test_apply_settings_accepts_loader_output_for_full_doctor_plan(self):
        self._setup_project()
        output = io.StringIO()
        try:
            with redirect_stdout(output):
                code = cli.main([
                    "apply-settings", "--project-root", str(self.root),
                    "--loader-output", "not-a-real-loader-result",
                ])
        except SystemExit as exc:
            code = int(exc.code)
        self.assertEqual(code, 0)
        plan = json.loads(output.getvalue())
        self.assertFalse(plan["production_ready"])
        self.assertIn("doctor", plan)

    def test_repair_lists_only_valid_backups_and_explicit_choice_never_overwrites_authority(self):
        workflow, valid = self._published_backup_fixture()
        authority = workflow / "ledger" / "Skills主台账.xlsx"
        archive = workflow / "ledger" / "archive"
        (archive / "Skills主台账_20260829_010204.xlsx").write_bytes(b"not-an-xlsx")
        original_sha = sha256(authority.read_bytes()).hexdigest()

        listed = cli.repair_ledger(self.root)
        self.assertEqual(listed.exit_code, 3)
        self.assertEqual(listed.valid_backups, (valid.absolute(),))
        chosen = cli.repair_ledger(self.root, backup=valid)
        self.assertEqual(chosen.exit_code, 0)
        self.assertEqual(sha256(authority.read_bytes()).hexdigest(), original_sha)
        self.assertIsNotNone(chosen.recovery_candidate)
        self.assertEqual(sha256(chosen.recovery_candidate.read_bytes()).hexdigest(), sha256(valid.read_bytes()).hexdigest())

    def test_repair_excludes_valid_workbook_not_bound_to_successful_publication(self):
        workflow = self._setup_project()
        authority = workflow / "ledger" / "Skills主台账.xlsx"
        unbound = workflow / "ledger" / "archive" / "Skills主台账_20260829_010203.xlsx"
        shutil.copyfile(authority, unbound)

        result = cli.repair_ledger(self.root)
        self.assertEqual(result.exit_code, 3)
        self.assertEqual(result.valid_backups, ())

    def test_repair_corrupt_authority_uses_only_backup_self_bound_to_verified_generation(self):
        workflow = self._complete_protocol_run()
        authority = workflow / "ledger" / "Skills主台账.xlsx"
        archive = workflow / "ledger" / "archive"
        committed_backup = archive / "Skills主台账_20260829_010203_1.xlsx"
        shutil.copyfile(authority, committed_backup)
        arbitrary = archive / "Skills主台账_20260829_010203_2.xlsx"
        store = LedgerStore.create(arbitrary)
        store.workbook.close()
        authority.write_bytes(b"corrupt-current-authority")

        listed = cli.repair_ledger(self.root)
        self.assertEqual(listed.exit_code, 3, listed.message)
        self.assertEqual(listed.valid_backups, (committed_backup.absolute(),))
        recovered = cli.repair_ledger(self.root, backup=committed_backup)
        self.assertEqual(recovered.exit_code, 0, recovered.message)
        self.assertIsNotNone(recovered.recovery_candidate)
        self.assertEqual(recovered.recovery_candidate.read_bytes(), committed_backup.read_bytes())

    def test_repair_corrupt_authority_rechecks_identity_after_backup_self_binding(self):
        workflow = self._complete_protocol_run()
        authority = workflow / "ledger" / "Skills主台账.xlsx"
        backup = workflow / "ledger" / "archive" / "Skills主台账_20260829_010203_1.xlsx"
        shutil.copyfile(authority, backup)
        replacement = backup.with_name("same-bytes-self-binding-replacement.xlsx")
        shutil.copyfile(backup, replacement)
        authority.write_bytes(b"corrupt-current-authority")
        real_capture = cli._capture_backup_snapshot
        attacked = False

        def capture_then_replace(candidate):
            nonlocal attacked
            result = real_capture(candidate)
            if Path(candidate) == backup and not attacked:
                attacked = True
                os.replace(replacement, backup)
            return result

        with patch(
            "skill_maintainer.cli._capture_backup_snapshot",
            side_effect=capture_then_replace,
        ):
            listed = cli.repair_ledger(self.root)
        self.assertTrue(attacked)
        self.assertEqual(listed.exit_code, 3)
        self.assertNotIn(backup.absolute(), listed.valid_backups)

    def test_repair_rejects_backup_aba_swap_during_self_binding(self):
        workflow = self._complete_protocol_run()
        authority = workflow / "ledger" / "Skills主台账.xlsx"
        archive = workflow / "ledger" / "archive"
        backup = archive / "Skills主台账_20260829_010203_7.xlsx"
        arbitrary = LedgerStore.create(backup)
        arbitrary.workbook.close()
        original_bytes = backup.read_bytes()
        genuine = archive / "genuine-self-binding-source.xlsx"
        shutil.copyfile(authority, genuine)
        held_original = archive / "held-original-during-aba.xlsx"
        authority.write_bytes(b"corrupt-current-authority")
        real_self_binding = cli._backup_self_binds_verified_generation
        attacked = False

        def aba_bind(workflow_path, evidence):
            nonlocal attacked
            targets_backup = (
                isinstance(evidence, (str, os.PathLike)) and Path(evidence) == backup
            ) or (
                not isinstance(evidence, (str, os.PathLike)) and not tuple(evidence)
            )
            if not targets_backup or attacked:
                return real_self_binding(workflow_path, evidence)
            attacked = True
            os.replace(backup, held_original)
            os.replace(genuine, backup)
            try:
                return real_self_binding(workflow_path, evidence)
            finally:
                os.replace(backup, genuine)
                os.replace(held_original, backup)

        with patch(
            "skill_maintainer.cli._backup_self_binds_verified_generation",
            side_effect=aba_bind,
        ):
            listed = cli.repair_ledger(self.root)
        self.assertTrue(attacked)
        self.assertEqual(backup.read_bytes(), original_bytes)
        self.assertEqual(listed.exit_code, 3)
        self.assertNotIn(backup.absolute(), listed.valid_backups)

    def test_repair_accepts_publishers_same_second_backup_suffix(self):
        workflow, first = self._published_backup_fixture()
        collision = first.with_name(first.stem + "_1.xlsx")
        shutil.copyfile(first, collision)

        result = cli.repair_ledger(self.root)
        self.assertEqual(result.exit_code, 3)
        self.assertEqual(result.valid_backups, (first.absolute(), collision.absolute()))

    def _assert_repair_rejects_live_backup_attack(self, *, replace_path: bool) -> None:
        workflow, backup = self._published_backup_fixture()
        authority = workflow / "ledger" / "Skills主台账.xlsx"
        authority_sha = sha256(authority.read_bytes()).hexdigest()
        replacement_bytes = authority.read_bytes()
        original_open = Path.open
        attacked = False
        replacement = backup.parent / "replacement.xlsx"
        replacement.write_bytes(replacement_bytes)

        def attack():
            nonlocal attacked
            attacked = True
            if replace_path:
                os.replace(replacement, backup)
            else:
                with original_open(backup, "wb") as handle:
                    handle.write(replacement_bytes)
                    handle.flush()
                    os.fsync(handle.fileno())

        def open_with_attack(path, mode="r", *args, **kwargs):
            handle = original_open(path, mode, *args, **kwargs)
            if Path(path).absolute() == backup.absolute() and mode == "rb" and not attacked:
                return _AttackingReader(handle, attack, on_exit=replace_path)
            return handle

        with patch.object(Path, "open", new=open_with_attack):
            result = cli.repair_ledger(self.root, backup=backup)
        self.assertTrue(attacked)
        self.assertEqual(result.exit_code, 1)
        self.assertIsNone(result.recovery_candidate)
        self.assertEqual(sha256(authority.read_bytes()).hexdigest(), authority_sha)
        recovery = workflow / "ledger" / "recovery"
        self.assertFalse(recovery.exists() and any(recovery.iterdir()))

    def test_repair_rejects_selected_backup_replaced_after_handle_read(self):
        self._assert_repair_rejects_live_backup_attack(replace_path=True)

    def test_repair_rejects_selected_backup_rewritten_in_place_during_handle_read(self):
        self._assert_repair_rejects_live_backup_attack(replace_path=False)

    def test_repair_rejects_backup_replaced_between_validation_and_hash_binding(self):
        workflow, backup = self._published_backup_fixture()
        replacement = backup.parent / "same-bytes-new-identity.xlsx"
        shutil.copyfile(backup, replacement)
        real_validation = cli._validated_ledger_records
        attacked = False

        def validate_then_replace(content):
            nonlocal attacked
            result = real_validation(content)
            if not attacked:
                attacked = True
                os.replace(replacement, backup)
            return result

        with patch("skill_maintainer.cli._validated_ledger_records", side_effect=validate_then_replace):
            result = cli.repair_ledger(self.root, backup=backup)
        self.assertTrue(attacked)
        self.assertEqual(result.exit_code, 1)
        self.assertIsNone(result.recovery_candidate)
        recovery = workflow / "ledger" / "recovery"
        self.assertFalse(recovery.exists() and any(recovery.iterdir()))

    def test_repair_write_failure_leaves_no_partial_recovery_candidate(self):
        workflow, backup = self._published_backup_fixture()
        with patch("skill_maintainer.cli.os.fsync", side_effect=OSError("injected recovery flush failure")):
            result = cli.repair_ledger(self.root, backup=backup)
        self.assertEqual(result.exit_code, 1)
        recovery = workflow / "ledger" / "recovery"
        self.assertFalse(recovery.exists() and any(recovery.iterdir()))

    def test_rebuild_report_reads_only_current_ledger_and_never_calls_network(self):
        workflow = self._setup_project()
        output = workflow / "output" / "rebuild-test"
        calls: list[tuple[str, Path]] = []

        def word_builder(summary, path):
            calls.append(("word", Path(path)))
            Document().save(path)
            return Path(path)

        def excel_builder(summary, path):
            calls.append(("excel", Path(path)))
            Workbook().save(path)
            return Path(path)

        authority = workflow / "ledger" / "Skills主台账.xlsx"
        before = sha256(authority.read_bytes()).hexdigest()
        with patch("urllib.request.urlopen", side_effect=AssertionError("rebuild must be offline")):
            result = cli.rebuild_reports(self.root, output=output, word_builder=word_builder, excel_builder=excel_builder)
        self.assertEqual(result.exit_code, 0)
        self.assertEqual([name for name, _ in calls], ["word", "excel"])
        self.assertEqual(sha256(authority.read_bytes()).hexdigest(), before)
        self.assertEqual({path.suffix for path in result.outputs}, {".docx", ".xlsx"})

    def test_rebuild_builder_failure_leaves_requested_output_absent(self):
        workflow = self._setup_project()
        output = workflow / "output" / "atomic-rebuild"

        def word_builder(summary, path):
            Document().save(path)
            return Path(path)

        def excel_builder(summary, path):
            raise OSError("injected Excel builder failure")

        result = cli.rebuild_reports(
            self.root,
            output=output,
            word_builder=word_builder,
            excel_builder=excel_builder,
        )
        self.assertEqual(result.exit_code, 1)
        self.assertFalse(output.exists())
        self.assertFalse(any("pending" in path.name for path in output.parent.iterdir()))

    def test_rebuild_builder_failure_removes_only_new_empty_ancestors(self):
        workflow = self._setup_project()
        output_base = workflow / "output"
        existing = output_base / "existing-parent"
        existing.mkdir()
        output = existing / "new-one" / "new-two" / "report"

        def failing_word_builder(summary, path):
            raise OSError("injected nested rebuild failure")

        result = cli.rebuild_reports(self.root, output=output, word_builder=failing_word_builder)
        self.assertEqual(result.exit_code, 1)
        self.assertTrue(existing.is_dir())
        self.assertEqual(list(existing.iterdir()), [])

    def _protocol_fixture(
        self,
        *,
        visual_approved: bool = True,
        eof_at: str | None = None,
        exception_at: tuple[str, BaseException] | None = None,
        output_exception: tuple[str, BaseException] | None = None,
    ):
        workflow = self._setup_project()
        output = _RecordingOutput(exception_on_type=output_exception)
        input_stream = _ReactiveInput(
            output,
            visual_approved=visual_approved,
            eof_at=eof_at,
            exception_at=exception_at,
        )
        gate = cli.InteractiveOfficeGate(input_stream=input_stream, output_stream=output, renderer=_UnitRenderer())

        def build_reports(prepared, staging):
            delivery = staging / "deliveries"
            delivery.mkdir()
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "执行概览"
            sheet.append(["项目", "结果"])
            sheet.append(["离线四源", "完成"])
            xlsx = delivery / "离线查验.xlsx"
            workbook.save(xlsx)
            word = Document()
            word.add_paragraph("离线四源查验；候选未安装、未运行。")
            docx = delivery / "离线查验.docx"
            word.save(docx)
            return (xlsx, docx)

        sources = (
            SourceRun("SkillHub", "complete"),
            SourceRun("ClawHub", "complete"),
            SourceRun("GitHub", "complete"),
            SourceRun("Hugging Face Spaces", "complete"),
        )
        coordinator = RunCoordinator(
            root=workflow,
            discover=lambda request, staging: sources,
            report_builder=build_reports,
            office_verifier=gate,
        )
        request = RunRequest(
            settings_path=workflow / "workflow-settings.toml",
            catalog_loader=lambda: {"fixture": "offline-four-source"},
            requested_run_id="run-task13-offline",
        )
        return workflow, coordinator, request, input_stream, output

    def _complete_protocol_run(self):
        workflow, coordinator, request, input_stream, output = self._protocol_fixture()
        with patch("skill_maintainer.office._run_office", side_effect=lambda *args: self._fake_office_result(list(args))):
            code = cli.run_interactive_protocol(coordinator, request, input_stream=input_stream, output_stream=output)
        self.assertEqual(code, 0, output.text)
        return workflow

    def _published_backup_fixture(self):
        workflow = self._complete_protocol_run()
        backups = sorted((workflow / "ledger" / "archive").glob("Skills主台账_*.xlsx"))
        self.assertEqual(len(backups), 1)
        return workflow, backups[0]

    def _assert_runtime_released(self, workflow: Path) -> None:
        self.assertFalse(any((workflow / ".runtime" / "staging").iterdir()))
        lock = SingleWriterLock(workflow / ".runtime" / "writer.lock")
        self.assertTrue(lock.acquire())
        lock.release()

    @staticmethod
    def _force_coordinator_cleanup(coordinator: RunCoordinator) -> None:
        for state in tuple(coordinator._states.values()):
            try:
                coordinator.abandon(state.prepared)
            except BaseException:
                pass

    @staticmethod
    def _fake_office_result(arguments):
        if "-Excel" in arguments:
            role = arguments[arguments.index("-ExcelRole") + 1]
            return {
                "passed": True,
                "office_opened": True,
                "read_only": True,
                "key_sheet": "运行记录" if role == "ledger" else "执行概览",
                "last_row": 2,
                "last_column": 2,
                "last_value": "verified",
                "process_count_before": 0,
                "process_count_after": 0,
                "error": None,
            }
        render = Path(arguments[arguments.index("-RenderDirectory") + 1])
        pdf = render / "office.pdf"
        pdf.write_bytes(b"unit-pdf")
        return {
            "passed": True,
            "office_opened": True,
            "read_only": True,
            "pdf_path": str(pdf),
            "page_count": 1,
            "process_count_before": 0,
            "process_count_after": 0,
            "error": None,
        }

    def test_long_lived_protocol_keeps_one_process_through_review_visual_and_success(self):
        workflow, coordinator, request, input_stream, output = self._protocol_fixture()
        with patch("skill_maintainer.office._run_office", side_effect=lambda *args: self._fake_office_result(list(args))):
            code = cli.run_interactive_protocol(coordinator, request, input_stream=input_stream, output_stream=output)
        self.assertEqual(code, 0, output.text)
        self.assertEqual(
            [frame["type"] for frame in output.frames],
            ["review_required", "word_visual_review_required", "run_complete"],
        )
        self.assertGreaterEqual(output.flush_count, 3)
        self.assertTrue((workflow / "ledger" / "Skills主台账.xlsx").is_file())
        self.assertTrue(any((workflow / "output").rglob("*.docx")))
        lock = SingleWriterLock(workflow / ".runtime" / "writer.lock")
        self.assertTrue(lock.acquire())
        lock.release()

    def test_protocol_eof_after_prepare_abandons_staging_and_releases_capabilities(self):
        workflow, coordinator, request, input_stream, output = self._protocol_fixture(eof_at="review_required")
        code = cli.run_interactive_protocol(coordinator, request, input_stream=input_stream, output_stream=output)
        self.assertEqual(code, 1)
        self.assertEqual([frame["type"] for frame in output.frames], ["review_required", "run_failed"], output.text)
        self.assertFalse(any((workflow / ".runtime" / "staging").iterdir()))
        lock = SingleWriterLock(workflow / ".runtime" / "writer.lock")
        self.assertTrue(lock.acquire())
        lock.release()

    def test_protocol_keyboard_interrupt_keeps_original_exception_and_cleans_runtime(self):
        original = KeyboardInterrupt("operator interrupted review")
        workflow, coordinator, request, input_stream, output = self._protocol_fixture(
            exception_at=("review_required", original)
        )
        real_abandon = coordinator.abandon

        def clean_then_report(prepared):
            real_abandon(prepared)
            raise RuntimeError("cleanup diagnostic")

        try:
            with patch.object(coordinator, "abandon", side_effect=clean_then_report):
                with self.assertRaises(KeyboardInterrupt) as caught:
                    cli.run_interactive_protocol(coordinator, request, input_stream=input_stream, output_stream=output)
            self.assertIs(caught.exception, original)
            self.assertIn("cleanup diagnostic", " ".join(getattr(caught.exception, "__notes__", ())))
            self._assert_runtime_released(workflow)
        finally:
            self._force_coordinator_cleanup(coordinator)

    def test_protocol_system_exit_and_custom_base_exception_clean_runtime(self):
        for exception in (SystemExit(19), _ProtocolAbort("fatal control abort")):
            with self.subTest(exception=type(exception).__name__):
                with tempfile.TemporaryDirectory() as temporary:
                    old_root, old_temporary = self.root, self.temporary
                    self.root = Path(temporary) / "中文 项目"
                    self.root.mkdir()
                    coordinator = None
                    try:
                        workflow, coordinator, request, input_stream, output = self._protocol_fixture(
                            exception_at=("review_required", exception)
                        )
                        with self.assertRaises(type(exception)) as caught:
                            cli.run_interactive_protocol(
                                coordinator, request, input_stream=input_stream, output_stream=output
                            )
                        self.assertIs(caught.exception, exception)
                        self._assert_runtime_released(workflow)
                    finally:
                        if coordinator is not None:
                            self._force_coordinator_cleanup(coordinator)
                        self.root, self.temporary = old_root, old_temporary

    def test_protocol_base_exception_after_commit_does_not_remove_committed_authority(self):
        exception = SystemExit(23)
        workflow, coordinator, request, input_stream, output = self._protocol_fixture(
            output_exception=("run_complete", exception)
        )
        authority = workflow / "ledger" / "Skills主台账.xlsx"
        before = sha256(authority.read_bytes()).hexdigest()
        with patch("skill_maintainer.office._run_office", side_effect=lambda *args: self._fake_office_result(list(args))):
            with self.assertRaises(SystemExit) as caught:
                cli.run_interactive_protocol(coordinator, request, input_stream=input_stream, output_stream=output)
        self.assertIs(caught.exception, exception)
        self.assertNotEqual(sha256(authority.read_bytes()).hexdigest(), before)
        self.assertTrue(any((workflow / "output" / "generations").iterdir()))
        self._assert_runtime_released(workflow)

    def test_protocol_requires_one_decision_for_every_review_packet_before_publish(self):
        workflow, coordinator, request, input_stream, output = self._protocol_fixture()
        source = workflow / "trusted-fixture-candidate"
        source.mkdir()
        (source / "SKILL.md").write_text("# trusted offline fixture\n", encoding="utf-8")
        evidence = workflow / "trusted-fixture-evidence.json"
        evidence.write_text('{"source":"offline"}', encoding="utf-8")
        snapshot = build_snapshot(
            SnapshotCandidate("candidate-1", "a" * 40, source, (str(evidence),)),
            workflow / "trusted-fixture-snapshot",
        )
        packet = build_review_packet({
            "candidate_id": "candidate-1", "canonical_source": "https://github.com/example/candidate-1",
            "upstream_repository": "https://github.com/example/candidate-1", "skill_entry_path": "SKILL.md",
            "license": "MIT", "security_grade": "SA",
        }, snapshot)
        request = replace(request, review_packets={"candidate-1": packet})
        authority = workflow / "ledger" / "Skills主台账.xlsx"
        before = sha256(authority.read_bytes()).hexdigest()
        code = cli.run_interactive_protocol(coordinator, request, input_stream=input_stream, output_stream=output)
        self.assertEqual(code, 2)
        self.assertEqual(sha256(authority.read_bytes()).hexdigest(), before)
        self.assertEqual([frame["type"] for frame in output.frames], ["review_required", "run_failed"])
        self.assertFalse(any((workflow / ".runtime" / "staging").iterdir()))

    def test_protocol_eof_at_visual_gate_preserves_authority_and_cleans_runtime(self):
        workflow, coordinator, request, input_stream, output = self._protocol_fixture(eof_at="word_visual_review_required")
        authority = workflow / "ledger" / "Skills主台账.xlsx"
        before = sha256(authority.read_bytes()).hexdigest()
        with patch("skill_maintainer.office._run_office", side_effect=lambda *args: self._fake_office_result(list(args))):
            code = cli.run_interactive_protocol(coordinator, request, input_stream=input_stream, output_stream=output)
        self.assertEqual(code, 1)
        self.assertEqual(sha256(authority.read_bytes()).hexdigest(), before)
        self.assertEqual(
            [frame["type"] for frame in output.frames],
            ["review_required", "word_visual_review_required", "run_failed"],
        )
        self.assertFalse(any((workflow / ".runtime" / "staging").iterdir()))
        lock = SingleWriterLock(workflow / ".runtime" / "writer.lock")
        self.assertTrue(lock.acquire())
        lock.release()

    def test_protocol_rejected_word_page_preserves_old_authority(self):
        workflow, coordinator, request, input_stream, output = self._protocol_fixture(visual_approved=False)
        authority = workflow / "ledger" / "Skills主台账.xlsx"
        before = sha256(authority.read_bytes()).hexdigest()
        with patch("skill_maintainer.office._run_office", side_effect=lambda *args: self._fake_office_result(list(args))):
            code = cli.run_interactive_protocol(coordinator, request, input_stream=input_stream, output_stream=output)
        self.assertEqual(code, 1)
        self.assertEqual(sha256(authority.read_bytes()).hexdigest(), before)
        self.assertEqual(
            [frame["type"] for frame in output.frames],
            ["review_required", "word_visual_review_required", "run_failed"],
            output.text,
        )
        self.assertFalse(any((workflow / ".runtime" / "staging").iterdir()))

    def test_public_run_commands_fail_before_prepare_when_runtime_inputs_are_missing(self):
        workflow = self._setup_project()
        output = io.StringIO()
        with redirect_stdout(output):
            code = cli.main(["run-now", "--project-root", str(self.root)])
        self.assertEqual(code, 1)
        payload = json.loads(output.getvalue().splitlines()[-1])
        self.assertEqual(payload["type"], "run_failed")
        self.assertIn("加载器", payload["error"])
        output = io.StringIO()
        with redirect_stdout(output):
            code = cli.main(["scheduled-run", "--project-root", str(self.root)])
        self.assertEqual(code, 2)
        self.assertIn("SHA-256", json.loads(output.getvalue().splitlines()[-1])["error"])
        self.assertFalse((workflow / ".runtime").exists())

    def test_independent_stage_commands_refuse_cross_process_capability_reconstruction(self):
        self._setup_project()
        for command in ("prepare", "apply-reviews", "finalize"):
            with self.subTest(command=command):
                output = io.StringIO()
                with redirect_stdout(output):
                    code = cli.main([command, "--project-root", str(self.root)])
                self.assertEqual(code, 1)
                self.assertIn("同一长驻进程", output.getvalue())

    def test_installer_fails_before_creating_environment_when_target_rules_are_missing(self):
        workflow = self.root / WORKFLOW
        shutil.copytree(self.source_workflow, workflow)
        skills_root = Path(self.temporary.name) / "Codex Skills 目录"
        command = [
            "powershell", "-NoProfile", "-File", str(workflow / "install.ps1"),
            "-ProjectRoot", str(self.root), "-PythonExe", sys.executable,
            "-CodexSkillsRoot", str(skills_root),
        ]
        result = subprocess.run(command, capture_output=True, text=True, encoding="utf-8", errors="replace", timeout=60)
        self.assertNotEqual(result.returncode, 0)
        self.assertFalse((workflow / ".venv").exists())
        self.assertFalse((workflow / "workflow-settings.toml").exists())
        self.assertFalse(skills_root.exists())

    def test_installer_rejects_python_with_reparse_ancestor_before_any_state_write(self):
        self._authority_fixture()
        workflow = self.root / WORKFLOW
        shutil.copytree(self.source_workflow, workflow)
        linked_python_parent = Path(self.temporary.name) / "linked-python-parent"
        linked_python_parent.symlink_to(Path(sys.executable).parent, target_is_directory=True)
        skills_root = Path(self.temporary.name) / "Codex Skills 目录"
        result = subprocess.run(
            [
                "powershell", "-NoProfile", "-File", str(workflow / "install.ps1"),
                "-ProjectRoot", str(self.root),
                "-PythonExe", str(linked_python_parent / Path(sys.executable).name),
                "-CodexSkillsRoot", str(skills_root),
            ],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=60,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertFalse((workflow / ".venv").exists())
        self.assertFalse((workflow / "workflow-settings.toml").exists())
        self.assertFalse(skills_root.exists())

    def test_installer_rejects_skills_root_reparse_ancestor_before_any_state_write(self):
        self._authority_fixture()
        workflow = self.root / WORKFLOW
        shutil.copytree(self.source_workflow, workflow)
        outside = Path(self.temporary.name) / "outside-skills"
        outside.mkdir()
        linked = Path(self.temporary.name) / "linked-skills"
        linked.symlink_to(outside, target_is_directory=True)
        result = subprocess.run(
            [
                "powershell", "-NoProfile", "-File", str(workflow / "install.ps1"),
                "-ProjectRoot", str(self.root), "-PythonExe", sys.executable,
                "-CodexSkillsRoot", str(linked / "missing" / "nested"),
            ],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=180,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertFalse((workflow / ".venv").exists())
        self.assertFalse((workflow / "workflow-settings.toml").exists())
        self.assertEqual(list(outside.iterdir()), [])

    def test_installer_runs_offline_in_chinese_space_path_and_is_idempotent(self):
        self._authority_fixture()
        workflow = self.root / WORKFLOW
        shutil.copytree(self.source_workflow, workflow)
        skills_root = Path(self.temporary.name) / "Codex Skills 目录"
        python = self._bundled_python()
        installer = workflow / "install.ps1"
        environment = os.environ.copy()
        environment["PIP_NO_INDEX"] = "1"
        poison = Path(self.temporary.name) / "poison-pythonpath"
        poison_package = poison / "skill_maintainer"
        poison_package.mkdir(parents=True)
        (poison_package / "__init__.py").write_text("ORIGIN = 'poison'\n", encoding="utf-8")
        pythonpath_marker = Path(self.temporary.name) / "external-pythonpath-executed.txt"
        (poison / "sitecustomize.py").write_text(
            "from pathlib import Path\n"
            f"Path({str(pythonpath_marker)!r}).write_text('executed', encoding='utf-8')\n",
            encoding="utf-8",
        )
        inherited_pythonpath = environment.get("PYTHONPATH")
        environment["PYTHONPATH"] = str(poison) + (os.pathsep + inherited_pythonpath if inherited_pythonpath else "")
        command = [
            "powershell", "-NoProfile", "-File", str(installer),
            "-ProjectRoot", str(self.root), "-PythonExe", str(python),
            "-CodexSkillsRoot", str(skills_root),
        ]
        first = subprocess.run(command, capture_output=True, text=True, encoding="utf-8", errors="replace", env=environment, timeout=180)
        self.assertEqual(first.returncode, 0, first.stdout + first.stderr)
        self.assertFalse(pythonpath_marker.exists(), "installer Python inherited and executed external PYTHONPATH")
        settings_path = workflow / "workflow-settings.toml"
        settings = load_settings(settings_path)
        self.assertFalse(settings.workflow.enabled)
        self.assertEqual(settings.schedule.mode, "manual")
        venv_python = workflow / ".venv" / "Scripts" / "python.exe"
        self.assertTrue(venv_python.is_file())
        imported = subprocess.run(
            [
                str(venv_python), "-c",
                "from pathlib import Path; import skill_maintainer,sys; "
                "raise SystemExit(0 if Path(skill_maintainer.__file__).resolve().is_relative_to(Path(sys.argv[1]).resolve()) else 1)",
                str(workflow / "src"),
            ],
            capture_output=True,
            env={key: value for key, value in environment.items() if key != "PYTHONPATH"},
            timeout=30,
        )
        self.assertEqual(imported.returncode, 0, imported.stderr)
        self.assertTrue(
            (workflow / ".venv" / "Scripts" / "skill-maintainer.exe").is_file()
            or (workflow / ".venv" / "Scripts" / "skill-maintainer.cmd").is_file()
        )
        fallback_link = workflow / ".venv" / "Lib" / "site-packages" / "university_skill_library_maintainer.pth"
        if fallback_link.exists():
            self.assertTrue(fallback_link.read_bytes().isascii())
        self.assertTrue((skills_root / "university-skill-library-maintainer" / "SKILL.md").is_file())
        self.assertFalse(any("automation" in path.name.casefold() for path in workflow.rglob("*") if path.name != "automation-prompt.md"))

        custom = settings_path.read_text(encoding="utf-8").replace("include_generic_skills = false", "include_generic_skills = true")
        settings_path.write_text(custom, encoding="utf-8")
        second = subprocess.run(command, capture_output=True, text=True, encoding="utf-8", errors="replace", env=environment, timeout=180)
        self.assertEqual(second.returncode, 0, second.stdout + second.stderr)
        self.assertEqual(settings_path.read_text(encoding="utf-8"), custom)

    def test_installer_uses_isolated_python_for_simulated_user_site_attack(self):
        self._authority_fixture()
        workflow = self.root / WORKFLOW
        shutil.copytree(self.source_workflow, workflow)
        skills_root = Path(self.temporary.name) / "Codex Skills 目录"
        marker = Path(self.temporary.name) / "simulated-user-site-executed.txt"
        wrapper = Path(self.temporary.name) / "python-user-site-wrapper.cmd"
        bundled_python = self._bundled_python()
        wrapper.write_text(
            "@echo off\r\n"
            f'if /I not "%~1"=="-I" "{sys.executable}" -c "open(r\'{marker}\',\'a\').write(\'executed\')"\r\n'
            f'"{bundled_python}" %*\r\n',
            encoding="ascii",
        )
        environment = os.environ.copy()
        environment["PIP_NO_INDEX"] = "1"
        result = subprocess.run(
            [
                "powershell", "-NoProfile", "-File", str(workflow / "install.ps1"),
                "-ProjectRoot", str(self.root), "-PythonExe", str(wrapper),
                "-CodexSkillsRoot", str(skills_root),
            ],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            env=environment,
            timeout=180,
        )
        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        self.assertFalse(marker.exists(), "an install-time Python process was not isolated")
        venv_python = workflow / ".venv" / "Scripts" / "python.exe"
        isolated_import = subprocess.run(
            [
                str(venv_python), "-I", "-c",
                "from pathlib import Path; import skill_maintainer,sys; "
                "raise SystemExit(0 if Path(skill_maintainer.__file__).resolve().is_relative_to(Path(sys.argv[1]).resolve()) else 1)",
                str(workflow / "src"),
            ],
            capture_output=True,
            env={key: value for key, value in environment.items() if key != "PYTHONPATH"},
            timeout=30,
        )
        self.assertEqual(isolated_import.returncode, 0, isolated_import.stderr)
        launcher = workflow / ".venv" / "Scripts" / "skill-maintainer.cmd"
        self.assertTrue(launcher.is_file())
        self.assertIn(" -I -m skill_maintainer.cli", launcher.read_text(encoding="ascii"))

    def test_runtime_dependency_pins_are_complete_and_match_project_metadata(self):
        expected = (
            "openpyxl==3.1.5",
            "python-docx==1.2.0",
            "et-xmlfile==2.0.0",
            "lxml==6.1.1",
            "typing-extensions==4.16.0",
        )
        requirements = tuple(
            line.strip()
            for line in (self.source_workflow / "requirements.txt").read_text(encoding="utf-8").splitlines()
            if line.strip() and not line.lstrip().startswith("#")
        )
        metadata = tomllib.loads((self.source_workflow / "pyproject.toml").read_text(encoding="utf-8"))
        self.assertEqual(requirements, expected)
        self.assertEqual(tuple(metadata["project"]["dependencies"]), expected)

    @unittest.skipUnless(importlib.util.find_spec("wheel") is None, "requires the offline editable fallback")
    def test_installer_fallback_refuses_non_owned_launcher_without_writing_link(self):
        self._authority_fixture()
        workflow = self.root / WORKFLOW
        shutil.copytree(self.source_workflow, workflow)
        subprocess.run(
            [sys.executable, "-m", "venv", "--system-site-packages", str(workflow / ".venv")],
            check=True,
            capture_output=True,
            timeout=60,
        )
        launcher = workflow / ".venv" / "Scripts" / "skill-maintainer.cmd"
        launcher.write_text("user-owned launcher", encoding="utf-8")
        environment = os.environ.copy()
        environment["PIP_NO_INDEX"] = "1"
        result = subprocess.run(
            [
                "powershell", "-NoProfile", "-File", str(workflow / "install.ps1"),
                "-ProjectRoot", str(self.root), "-PythonExe", sys.executable,
                "-CodexSkillsRoot", str(Path(self.temporary.name) / "Codex Skills 目录"),
            ],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            env=environment,
            timeout=180,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertEqual(launcher.read_text(encoding="utf-8"), "user-owned launcher")
        self.assertFalse(
            (workflow / ".venv" / "Lib" / "site-packages" / "university_skill_library_maintainer.pth").exists()
        )


if __name__ == "__main__":
    unittest.main()
